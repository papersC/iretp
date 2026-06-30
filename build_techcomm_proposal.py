"""
Generate the full Technical & Commercial Proposal for DLD-IRETP-2026-001.

Writes one .docx that mirrors the depth of the reference layout:
17 numbered sections + annexes + appendices + functional compliance matrix.

Pricing is drawn verbatim from IRETP_Internal_Costs.pptx.
Vendor name is left as the placeholder [Vendor Legal Entity Name] so the
contracting entity can be inserted at submission time.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r"C:\Users\kalmi\IRETP\DLD-IRETP-2026-TechnicalCommercialProposal.docx"

# ---------- palette + helpers ----------
COL_INK   = RGBColor(0x0F, 0x1B, 0x2D)
COL_DLD   = RGBColor(0x06, 0x67, 0x35)   # DLD green
COL_TEXT  = RGBColor(0x1A, 0x1A, 0x1A)
COL_MUTED = RGBColor(0x6A, 0x72, 0x80)
COL_RULE  = RGBColor(0xD9, 0xDD, 0xD7)
COL_HEAD  = RGBColor(0xFF, 0xFF, 0xFF)
COL_BAND  = RGBColor(0xF1, 0xF5, 0xF1)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D9DDD7", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)

def add_para(doc, text, bold=False, italic=False, size=10.5, color=COL_TEXT, align=None, space_after=4, space_before=0, font="Calibri"):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def add_heading(doc, text, level=1):
    # Custom heading rendering for tighter control
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: COL_DLD, 2: COL_INK, 3: COL_INK}
    space_before = {1: 14, 2: 10, 3: 6}
    space_after = {1: 6, 2: 4, 3: 3}
    size = sizes.get(level, 11)
    color = colors.get(level, COL_INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before.get(level, 8))
    p.paragraph_format.space_after = Pt(space_after.get(level, 4))
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = True
    # Mark as heading style for outline/TOC compatibility
    p.style = doc.styles[f'Heading {min(level,3)}']
    for r in p.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.bold = True
    return p

def add_bullets(doc, items, indent_cm=0.5, size=10.5):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0] if p.runs else p.add_run('')
        # Re-add text properly
        p.text = ''
        run = p.add_run(it)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.color.rgb = COL_TEXT
    return None

def add_table(doc, headers, rows, col_widths_in=None, header_fill="066735", band=True, font_size=9.5):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for c in table.columns[i].cells:
                c.width = Inches(w)
    # header
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        c.text = ''
        shade(c, header_fill)
        set_cell_borders(c, color="066735")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h)
        r.font.name = 'Calibri'; r.font.size = Pt(font_size+0.5); r.bold = True
        r.font.color.rgb = COL_HEAD
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # body
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[1+i].cells[j]
            c.text = ''
            if band and i % 2 == 1:
                shade(c, "F1F5F1")
            set_cell_borders(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(str(val))
            r.font.name = 'Calibri'; r.font.size = Pt(font_size)
            r.font.color.rgb = COL_TEXT
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '066735')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---------- build the document ----------
doc = Document()

# Page setup (A4, narrow margins)
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# Heading defaults
for lvl, size, color in [(1, 18, COL_DLD), (2, 14, COL_INK), (3, 12, COL_INK)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.color.rgb = color

# =====================================================================
# COVER
# =====================================================================
for _ in range(2):
    doc.add_paragraph()
add_para(doc, "TECHNICAL & COMMERCIAL PROPOSAL", bold=True, size=24, color=COL_DLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc, "In Response to RFP No. DLD-IRETP-2026-001", size=14, color=COL_INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "Integrated Real Estate Transparency Platform (IRETP)", italic=True, size=14, color=COL_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_para(doc, "Submitted to:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Dubai Land Department", size=14, color=COL_DLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Government of Dubai", size=11, color=COL_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para(doc, "Submitted by:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "[Vendor Legal Entity Name]", size=14, color=COL_INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "UAE Registered | Dubai, United Arab Emirates", italic=True, size=11, color=COL_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_para(doc, "Submission Date: 18 May 2026", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Proposal Validity: 120 days from submission date", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Document Version: 1.0", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para(doc, "Classification: CONFIDENTIAL", bold=True, size=11, color=COL_DLD, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break(doc)

# =====================================================================
# CONFIDENTIALITY
# =====================================================================
add_heading(doc, "Confidentiality Statement", 1)
add_para(doc,
    "This proposal and all information contained herein are submitted in strict confidence by "
    "[Vendor Legal Entity Name] in response to RFP No. DLD-IRETP-2026-001 issued by the Dubai Land Department. "
    "The contents of this document constitute confidential and proprietary information and are intended solely "
    "for the use of the Dubai Land Department's evaluation team. Disclosure of any part of this document to any "
    "third party without the prior written consent of [Vendor Legal Entity Name] is strictly prohibited.")
add_para(doc,
    "[Vendor Legal Entity Name] acknowledges and accepts the confidentiality obligations set out in Section 4.3 "
    "of the RFP. All information provided by DLD in connection with this procurement shall be treated as strictly "
    "confidential and shall not be disclosed or used for any purpose other than the preparation and submission of "
    "this proposal.")

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================
add_heading(doc, "Table of Contents", 1)
toc = [
    "Executive Summary",
    "1. Understanding of the Requirement",
    "2. Why This Proposal",
    "3. Proposed Solution Architecture",
    "4. Real Estate AI Agent — Architecture & Design",
    "5. Internal Platform — EWRS & Developer Performance",
    "6. Phased Delivery Plan",
    "7. Security & DESC Compliance",
    "8. Visualisation Strategy — Public and Internal",
    "9. Data Understanding & Accuracy Methodology",
    "10. Data Integration Reference Architecture",
    "11. Project Team & Qualifications",
    "12. Cost Structure",
    "13. Warranty & Technical Support — 12 Months Post Go-Live",
    "14. AI Agent Knowledge Update Approach",
    "15. Risk Register",
    "16. Knowledge Transfer & Exit Plan",
    "17. Beyond the Brief — GRETI Score Maximisation",
    "Annex A — System Architecture Diagram",
    "Annex B — Phased Delivery Gantt (90 Calendar Days)",
    "Annex C — Functional Requirements Compliance Matrix",
    "Appendix A — Working System Evidence",
    "Appendix B — Mandatory Proposal Contents Compliance",
]
for entry in toc:
    add_para(doc, entry, size=11, space_after=3)

page_break(doc)

# =====================================================================
# EXECUTIVE SUMMARY
# =====================================================================
add_heading(doc, "Executive Summary", 1)
add_para(doc,
    "[Vendor Legal Entity Name] is pleased to present this comprehensive proposal to the Dubai Land Department "
    "(DLD) for the design, development, deployment, and post-go-live support of the Integrated Real Estate "
    "Transparency Platform (IRETP). We have read RFP No. DLD-IRETP-2026-001 v1.3 in full — every functional, "
    "non-functional, hosting, security, and compliance requirement — and we submit this proposal in the belief "
    "that [Vendor Legal Entity Name] is uniquely positioned to deliver this platform on time, to specification, "
    "and to a standard that will measurably advance Dubai's standing on the JLL Global Real Estate Transparency "
    "Index (GRETI).")
add_para(doc,
    "IRETP is not a static portal. It is a mission-critical, multi-stakeholder data platform that must "
    "simultaneously serve global investors who demand institutional-grade transparency, DLD leadership who "
    "require proactive risk intelligence, and the general public who expect a modern, responsive, multilingual "
    "experience. The stakes — Dubai's international competitiveness as a real estate destination — could not "
    "be higher.")
add_para(doc,
    "Our response is grounded in a working reference implementation already running today against the full "
    "RFP scope. The platform is a Clean-Architecture .NET 9 solution: Blazor Server portal, public WebAPI, "
    "internal AdminAPI, Hangfire orchestration, EF Core on SQL Server, Microsoft Fabric / OneLake as the "
    "analytical source of truth, MapLibre for GIS, and an Arabic-first AI agent with a UAE-resident tiered "
    "LLM orchestrator. The build is verified by an automated test suite that currently passes at 88/88 with "
    "zero warnings, and the solution exposes 18 public pages and 10 admin pages that respond 200 in production "
    "build today. What DLD will receive is the productionisation and full-scope completion of a working "
    "foundation — not a build from scratch.", space_after=8)

add_heading(doc, "Headline Commitments", 3)
add_table(doc,
    ["Commitment", "Target"],
    [
        ["Functional coverage of RFP mandatory requirements", "100% at Go-Live"],
        ["Bilingual EN/AR parity", "Full UI, Open Data, AI Agent, exports"],
        ["Data residency", "100% UAE (Azure UAE North primary, UAE Central DR)"],
        ["Public-tier availability SLA", "99.9% measured quarterly"],
        ["Internal-tier availability SLA", "99.95% measured quarterly"],
        ["AI Agent P95 response latency", "< 3.5 seconds (text), < 8 s (chart)"],
        ["KPI dashboard freshness", "≤ 15 minutes (cached, signalled to UI)"],
        ["EWRS alert-to-notification (Level 1)", "< 5 minutes — auto-escalation L1→L4"],
        ["Open Data API availability", "99.9% with partitioned rate limits"],
        ["Final Go-Live", "Day 90 from contract signing"],
        ["Source code & IP ownership", "100% transferred to DLD on Day 1 of delivery"],
        ["5-year Total Cost of Ownership", "AED ≈ 5.7M (excluding VAT, all-in)"],
    ],
    col_widths_in=[3.0, 4.0])

add_heading(doc, "Why This Proposal", 3)
add_bullets(doc, [
    "De-risked delivery. The reference build already implements the External Portal, the AI Agent with RAG and "
    "guardrails, the Open Data API, the EWRS 10-indicator risk engine, the Developer Scorecard, escrow "
    "monitoring, and the GRETI beyond-brief modules. The 90-day window is spent on DLD data integration, "
    "UAT, VAPT, and language extension — not on building the platform from first principles.",
    "Microsoft-native and Fabric-aligned. The stack is .NET 9, EF Core on SQL Server, Microsoft Fabric / "
    "OneLake (Gold layer as single source of analytical truth), Azure Front Door + Key Vault + App Insights "
    "+ Log Analytics, Bicep IaC pinned to UAE North. There is no proprietary platform layer between DLD's "
    "data and DLD's portal.",
    "Source ownership from Day 1. On contract signature, the complete .NET source, EF Core migrations, "
    "Bicep IaC, Hangfire job definitions, and the five technical documentation pillars are transferred to "
    "DLD's GitHub organisation. No platform licence is retained by any third party. There is no per-user "
    "fee at any scale.",
    "AI governance baked in. The orchestrator routes each query across four tiers (nav / analytics / "
    "primary / secondary), every inference carries UAE-residency metadata, every response is grounded in "
    "RAG over OneLake Gold with deterministic statistical analysis appended (anomaly, trend, correlation), "
    "and a two-layer non-advisory guardrail (keyword catalogue + regex pattern) is enforced at system-prompt "
    "level and validated by a 14-question accuracy harness (EN+AR) that extends to DLD's 100-question UAT "
    "catalogue.",
    "Compliance posture, not promises. Audit-log immutability is enforced at the DbContext layer; security "
    "CI runs vulnerability scans, OWASP ZAP baseline, and Bicep linting on every commit; SLO health checks "
    "publish AI P95, UAE-residency state, and KPI freshness against §10.1 thresholds. DESC ISR v3 alignment "
    "is structural, not retrofitted.",
    "GRETI extras already shipping. Beneficial Ownership disclosure (with Arabic owner names), Mortgage & "
    "Debt Market Transparency (5 KPIs including average mortgage size), and ESG coverage heatmap on the "
    "interactive map are live in the reference build — not roadmap items.",
])

add_heading(doc, "Investment Snapshot", 3)
add_table(doc,
    ["Cost category", "Amount (AED)", "Type"],
    [
        ["Infrastructure CAPEX (Azure UAE North, recommended)", "335,500", "One-time"],
        ["Software Licences (Year 1)", "235,200", "One-time"],
        ["Application Development (M1–M4, 1,095 person-days)", "2,007,000", "One-time"],
        ["Total one-time CAPEX", "2,577,700", "—"],
        ["Annual OPEX (Year 1, includes 12-month warranty)", "629,200", "Recurring"],
        ["Indicative 5-Year Total Cost of Ownership", "≈ 5,723,500", "Reference only"],
    ],
    col_widths_in=[3.5, 1.7, 1.8])
add_para(doc, "All figures are AED, exclusive of VAT. Detailed breakdowns are in Section 12.",
    italic=True, size=9, color=COL_MUTED)

page_break(doc)

# =====================================================================
# SECTION 1 — UNDERSTANDING
# =====================================================================
add_heading(doc, "1. Understanding of the Requirement", 1)

add_heading(doc, "1.1 Strategic Context", 2)
add_para(doc,
    "DLD's objective extends well beyond building a website. IRETP is a strategic instrument designed to "
    "advance Dubai's position on the JLL GRETI — a biennial index that directly influences where over 80% of "
    "global direct commercial real estate investment flows. Every functional and technical decision in this "
    "proposal has been made with that strategic objective in mind.")
add_para(doc,
    "We have mapped every identified GRETI gap against the platform components we are proposing. We are "
    "confident that a correct implementation of IRETP will create measurable, auditable improvements across "
    "all six GRETI sub-indices within the first assessment cycle following Go-Live.", space_after=6)

add_table(doc,
    ["GRETI Sub-Index", "Identified Gap", "IRETP Component Addressing It"],
    [
        ["Performance Measurement",
         "Inconsistent quarterly KPIs, no public price/rental index, no developer accountability data.",
         "Public Price Index + Rental Index + Yield Calculator (Section 3.5); Public Developer Scorecards (Section 5.2); KPI dashboard with 15-minute cache (Section 3.1)."],
        ["Market Fundamentals",
         "Property-type breakdowns and transaction-type splits not surfaced publicly.",
         "External Portal: Transactions page (FR-007), Property Distribution, Top-Zones bar, Monthly Volume chart (FR-004)."],
        ["Governance",
         "Beneficial ownership opacity; absence of public developer rating.",
         "Beneficial Ownership module (Arabic owner names rendered RTL); Public Developer Scorecard with on-time delivery %, RERA compliance, units delivered (already shipping)."],
        ["Sustainability",
         "ESG indicators not surfaced; no LEED/Estidama Pearl/BREEAM mapping.",
         "ESG module with rubric explainer (LEED ↔ Estidama Pearl ↔ BREEAM ↔ investor signal); ESG coverage as 4th heatmap mode on the GIS map."],
        ["Regulatory & Legal Framework",
         "Process for escrow, AML, and dispute resolution not transparent to investors.",
         "Escrow Monitoring page; Beneficial Ownership module; CMS-served regulatory content (Section 3.1)."],
        ["Transaction Process",
         "Real-time pricing, debt market intelligence, and yield analytics not available.",
         "Live Price/Rental indices; Mortgage & Debt Market Transparency (5 KPIs including average mortgage size, mortgage-to-transaction ratio, LTV distribution)."],
    ],
    col_widths_in=[1.8, 2.4, 2.8])

add_heading(doc, "1.2 Platform Architecture Summary", 2)
add_para(doc,
    "IRETP comprises two integrated pillars, delivered across four phases within a 90 calendar day window "
    "from contract signing. We have read and understood every section of RFP No. DLD-IRETP-2026-001 v1.3 in "
    "full, including all 55 deliverables in Section 13, all 10 UAT test categories in Section 17.3, all 14 "
    "cost-table line items in Section 14, and all architectural requirements in Section 11.", space_after=6)
add_table(doc,
    ["Platform Pillar", "Description"],
    [
        ["External Public Portal",
         "Public-facing investor portal: Market Overview Dashboard, Transactions, Interactive GIS Map, Price/Rental Indices, Slice-and-Dice Analytics, Real Estate AI Agent, Investor Watchlist & Notifications, Open Data API, ESG Module, Public Developer Scorecard."],
        ["Internal Management Platform",
         "DLD-staff platform: Early Warning Risk System (EWRS) with L1–L4 auto-escalation, Escrow Monitoring, Developer Performance & Rating, AI Model Performance Dashboard, Audit Logs, CMS, AI orchestration admin, RBAC for six internal roles."],
        ["Open Data Layer",
         "Public Open Data API, OpenAPI 3.0 dual specifications (public-filtered + full-v1), interactive console at /api-docs, per-key rate limits via ApiKey middleware."],
    ],
    col_widths_in=[2.0, 5.0])

add_heading(doc, "1.3 Critical Constraints We Have Fully Absorbed", 2)
add_para(doc, "We draw attention to the following constraints, which are non-negotiable and which our proposed solution has been explicitly designed to satisfy:", space_after=4)
add_bullets(doc, [
    "Pre-Publication Analytics Assessment gate. No data is published publicly until the mandatory assessment "
    "is completed and signed off in writing by the DLD Data Liaison Officer. This is hard-coded into our Phase 1 "
    "delivery plan as a go/no-go gate (Days 32–37).",
    "AI non-advisory constraint. The AI Agent is prohibited from providing investment advice, recommending "
    "specific properties, or making price forecasts. The constraint is enforced at three layers: system-prompt "
    "(orchestrator), runtime input filter (keyword + regex catalogue), and post-generation validation (response "
    "scan with bilingual refusal copy). All three are exercised by the accuracy harness as mandatory UAT cases.",
    "Microsoft Fabric / OneLake. DLD's existing data architecture is the foundation of our data layer, not an "
    "afterthought. We do not duplicate data outside the Fabric ecosystem. All platform visualisations and APIs "
    "are served from the OneLake Gold layer.",
    "UAE data residency. All AI inference, data storage, processing, backups, and DR replicas remain within UAE "
    "sovereign cloud regions regardless of which AI model is active. UAE-residency metadata is stamped on every "
    "AI snapshot for audit.",
    "DESC ISR v3. The security plan is built to this standard from Day 1, not retrofitted after delivery. "
    "Audit-log immutability is enforced at the DbContext layer; central audit-trail rows are written on every "
    "scoring-weight or risk-threshold change.",
])

page_break(doc)

# =====================================================================
# SECTION 2 — WHY THIS PROPOSAL
# =====================================================================
add_heading(doc, "2. Why This Proposal", 1)

add_heading(doc, "2.1 A Working Reference Implementation, Not a Concept", 2)
add_para(doc,
    "[Vendor Legal Entity Name] does not approach this RFP with a concept paper or a design file. We have "
    "already built a functional reference implementation that covers approximately 88% of the RFP scope today. "
    "The platform builds clean (0 warnings, 0 errors), passes 88 of 88 xUnit tests, and renders 18 public "
    "pages and 10 admin pages with HTTP 200 against a real SQL Server backend. What DLD will be commissioning "
    "is the productionisation, the DLD-specific data integration, and the formal compliance certification of "
    "a working foundation.")
add_para(doc, "What is already built and working in the reference codebase:", space_after=4)
add_bullets(doc, [
    "External Public Portal — Market Overview Dashboard with KPI cards (15-minute cache), Transactions page "
    "with URL-persisted filters (FR-007), Interactive GIS Map with four heatmap modes (volume / PSF / yield / "
    "ESG) and 183 individual project pins colour-coded by status (FR-011), Price Index, Rental Index with "
    "live yield calculator (FR-014), Slice-and-Dice Analytics, Zone Comparison up to five zones, ESG dashboard, "
    "Public Developer Scorecard.",
    "Real Estate AI Agent — tier-routed orchestrator (nav / analytics / primary / secondary), RAG against "
    "OneLake Gold, deterministic statistical analysis appended to context, two-layer non-advisory guardrail "
    "(keyword + regex), 14-question accuracy harness in EN+AR, cross-session memory with consent (AI-006), "
    "Arabic NLP with explicit Gulf dialect handling, citation chips and retry on the chat surface.",
    "Investor surfaces — Watchlist v2, Notifications with quiet hours, six alert types, RFC 8058 List-Unsubscribe "
    "header on emails, HMAC-signed unsubscribe tokens (§6.2), Account preferences with PDPL Article 19.2 data "
    "erasure and consent toggles, MFA enrolment.",
    "Open Data layer — OpenAPI 3.0 dual specifications (public-filtered + full-v1), interactive console at "
    "/api-docs, ApiKey middleware with per-key rate limits, IP fallback for anonymous, API key portal page.",
    "Early Warning Risk System — 10-indicator engine (project delivery delay, escrow shortfall, construction "
    "suspension, transaction-volume decline, price decline, developer-score deterioration, high-risk concentration, "
    "severe regulatory violation, manual flag, watchlist trigger), L1–L4 SLA-driven auto-escalation via Hangfire, "
    "central audit log on threshold edits, EWRS dashboard with SLA badges.",
    "Developer Performance & Rating — composite scoring with configurable weights, on-time delivery percentage, "
    "RERA compliance, total units delivered, quarter-over-quarter trend, central audit trail on weight changes, "
    "public scorecard surface.",
    "Escrow Monitoring — admin dashboard, balance and outflow checks, EWRS hook for shortfall thresholds.",
    "GRETI beyond-brief modules — Beneficial Ownership (English + Arabic owner names rendered RTL), Mortgage & "
    "Debt Market Transparency (5 KPIs including average mortgage size), ESG coverage map layer.",
    "Operations — 9 Hangfire jobs (KPI refresh, EWRS escalation, notification delivery, accuracy harness, "
    "knowledge index refresh, etc.), three health endpoints (/healthz/live, /healthz/ready, /healthz/sla), "
    "five-pillar technical documentation under /docs (ARCHITECTURE, API_REFERENCE, RUNBOOK, DISASTER_RECOVERY, "
    "DATA_DICTIONARY), Bicep IaC pinned to UAE North, Docker compose hardened with cap_drop ALL and "
    "no-new-privileges.",
])
add_para(doc,
    "DLD can review the working system at any point during the evaluation. The phasing, team structure, and "
    "Day-1 VAPT scheduling in this proposal are informed by what has already been built, tested, and "
    "documented — not by what we hope to build.", italic=True)

add_heading(doc, "2.2 Clean Architecture on a Microsoft-Native Stack", 2)
add_para(doc,
    "The platform is architected as a Clean-Architecture .NET 9 solution. Outer layers depend inward; the "
    "Domain layer has zero framework references. This is the architecture the Microsoft developer ecosystem "
    "treats as the long-term default for enterprise-grade business systems, and it is the architecture that "
    "delivers the longest maintainable lifetime for a public-sector platform with a 10-year horizon.", space_after=6)
add_table(doc,
    ["Project", "Responsibility"],
    [
        ["IRETP.Domain", "Framework-free entities, value objects, enums, domain helpers. Depends on nothing."],
        ["IRETP.Application", "MediatR commands, queries, DTOs, feature-level interfaces. No infrastructure or UI concerns."],
        ["IRETP.Infrastructure", "EF Core IretpDbContext, migrations, repository implementations, ASP.NET Identity, Hangfire job services, AI orchestration, notification senders."],
        ["IRETP.WebAPI", "Public read/write HTTP surface. Hosts SignalR for real-time alerts and the Hangfire dashboard."],
        ["IRETP.AdminAPI", "DLD-internal endpoints (EWRS threshold editing, scoring-weight configuration, AI model status, audit). Protected by the policy registry in IRETP.Infrastructure/Identity/AuthorizationPolicies.cs."],
        ["IRETP.Web", "Blazor Server frontend. Public portal + investor pages + admin pages. Routes WebAPI + AdminAPI calls through typed HttpClient services."],
        ["IRETP.Tests", "xUnit coverage for domain rules, application handlers, and infrastructure services. 88 tests passing today."],
    ],
    col_widths_in=[2.0, 5.0])

add_heading(doc, "2.3 Source Ownership From Day One", 2)
add_para(doc,
    "Every line of code, every EF Core migration, every Bicep template, every Hangfire job definition, every "
    "test, and every architectural document becomes the exclusive intellectual property of the Dubai Land "
    "Department on the date of contract signature — not on Day 90 of delivery, not at the expiry of any "
    "warranty period.")
add_para(doc,
    "There is no proprietary platform layer hidden inside the stack. The components are all standard "
    "Microsoft and open-source technologies — ASP.NET Core, Blazor Server, EF Core, SQL Server, Hangfire, "
    "SignalR, MapLibre, Chart.js, OpenAPI. Any qualified .NET team in the UAE can maintain this codebase "
    "without vendor involvement. DLD is not paying a platform subscription, a per-user licence, a per-tenant "
    "fee, or a knowledge-base index licence at any point during the platform's operational lifetime.")
add_para(doc,
    "This commercial structure is materially different from a SaaS platform engagement. Over a five-year "
    "horizon, the source-owned model converts what would otherwise be an ongoing platform subscription into "
    "a one-time CAPEX paid against verified UAT acceptance. The five-year TCO advantage is detailed in "
    "Section 12.7.")

add_heading(doc, "2.4 De-Risked Delivery Pathway", 2)
add_para(doc,
    "Three structural advantages reduce delivery risk on a 90-day timeline:", space_after=4)
add_bullets(doc, [
    "The reference implementation removes the construction risk on every Phase 1 feature. Phase 1 of the "
    "delivery plan is not building the External Portal; it is integrating the working External Portal with "
    "DLD's OneLake Gold layer, running the Pre-Publication Analytics Assessment, completing UAT, and going live.",
    "Microsoft Fabric / OneLake is consumed via the standard Microsoft Fabric SQL endpoint and semantic model "
    "APIs. There is no proprietary connector to license, configure, or wait for. The data layer is a SQL "
    "connection string and a credential — both are configured on Day 1.",
    "The 9 Hangfire jobs that drive KPI refresh, EWRS escalation, notification delivery, knowledge re-indexing, "
    "and the accuracy harness are already operating in the reference build. The Phase 1 work on these is "
    "schedule configuration, not implementation.",
])

add_heading(doc, "2.5 UAE Government Operating Posture", 2)
add_para(doc,
    "All AI inference, data storage, telemetry, backups, and DR replicas are pinned to UAE sovereign cloud "
    "regions. The Bicep IaC under /infra/bicep/ pins every resource to Azure UAE North; the DR plan replicates "
    "to a physically separate UAE region. UAE Pass OIDC federation is wired into the authentication stack as "
    "an opt-in scheme. MFA is mandatory for every internal DLD user via the policy registry. PDPL Article 19.2 "
    "(right to erasure) is implemented end-to-end: opting out of AI memory purges UserAiMemory rows in the "
    "same transaction.")

page_break(doc)

# =====================================================================
# SECTION 3 — SOLUTION ARCHITECTURE
# =====================================================================
add_heading(doc, "3. Proposed Solution Architecture", 1)

add_heading(doc, "3.1 Architecture Overview", 2)
add_para(doc,
    "The IRETP architecture is structured across four tiers with clearly defined responsibilities, data flows, "
    "and security boundaries. The architecture fully leverages DLD's existing Microsoft Fabric ecosystem: the "
    "External Public Portal's visualisation layer is a Blazor Server frontend consuming data from the OneLake "
    "Gold layer, consistent with Section 11.4.2 of the RFP.", space_after=4)
add_table(doc,
    ["Tier", "Responsibility"],
    [
        ["Tier 1 — Edge & Access",
         "Azure Front Door Premium (WAF, CDN, TLS 1.2 minimum), Azure DDoS Standard, geo-routing pinned to UAE North/Central. UAE Pass OIDC federation for investors; ASP.NET Identity + JWT + MFA for DLD staff."],
        ["Tier 2 — Data & Integration",
         "Microsoft Fabric / OneLake Gold layer as the analytical source of truth, consumed via the Fabric SQL endpoint. SQL Server (Azure SQL with elastic pool) for OLTP. Hangfire orchestration for scheduled refresh, EWRS escalation, accuracy harness, knowledge re-indexing."],
        ["Tier 3 — Application & AI",
         "ASP.NET Core 9 on three hosts: public WebAPI (:5000), AdminAPI (:5002), Blazor Server Web (:5010). MediatR CQRS. AIOrchestrator with four-tier model routing. SignalR for real-time alerts. Five Hangfire jobs run inside WebAPI, four inside AdminAPI."],
        ["Tier 4 — Presentation & Surfaces",
         "Public users and registered investors access the External Portal via browser. DLD staff access the Internal Platform via browser with MFA. The Open Data API surfaces machine-readable data for researchers, PropTech developers, and financial institutions."],
    ],
    col_widths_in=[1.7, 5.3])

add_heading(doc, "3.2 Microsoft Fabric / OneLake Alignment", 2)
add_para(doc,
    "DLD's existing investment in the Microsoft Fabric ecosystem is the data foundation of this platform. The "
    "IRETP platform treats the OneLake Gold layer as the single, authoritative source of truth for all "
    "analytical data. Our approach is explicitly non-duplicative: we do not introduce parallel data stores, "
    "shadow databases, or separate ETL layers outside the Fabric ecosystem.")
add_para(doc,
    "Connectivity to OneLake is achieved via the standard Fabric SQL endpoint, which exposes the Gold-layer "
    "lakehouse and warehouse tables as a SQL surface accessible by any .NET application through the standard "
    "Microsoft.Data.SqlClient driver. From the IRETP service tier, OneLake is one entry in the connection-string "
    "configuration — no proprietary connector, no platform layer, no third-party intermediary.", space_after=6)
add_table(doc,
    ["DLD Fabric Component", "IRETP Integration Approach", "What We Add"],
    [
        ["OneLake Gold layer (analytical)",
         "Consumed read-only via Fabric SQL endpoint. EF Core context maps Gold-layer views to projection types in IRETP.Application.",
         "Derived metric pipelines (price-per-sqft averages, YoY %, rental yields), KPI snapshot cache (15-min TTL), the Pre-Publication Analytics Assessment automation."],
        ["OneLake Bronze/Silver",
         "Not consumed by IRETP. The platform displays only Gold-layer data; data quality and reconciliation between Bronze/Silver and Gold remain DLD's data team responsibility.",
         "We do not add to or modify DLD's Bronze/Silver/Gold curation pipeline."],
        ["Data Factory / pipelines",
         "Out of scope for IRETP. The platform respects whatever update cadence DLD's Data Factory pipelines produce on Gold.",
         "We add a freshness probe (the /healthz/sla endpoint) that warns DLD if Gold has not refreshed within the agreed cadence."],
        ["Semantic models",
         "Power BI semantic models are reused unchanged for the Internal Platform dashboards (EWRS, Developer Performance, Escrow Monitoring).",
         "Internal Platform reports embed Power BI tiles via the standard report-server SDK."],
        ["UAE-residency tagging",
         "Every Gold-layer SQL query carries an explicit UAE-region tag in App Insights telemetry.",
         "We add SLO health checks that fail if any inference or data path leaves the UAE residency boundary."],
    ],
    col_widths_in=[1.7, 2.6, 2.7])

add_heading(doc, "3.3 External Public Portal — Visualisation Layer", 2)
add_para(doc,
    "We agree with DLD's assessment in Section 11.4.2 that Power BI is not suitable as the primary visualisation "
    "engine for the External Public Portal at the projected anonymous-user volumes. The licensing model does "
    "not scale cost-effectively to 10,000+ concurrent anonymous users; embed-token management adds operational "
    "overhead; and the UX customisation surface is constrained.")
add_para(doc, "Our External Portal visualisation strategy:", space_after=4)
add_bullets(doc, [
    "Blazor Server frontend (IRETP.Web) consuming the public WebAPI via typed HttpClient services. Server-side "
    "rendering means initial paint is fast on low-bandwidth connections and the SEO posture for the public "
    "portal is preserved.",
    "Chart.js for all interactive analytical charts (KPI cards, monthly volume line, transaction-type "
    "doughnut, top-zones bar, dual-axis PSF trend). Open-source, no per-tile licence, full RTL support.",
    "MapLibre GL JS for the interactive GIS map: four heatmap modes (volume / price-per-sqft / yield / ESG "
    "coverage), 183 individual project pins colour-coded by status (FR-011), zoom-driven radius interpolation, "
    "official Dubai zone GeoJSON overlay, popup detail on pin click.",
    "Shared Blazor components — ModuleBanner, Skeleton/SkeletonTable/SkeletonKpiGrid, NetworkError, EmptyState, "
    "BackToTop — provide a consistent loading-state and error-state experience across 29 retrofitted pages.",
    "Public Open Data API surfaces machine-readable payloads (JSON, CSV) for institutional consumers; "
    "OpenAPI 3.0 dual specifications are published at /openapi/public.json and /openapi/v1.json; the "
    "interactive console at /api-docs allows registered investors to test calls in the browser.",
])
add_para(doc,
    "Power BI is retained and recommended for all Internal Platform components — EWRS dashboards, Developer "
    "Performance dashboards, internal management reporting — where the named-user licensing model is "
    "cost-appropriate. This is a deliberate split, not a rejection of Power BI: the right tool for each "
    "audience.")

add_heading(doc, "3.4 Technology Stack Proposal", 2)
add_para(doc,
    "The stack below reflects our recommended Option A: Azure UAE North hosting. Option B (on-premises at "
    "DLD's data centre) substitutes the corresponding on-prem equivalents; both options run an identical .NET "
    "application stack with no capability difference. Section 7.4 covers the hosting decision.", space_after=4)
add_table(doc,
    ["Architectural Layer", "Proposed Technology", "Justification / NFR Alignment"],
    [
        ["Frontend / Presentation",
         "ASP.NET Core 9 + Blazor Server, MapLibre GL JS (GIS), Chart.js (analytics).",
         "Native RTL/LTR; WCAG 2.1 AA on shared components; SSR posture good for low-bandwidth public; zero charting licence cost."],
        ["Backend API Layer",
         "ASP.NET Core 9 minimal-API / controllers, MediatR CQRS, FluentValidation, OpenAPI 3.0 (Swashbuckle).",
         "P95 API response < 500 ms under 10,000 concurrent sessions; partitioned rate limiting on Open Data; standard .NET tooling."],
        ["AI / LLM Services",
         "AIOrchestrator with four tiers (nav / analytics / primary / secondary). BYO-LLM via configuration. Azure OpenAI primary, secondary fallback to alternate UAE-resident provider.",
         "Multi-model orchestration; model swap via configuration (no redeployment); UAE residency metadata on every snapshot; accuracy harness."],
        ["Database & OLTP",
         "Microsoft SQL Server (Azure SQL Database, elastic pool) for IRETP OLTP; EF Core 9 migrations.",
         "Encryption at rest (TDE), Always Encrypted for PDPL-classified columns, geo-redundant backup, point-in-time restore."],
        ["Analytical Source of Truth",
         "Microsoft Fabric OneLake Gold layer via Fabric SQL endpoint.",
         "DLD's single source of truth, non-duplicated. P95 query budget enforced by the KPI snapshot cache."],
        ["GIS / Mapping",
         "MapLibre GL JS + official Dubai zone GeoJSON.",
         "Map initial load < 5 s on 10 Mbps; four heatmap layers switch without page reload; zoom-based radius interpolation for 183 project pins."],
        ["CMS",
         "ASP.NET Core minimal CMS surface inside IRETP.AdminAPI + Blazor admin pages.",
         "Eliminates CMS licence cost; bilingual content storage; RBAC-gated publish."],
        ["Authentication & Identity",
         "ASP.NET Identity + JWT + TOTP MFA for DLD staff; UAE Pass OIDC federation for investors; refresh-token rotation; internal idle timeout 30 minutes (RFP §10.2.1).",
         "Centralised identity; MFA mandatory at policy level; refresh-token reuse detection; no dependency on external identity cloud."],
        ["Job Orchestration",
         "Hangfire (open-source) — 9 scheduled jobs (KPI cache refresh, EWRS escalation, notification delivery, accuracy harness, knowledge index refresh).",
         "Persisted job state in SQL Server; dashboard at /hangfire; durable retries; cron schedules."],
        ["Real-Time Messaging",
         "SignalR over WebSocket (HTTPS fallback) inside the public WebAPI.",
         "EWRS alert push to internal dashboards; investor notification ping; bandwidth-bounded."],
        ["Infrastructure as Code",
         "Bicep — main.bicep (subscription) + platform.bicep (platform components), all UAE North-pinned, MSI-enabled, Key Vault purge-protected.",
         "Reproducible deployments; CI-tested; delivered to DLD on Day 1."],
        ["Monitoring & Observability",
         "Azure Application Insights + Log Analytics Workspace. Health endpoints /healthz/live, /healthz/ready, /healthz/sla.",
         "P95 latency, error rate, AI residency state, KPI freshness — all measured against §10.1 thresholds; SLA probe is the single source of truth."],
    ],
    col_widths_in=[1.7, 2.5, 2.8])

add_heading(doc, "3.5 Scalability & Performance Targets", 2)
add_para(doc,
    "IRETP is architected for horizontal scale from Day 1. Every application tier runs as a stateless ASP.NET "
    "Core host — additional capacity is added by deploying more instances, not by upgrading individual hosts. "
    "Scale-out does not require re-architecture at any user volume.")
add_para(doc,
    "The platform is sized for Year-1 projections of 5,000 registered investors and 98 internal DLD staff. "
    "It is designed and tested to handle 3× that volume — 15,000 concurrent registered users, 10,000+ peak "
    "anonymous sessions, and triple the API call volume — without infrastructure change. The Azure App Service "
    "Plan auto-scales on CPU and memory pressure; DLD does not intervene.", space_after=4)
add_para(doc, "Mandatory performance thresholds (per §10.1), all of which are measured by the SLA probe:")
add_bullets(doc, [
    "Public-tier availability: 99.9% measured quarterly.",
    "Internal-tier availability: 99.95% measured quarterly.",
    "P95 API response: < 500 ms under 10,000 concurrent sessions.",
    "P90 AI text response: < 8 seconds.",
    "P95 AI chart response: < 12 seconds.",
    "KPI dashboard data freshness: ≤ 15 minutes (cache TTL, signalled to the UI via a freshness badge).",
    "EWRS Level-1 alert delivery: < 5 minutes (email/SMS/in-platform); L2 escalation if unacknowledged within 4 business hours; L3 within 1 business hour; L4 immediate.",
])
add_para(doc,
    "Load testing is a mandatory Phase 1 go-live acceptance criterion. Before the External Portal goes live, "
    "the platform must demonstrate every threshold above under simulated peak load. DLD signs off on the load "
    "test report as part of Phase 1 UAT acceptance. The same suite runs again before Phase 3 go-live to cover "
    "the expanded internal-platform load.")

page_break(doc)

# =====================================================================
# SECTION 4 — AI AGENT
# =====================================================================
add_heading(doc, "4. Real Estate AI Agent — Architecture & Design", 1)

add_heading(doc, "4.1 Our Position on AI in This Platform", 2)
add_para(doc,
    "Artificial intelligence is among the most over-used terms in technology proposals. The distinction between "
    "a credible AI proposal and a marketing document rests not on the presence of AI, but on whether each AI "
    "capability has been placed where it creates demonstrable value — and whether the temptation to apply AI "
    "indiscriminately has been resisted.")
add_para(doc,
    "We make one commitment that we consider the most important in this entire proposal: the non-advisory "
    "constraint is non-negotiable, technically enforceable, and UAT-verified. The AI Agent will retrieve, "
    "present, and analyse DLD data. It will never recommend properties, forecast prices as facts, or provide "
    "investment guidance. This is not a policy choice we might revisit — it is a hard architectural constraint "
    "enforced at three independent layers and exercised by an automated accuracy harness on every release.")

add_heading(doc, "4.2 Multi-Model Orchestration with Task-Aware Tier Routing", 2)
add_para(doc,
    "The AIOrchestrator routes each incoming query across four named tiers, allowing different model selections "
    "for different task profiles. Tier assignment is computed by a query-classifier that inspects the request "
    "text, the user context, and the conversational state; routing is auditable in every snapshot stored in the "
    "AI accuracy log.", space_after=4)
add_table(doc,
    ["Tier", "Use Case", "Default Model Posture", "Latency Budget"],
    [
        ["nav",       "DLD service navigation (\"How do I file a complaint?\")", "Lightweight, low-cost UAE-resident model", "P95 < 2 s"],
        ["analytics", "Numerical or data-heavy retrieval (\"Top 5 zones by yield in Q1?\")", "Higher-capability UAE-resident model + deterministic stats appended", "P95 < 5 s"],
        ["primary",   "Open-ended bilingual conversation, RAG over Gold layer", "Top-tier UAE-resident model (Azure OpenAI)", "P90 < 8 s"],
        ["secondary", "Fallback if primary fails", "Alternate UAE-resident provider", "P90 < 10 s"],
    ],
    col_widths_in=[0.9, 2.6, 2.5, 1.0])
add_para(doc,
    "Tier selection and model identity are configurable from the AI Model admin surface in IRETP.AdminAPI. "
    "Changing the active model on any tier is a configuration update — no redeployment, no code change, no "
    "vendor involvement. New model providers are added by adding a configuration entry plus a credential to "
    "Key Vault; the orchestrator picks them up at next session start.")

add_heading(doc, "4.3 RAG Architecture — Grounding Every Response in DLD Data", 2)
add_para(doc,
    "Retrieval-Augmented Generation is mandatory for every response the AI Agent delivers. The Agent never "
    "generates an answer from general training data alone — every factual claim is grounded in retrieved, "
    "timestamped DLD data. The RAG pipeline:", space_after=4)
add_bullets(doc, [
    "Step 1 — Query parsing. The natural-language input is classified (data query, service navigation, "
    "analytical request, comparative query) and routed to the appropriate tier.",
    "Step 2 — Retrieval. The retrieval layer queries the Gold-layer projection (Microsoft Fabric SQL endpoint) "
    "for relevant transactions, projects, developer profiles, zone boundaries, RERA regulations, and DLD "
    "service procedures. Vector retrieval is layered over keyword retrieval for hybrid recall.",
    "Step 3 — Deterministic statistics. For analytical queries, the deterministic TimeSeriesAnalyzer in "
    "IRETP.Infrastructure computes anomalies (z-score), trends (least-squares with R² and flat-threshold), and "
    "correlations (Pearson with no-variance guard). These statistics are appended to the RAG context so the LLM "
    "describes computed numbers, not approximated ones.",
    "Step 4 — Augmentation. Retrieved data + deterministic statistics are injected into the system prompt, "
    "which carries the non-advisory constraint and the citation requirement.",
    "Step 5 — Generation. The model produces a response grounded exclusively in the retrieved data. Every "
    "data point is cited with its source identifier and timestamp.",
    "Step 6 — Validation. A post-generation validation pass scans the response for investment-advisory "
    "language. Flagged responses are reformulated or declined with bilingual refusal copy.",
])

add_heading(doc, "4.4 Non-Advisory Guardrail — Three Independent Layers", 2)
add_para(doc,
    "The non-advisory constraint is enforced by three independent layers, exactly because no single layer is "
    "sufficient.", space_after=4)
add_bullets(doc, [
    "Layer 1 — System prompt. The orchestrator's system prompt explicitly forbids advisory language. The "
    "constraint is reinforced by example before every conversation begins.",
    "Layer 2 — Input filter (KeywordAdvisoryGuardrail). A bilingual keyword catalogue plus a regex catalogue "
    "blocks queries that explicitly request recommendations or forecasts (\"recommend me a property in JVC\", "
    "\"will Dubai Hills prices rise in 2027?\"). Historical framing (\"appreciated 6% YoY\") still passes. "
    "Tightened in Pass 19 of the development cycle to catch forward-looking patterns the narrow phrase list "
    "missed: \"by 20XX prices will\", \"prices are expected to rise/climb/jump/grow\", \"yields are projected\", "
    "\"forecast of AED N\".",
    "Layer 3 — Output validation. The post-generation scan inspects the response text for advisory patterns "
    "even when the input passed the filter. Triggered responses are replaced with a bilingual refusal message "
    "and logged for audit.",
])
add_para(doc,
    "All three layers are exercised by the accuracy harness on every release. The harness catalogue includes "
    "adversarial cases designed to probe the guardrail across English, Arabic, and Gulf-dialect phrasings.")

add_heading(doc, "4.5 Arabic NLP Capability", 2)
add_para(doc,
    "Phase 1 requires full conversational capability in both English and Arabic. Arabic NLP accuracy must be "
    "within 5% of English accuracy on the same standardised test set. Our approach:")
add_bullets(doc, [
    "Arabic NLP is a weighted criterion in the LLM evaluation framework run in Phase 0. Candidate models are "
    "assessed against a Gulf-Arabic real-estate test set before any model recommendation is finalised with DLD.",
    "A 100-question Arabic-language test catalogue is developed in parallel with the English catalogue, "
    "covering zone names, transaction types, developer names, RERA procedures, and comparative market queries.",
    "Arabic zone names and developer names are validated against DLD's official records. The validator stores "
    "the canonical Arabic strings and rejects mismatches at retrieval time.",
    "All AI chat surfaces render RTL when the user's active language is Arabic: text alignment, axis labels on "
    "appended charts, citation chip order, and the Expand Insight / Retry controls all mirror correctly.",
])

add_heading(doc, "4.6 Accuracy Harness — Continuous Quality Measurement", 2)
add_para(doc,
    "The platform ships with an automated accuracy harness (IAiAccuracyHarness + AiAccuracyHarness in "
    "IRETP.Infrastructure). The harness runs a seed catalogue of 14 questions on every release covering "
    "data retrieval, service navigation, adversarial guardrail probes, and bilingual variants. The catalogue "
    "is extensible — for DLD UAT, the seed is extended to the full 100-question DLD catalogue. Each run "
    "produces a structured report scored by expected-keyword match plus a manual review queue for ambiguous "
    "cases. The dashboard surfaces P95 latency, residency state, accuracy %, and tier-level breakdowns. The "
    "POST /api/admin/ai-models/accuracy-test endpoint is gated to SystemAdministrator.")

add_heading(doc, "4.7 Cross-Session Memory with Consent (AI-006)", 2)
add_para(doc,
    "AI-006 requires the AI Agent to remember a registered investor's preferences across sessions when the "
    "investor opts in. The platform implements this via a UserAiMemory entity (UserId + Kind [zone|topic] + "
    "Key + Frequency + LastUsedAt) with explicit consent capture at registration and at any point thereafter. "
    "When the investor toggles AI memory off, every UserAiMemory row for the user is deleted in the same "
    "transaction (PDPL Article 19.2).")
add_para(doc,
    "Cross-session context is prepended only when consent is granted: the user's top 5 zones and top 3 topics "
    "are injected into the RAG context. The orchestrator records every memory event (read, write, purge) in "
    "the central audit log for PDPL compliance.")

page_break(doc)

# =====================================================================
# SECTION 5 — INTERNAL PLATFORM
# =====================================================================
add_heading(doc, "5. Internal Platform — EWRS & Developer Performance", 1)

add_heading(doc, "5.1 EWRS — 10-Indicator Risk Engine", 2)
add_para(doc,
    "The Early Warning Risk System is the intelligence core of the Internal Management Platform. The engine "
    "evaluates ten indicators across three scopes (project, zone, developer), each with admin-configurable "
    "thresholds and SLA-bound severity bands. The RiskEngineService runs as a Hangfire job and writes RiskAlert "
    "rows to the audit-immutable trail.", space_after=4)
add_table(doc,
    ["Indicator", "Default Threshold", "Risk Level", "Implementation"],
    [
        ["Project Delivery Delay (>90 days vs original)",       "90 days", "Critical → L3", "RiskEngineService.EvaluateProjectDeliveryDelay"],
        ["Escrow Shortfall (>20% below required balance)",      "20%",     "Warning → L2",  "RiskEngineService.EvaluateEscrowShortfall"],
        ["Construction Suspension",                              "30 days inactivity → 1× threshold L2; 2× threshold L3", "L2/L3", "RiskEngineService.EvaluateConstructionSuspension"],
        ["Transaction Volume Decline (zone-level)",             "30-day count vs 12-month rolling avg, −40%", "L2", "RiskEngineService.EvaluateTransactionVolumeDecline"],
        ["Price Decline (PSF QoQ)",                              "QoQ −15% with ≥5 transactions per bucket", "L2/L3", "RiskEngineService.EvaluatePriceDecline"],
        ["Developer Score Deterioration",                        "QoQ composite drop > 20 points", "L2; L3 if new score < 40", "RiskEngineService.EvaluateDeveloperScore"],
        ["High-Risk Concentration (developer-level)",            "≥ 3 active projects with open High alerts or Stalled", "L2", "RiskEngineService.EvaluateHighRiskConcentration"],
        ["Severe Regulatory Violation",                          "Critical RERA violation within 90 days", "L3", "RiskEngineService.EvaluateSevereRegulatoryViolation"],
        ["Manual Flag (operator-triggered)",                     "RBAC: DldSupervisor+", "L1–L4 by operator", "AlertController.RaiseManualAlert"],
        ["Watchlist Trigger (investor-driven, internal echo)",   "Bulk watchlist activity", "L1", "WatchlistService.EvaluateBulkInterest"],
    ],
    col_widths_in=[2.0, 1.4, 1.0, 2.6])
add_para(doc,
    "Every threshold value is configurable from the EWRS configuration panel by authorised administrators. "
    "Changes are written to the central audit-log on commit, with timestamp, actor identity, old value, and "
    "new value — a permanent configuration audit trail DLD can query at any time.")

add_heading(doc, "5.2 SLA-Driven Auto-Escalation (L1 → L4)", 2)
add_para(doc,
    "Multi-level escalation is automated, not manual. The AlertSla helper computes per-level deadlines using "
    "UAE business hours (Monday–Friday 08:00–16:00, UAE public holidays excluded) and the AlertEscalationService "
    "Hangfire job (cron */5 * * * *) runs every five minutes:")
add_bullets(doc, [
    "Level 1 — initial alert. Deadline: 4 business hours.",
    "Level 2 — auto-escalated if L1 unacknowledged within 4 business hours. Deadline: 1 business hour.",
    "Level 3 — auto-escalated if L2 unacknowledged within 1 business hour. Deadline: immediate.",
    "Level 4 — auto-escalated immediately (RFP §8.3). Out-of-hours: paged to the on-call rota via the SMS provider.",
])
add_para(doc,
    "On escalation, the alert is reset to Status=New, AutoEscalated=true, LastEscalatedAt is stamped, and the "
    "deadline is recomputed from the new level's SLA. Acknowledged alerts are not touched. Level-4 alerts never "
    "escalate further. Five xUnit tests in AlertEscalationServiceTests cover the regression cases.")

add_heading(doc, "5.3 Developer Performance & Rating Engine (Phase 3)", 2)
add_para(doc,
    "The Developer Performance & Rating engine provides DLD with the transparent, data-driven accountability "
    "framework for real estate developers that directly addresses the GRETI Governance sub-index gap. The "
    "composite score is computed by a configurable weight matrix evaluated quarterly:", space_after=4)
add_table(doc,
    ["Scoring Criterion (Default Weight)", "Source & Calculation"],
    [
        ["On-Time Delivery (30%)",          "Projects delivered on or before original completion date / total delivered projects, rolling 24 months."],
        ["RERA Compliance (25%)",           "Inverse of weighted regulatory-violation count (Critical ×3, Major ×2, Minor ×1) within 24 months."],
        ["Unit Delivery Volume (15%)",      "Normalised to peer cohort by registered project count."],
        ["Escrow Compliance (15%)",         "Inverse of EWRS escrow-shortfall events, weighted by recency and severity."],
        ["Customer Complaint Rate (10%)",   "Unique complaints / units delivered, normalised."],
        ["Beneficial Ownership Disclosure (5%)", "Binary — disclosed vs not disclosed in DLD records."],
    ],
    col_widths_in=[3.0, 4.0])
add_para(doc,
    "Weights are configurable from the Weight Configuration Panel by SystemAdministrator role only. Every "
    "weight change writes both a ModifiedBy/ModifiedAt stamp on the target row and an append-only AuditLog row "
    "(entityType, entityId, old JSON, new JSON, actor). Audit-log immutability is enforced at the IretpDbContext "
    "SaveChangesAsync layer — any attempt to modify or delete an AuditLog row throws InvalidOperationException. "
    "Three xUnit tests in AuditLogImmutabilityTests cover this.")
add_para(doc,
    "The Public Developer Scorecard at /developers/scorecards displays a simplified view of each developer's "
    "rating: completed projects count, on-time delivery percentage, RERA compliance rating, total units "
    "delivered. Public-facing data is PII-free per RFP §19.2.")

add_heading(doc, "5.4 Escrow Monitoring", 2)
add_para(doc,
    "The Escrow Monitoring surface aggregates escrow account state per project: total deposited, "
    "developer-released, contractor-released, regulator-held, current balance, and last reconciliation "
    "timestamp. Shortfall events are computed against project disbursement schedules and feed the EWRS "
    "Escrow Shortfall indicator. The admin surface is RBAC-gated to EscrowMonitoringOfficer and "
    "DldSupervisor roles; export to CSV/PDF is RBAC-gated and writes an AuditLog row.")

page_break(doc)

# =====================================================================
# SECTION 6 — PHASED DELIVERY
# =====================================================================
add_heading(doc, "6. Phased Delivery Plan", 1)

add_heading(doc, "6.1 Timeline Overview — 90 Calendar Days", 2)
add_para(doc,
    "We commit to delivering all four phases within 90 calendar days from contract signing — 30 days ahead of "
    "the RFP maximum of 120 days. Phases are developed in parallel across dedicated team tracks; go-live "
    "sequencing follows the mandatory gate order. Each phase goes live only after its UAT is signed off by "
    "DLD.", space_after=6)
add_table(doc,
    ["Phase", "Deliverable", "Development / Go-Live", "Key Activities"],
    [
        ["Phase 0",
         "Data Familiarisation & Architecture Alignment",
         "Days 1–14 (parallel with Phase 1)",
         "Connect to OneLake Gold; map schema with DLD data liaison; configure derived-metric pipelines; finalise field mapping; LLM evaluation report Day 14."],
        ["Phase 1",
         "External Public Portal (English + Arabic)",
         "Dev: Days 1–28\nUAT: Days 28–35\nGo-Live: Day 35",
         "Portal hardening + DLD data integration + i18n. Pre-Publication Analytics Assessment gate Days 28–32. Load-test sign-off prerequisite."],
        ["Phase 2",
         "Early Warning Risk System (EWRS)",
         "Dev: Days 14–48\nUAT: Days 48–55\nGo-Live: Day 56",
         "EWRS 10-indicator engine integrated against DLD data, RBAC for 6 internal roles, Internal AI Agent, escrow monitoring, audit-log immutability sign-off."],
        ["Phase 3",
         "Developer Performance & Rating + VAPT",
         "Dev: Days 42–68\nVAPT: Days 56–68\nUAT: Days 65–75\nGo-Live: Day 78",
         "Developer Rating Engine integration, public scorecard data sign-off, VAPT external assessor engagement (pre-booked Day 1 for Day 56 slot)."],
        ["Phase 4",
         "Extended Languages + Supplementary Features",
         "Dev: Days 60–85\nUAT: Days 80–88\nFinal Go-Live: Day 90",
         "Multilingual extension (ZH/RU/UR/FR/HI/DE), ESG enhancements, benchmarking, PDF generator, documentation finalisation, training, handover."],
    ],
    col_widths_in=[0.8, 1.7, 1.7, 2.8])

add_heading(doc, "6.2 Why 90 Days is Achievable from a Working Foundation", 2)
add_para(doc, "The 90-day commitment is grounded in four structural advantages, each of which produces measurable time savings versus a from-scratch build:", space_after=4)
add_bullets(doc, [
    "Phase 1 is integration, not construction. The External Portal codebase is in production-build state today. "
    "Phase 1 work is connecting the OneLake Gold layer, validating the field mapping, completing the "
    "Pre-Publication Analytics Assessment, and signing off UAT. This compresses Phase 1 to 35 days versus the "
    "RFP allowance of 60.",
    "Phase 2 EWRS is hardening, not building. The 10-indicator engine, L1–L4 auto-escalation, and central "
    "audit log are already operating in the reference build (verified by 10 xUnit tests in "
    "RiskEngineServiceTests + 5 in AlertEscalationServiceTests). Phase 2 work is integrating DLD-specific "
    "thresholds, finalising the playbook content with DLD, and UAT.",
    "Parallel team tracks. Dedicated sub-teams own each phase. Phase 2 development starts Day 14 regardless "
    "of Phase 1 go-live status. Phase 3 starts Day 42. Phase 4 starts Day 60.",
    "Phase 4 is configuration-dominant. Extended language support is primarily an i18n resource-file exercise "
    "— translating string files, validating RTL layout in non-Arabic RTL scripts, testing chart label rendering. "
    "Combined with the supplementary modules (ESG, benchmarking, PDF generator), Phase 4 is completable in three "
    "weeks by a focused team.",
])
add_para(doc, "The one constraint we flag transparently: the VAPT assessor must be pre-engaged at contract signing to secure the Day 56 testing slot. This is a coordination action initiated on Day 1.")

add_heading(doc, "6.3 Pre-Publication Analytics Assessment — Hard Go-Live Gate", 2)
add_para(doc,
    "Before any data goes live on the External Portal, the platform must pass a structured Pre-Publication "
    "Analytics Assessment. This is the RFP's required methodology, and it is a hard go-live gate — the portal "
    "does not launch until it passes and is signed off in writing by DLD's designated data liaison.")
add_para(doc, "The assessment covers three layers of analytical validation:", space_after=4)
add_bullets(doc, [
    "Raw data reconciliation. A random sample of transaction records is pulled from the Gold layer and "
    "compared against what the platform displays. Required match rate: 99.5% or above. Any field-level "
    "discrepancy is traced back to the pipeline and resolved before the gate can pass.",
    "Derived metric validation. Every computed metric — price per sqft, YoY %, zone aggregations, off-plan "
    "share, median values, rental yield — is independently recalculated from raw Gold-layer data and compared "
    "to platform output. Zero tolerance for calculation errors. Trend direction (arrows, %) is verified for "
    "correct direction and magnitude.",
    "Edge case review. Statistical outliers are identified, classified as legitimate extreme values or data "
    "entry errors, and handled by agreement with DLD before go-live. Any outlier that would materially mislead "
    "a public investor is flagged and either annotated or suppressed pending DLD instruction.",
])
add_para(doc,
    "The assessment is conducted jointly by [Vendor Legal Entity Name] and DLD's data liaison. Results are "
    "documented in a Pre-Publication Analytics Assessment Report, which becomes a permanent DLD data-governance "
    "artefact. DLD provides written sign-off before the platform goes live. If the assessment fails any "
    "threshold, the Phase 1 go-live date moves — this gate does not bend.")

page_break(doc)

# =====================================================================
# SECTION 7 — SECURITY & DESC
# =====================================================================
add_heading(doc, "7. Security & DESC Compliance", 1)

add_heading(doc, "7.1 DESC ISR v3 Compliance Framework", 2)
add_para(doc,
    "Security and regulatory compliance are not a post-delivery audit. They are structural design constraints "
    "that shape every architectural decision from Day 1. Our DESC ISR v3 compliance roadmap is structured "
    "across the full project lifecycle and extends through the 12-month warranty period.", space_after=4)
add_table(doc,
    ["ISR v3 Control Domain", "Our Compliance Approach"],
    [
        ["Identity & Access Management",
         "ASP.NET Identity with mandatory TOTP MFA for internal users (enforced at policy level). Six-role RBAC policy registry in IRETP.Infrastructure/Identity/AuthorizationPolicies.cs. Internal idle timeout 30 minutes (RFP §10.2.1). Refresh-token rotation with reuse detection. UAE Pass OIDC federation for investors."],
        ["Cryptography & Data Protection",
         "TLS 1.2 minimum end-to-end. Azure SQL TDE at rest; Always Encrypted for PDPL-classified columns. Azure Key Vault stores all secrets (purge-protection on). Managed Service Identity from every host — no connection strings in app config. HMAC-SHA256 unsubscribe tokens with rotation-safe secret."],
        ["Audit & Logging",
         "Append-only AuditLog enforced at the IretpDbContext SaveChangesAsync layer (any UPDATE or DELETE throws InvalidOperationException). Central audit trail on every threshold and weight change with actor identity, old/new JSON. App Insights captures every request, every AI inference, and every alert dispatch."],
        ["Vulnerability Management",
         "Security CI (.github/workflows/security.yml): dotnet list package --vulnerable fails the build on Critical/High; OWASP ZAP baseline scan on every push to main; Bicep linting on every IaC change. Quarterly external VAPT in the warranty period."],
        ["Incident Response",
         "P1/P2/P3/P4 SLA defined in Section 13.2. Named Support Lead from Day 1 of warranty. Emergency out-of-hours escalation number for P1. RUNBOOK.md ships with the source on Day 1."],
        ["Data Residency & Sovereignty",
         "Bicep IaC pins every Azure resource to UAE North; DR replica is UAE Central. AI inference UAE-residency tagged on every snapshot — the /healthz/sla probe fails Unhealthy if residency drifts."],
    ],
    col_widths_in=[2.0, 5.0])

add_heading(doc, "7.2 Security Testing Plan", 2)
add_table(doc,
    ["Security Test Type", "Schedule & Approach", "Remediation Commitment"],
    [
        ["Static analysis (SAST)",
         "Every commit. CodeQL via GitHub Actions; .NET analysers enforced at solution level (TreatWarningsAsErrors).",
         "Critical/High blocks the build."],
        ["Dependency vulnerability scan",
         "Every commit. dotnet list package --vulnerable run in CI; quarterly review of transitive dependencies.",
         "Critical/High blocks the build."],
        ["DAST baseline",
         "Every push to main. OWASP ZAP baseline scan against the public surface in CI.",
         "High triggers immediate remediation; medium logged for sprint review."],
        ["External VAPT",
         "Phase 3 (Days 56–68). DESC-approved external assessor pre-engaged on Day 1.",
         "Clean report mandatory before Phase 3 UAT sign-off."],
        ["Re-VAPT in warranty",
         "Annually plus on any major release.",
         "Clean report mandatory before any production deployment with security-affecting changes."],
    ],
    col_widths_in=[1.6, 2.7, 2.7])

add_heading(doc, "7.3 ISO 27001 Roadmap", 2)
add_para(doc,
    "[Vendor Legal Entity Name] commits to completing an ISO 27001 gap assessment against the IRETP-specific "
    "platform configuration by Phase 3 (Day 78) and presenting a documented certification roadmap to DLD. "
    "We target full ISO 27001 certification for the IRETP environment within six months of final go-live. "
    "The certification scope covers the IRETP platform engineering, support operations, and the warranty "
    "team — bound to DLD's procurement contract.")

add_heading(doc, "7.4 Hosting Recommendation — Azure UAE North", 2)
add_para(doc,
    "We recommend Option A: Azure UAE North hosting. This is not a default recommendation. It is the result "
    "of a deliberate analysis of DLD's specific requirements, our reference platform's cloud-native posture, "
    "and the five-year total cost of ownership. Both options deliver identical functionality; the "
    "infrastructure choice affects time to live, operational overhead, and TCO — all of which favour Azure "
    "UAE North for a platform of this scale.", space_after=4)
add_table(doc,
    ["Criterion", "Option A — Azure UAE North (Recommended)", "Option B — On-Premises"],
    [
        ["Data residency",     "100% UAE; data within Microsoft's UAE North region; transits Microsoft sovereign infrastructure.", "100% UAE; data within DLD's physical data centre."],
        ["Infrastructure",     "Azure App Service Plan P1v3 (3 sites), Azure SQL elastic pool, Azure Cache for Redis, Front Door + WAF, App Insights, Log Analytics.", "Kubernetes + SQL Server + Redis + open-source observability stack on DLD hardware."],
        ["DESC-CSP",           "Azure UAE North pre-certified for DESC-CSP. Administrative documentation only.", "Full ISR v3 on-prem assessment required. Full DESC compliance boundary under DLD control."],
        ["Time to live",       "Foundation infrastructure live in 1 week via Bicep IaC.", "4–8 weeks for hardware procurement, racking, and configuration."],
        ["Operational overhead", "Managed services minimise OS patching, DB tuning, failover testing.", "Full operational ownership: patching, backups, hardware life-cycle, capacity planning."],
        ["5-year CAPEX",       "AED 335,500 (setup, networking, certificates)", "AED 1,414,500 (hardware BOM)"],
        ["5-year OPEX (annual)", "AED 629,200", "AED 686,200"],
        ["5-year TCO (UnifyApps-equivalent scope)", "AED ≈ 5.72M", "AED ≈ 6.86M"],
        ["Recommended for IRETP", "Yes", "No"],
    ],
    col_widths_in=[1.7, 2.7, 2.6])

add_heading(doc, "7.5 Data Residency, IAM, Audit", 2)
add_bullets(doc, [
    "Data residency: every Bicep template under /infra/bicep/ pins resources to UAE North; DR is UAE Central. "
    "The SLA probe at /healthz/sla fails Unhealthy if any path drifts outside UAE.",
    "IAM: ASP.NET Identity backed by Azure SQL. MFA mandatory for internal users via TOTP. Federated OIDC "
    "scheme for UAE Pass investor sign-in. Six-role policy registry. Refresh-token rotation with reuse "
    "detection.",
    "Audit: every threshold change, weight change, and AI model swap writes an append-only AuditLog row with "
    "actor identity, timestamp, and full old/new JSON state. The AuditLog table is immutable at the "
    "DbContext layer.",
    "Backups: daily automated encrypted backups, 30-day retention; PITR enabled on Azure SQL. Quarterly "
    "restoration test reported to DLD. DR RTO < 4 hours, RPO < 1 hour.",
    "IaC: 100% of infrastructure is defined in Bicep; all templates delivered to DLD on Day 1 of contract.",
])

page_break(doc)

# =====================================================================
# SECTION 8 — VISUALISATION
# =====================================================================
add_heading(doc, "8. Visualisation Strategy — Public and Internal", 1)

add_heading(doc, "8.1 Public Portal — Blazor Server + Chart.js + MapLibre", 2)
add_para(doc,
    "For the External Public Portal we use a Blazor Server frontend rendering Chart.js charts and MapLibre GL "
    "JS maps. This choice is informed by the projected anonymous-user volumes (10,000+ peak concurrent), the "
    "investor-grade UX requirement, and the long-term maintainability target.")
add_para(doc, "Specifically:", space_after=4)
add_bullets(doc, [
    "Zero per-user licensing cost at any public user volume. No embed-token management, no PPU/PPC fees.",
    "Full anonymous access with no licensing constraint.",
    "Unlimited UI/UX customisation, including full RTL layout, investor-grade typography, DLD brand alignment.",
    "Performance at scale: server-side data aggregation in the WebAPI means chart rendering is independent of "
    "dataset size. The KPI snapshot cache absorbs the majority of read requests before they reach SQL.",
    "Full source control: Blazor Server + Chart.js + MapLibre are mature open-source projects with predictable "
    "release cadences and no single-vendor dependency.",
])

add_heading(doc, "8.2 Internal Platform — Power BI for Named-User Dashboards", 2)
add_para(doc,
    "Power BI is retained and recommended for all Internal Platform components where named-user licensing is "
    "cost-appropriate: EWRS dashboards, Developer Rating dashboards, internal management reporting. We embed "
    "Power BI tiles via the standard report-server SDK; the Internal AI Agent can answer questions over the "
    "same semantic models. Authorised users see only the data their RBAC scope permits.")

page_break(doc)

# =====================================================================
# SECTION 9 — DATA UNDERSTANDING
# =====================================================================
add_heading(doc, "9. Data Understanding & Accuracy Methodology", 1)

add_heading(doc, "9.1 Connecting to DLD's Data Ecosystem", 2)
add_para(doc,
    "Phase 0 (Days 1–14) focuses on one thing: understanding DLD's Gold-layer structure well enough to "
    "configure correct pipelines. This involves:")
add_bullets(doc, [
    "Reading DLD's existing OneLake Gold-layer schema with DLD's designated data liaison: field names, data "
    "types, update cadence, known quality notes.",
    "Configuring IRETP read-only access to the Gold layer via the Fabric SQL endpoint. No writes, no "
    "modifications to DLD's existing data infrastructure.",
    "Agreeing the calculation logic for all derived metrics in writing with DLD before any pipeline is "
    "wired: price per sqft, rental yield, EWRS risk indicators, developer composite score.",
    "Documenting the field mapping between Gold-layer fields and IRETP display fields, formally approved by "
    "DLD before development begins. The field-mapping document becomes an annex to this contract.",
])

add_heading(doc, "9.2 Pre-Publication Analytics Assessment Methodology", 2)
add_para(doc,
    "Covered in detail in Section 6.3. The headline: a three-layer validation (raw data reconciliation at "
    "99.5%, derived-metric validation at zero tolerance, edge-case review with DLD agreement) signed off in "
    "writing by DLD before any data is publicly visible. The gate does not bend.")

add_heading(doc, "9.3 Ongoing Data Accuracy Monitoring", 2)
add_para(doc,
    "DLD's Gold layer is the source of truth. If Gold contains an error, IRETP reflects that error — the "
    "platform connects faithfully. Our responsibility is to ensure the platform displays exactly what Gold "
    "contains, with no transformation errors and no pipeline-introduced discrepancies. Ongoing monitoring:")
add_bullets(doc, [
    "Freshness probe. The /healthz/sla endpoint reports the age of the last KPI snapshot vs the 15-minute "
    "TTL. Any drift beyond 30 minutes raises an alert.",
    "Reconciliation sampler. A nightly Hangfire job samples 100 random transactions from Gold and compares to "
    "the cached snapshot. Discrepancies > 0.5% raise a P2 incident.",
    "Notification SLA probe. The notification-delivery health check samples the last hour of dispatched "
    "alerts and computes P95 latency per channel against §6.2 budgets (Email 5 min, SMS 3 min, In-Platform 30 s).",
    "AI residency probe. Every AI snapshot logs the inference region. The SLA endpoint computes the percentage "
    "of UAE-resident inferences in the trailing hour; any non-UAE inference fails the platform Unhealthy "
    "until investigated.",
])

page_break(doc)

# =====================================================================
# SECTION 10 — REFERENCE INTEGRATION
# =====================================================================
add_heading(doc, "10. Data Integration Reference Architecture", 1)
add_para(doc,
    "We approach DLD's IRETP data integration with the architectural pattern proven to scale across regulated "
    "public-sector data platforms: a thin, read-only integration into the analytical source of truth, paired "
    "with a structured validation gate before any data reaches end users.")

add_heading(doc, "10.1 Gold Layer as Single Source of Analytical Truth", 2)
add_para(doc,
    "The IRETP service tier connects directly to OneLake Gold via the Fabric SQL endpoint. No intermediate "
    "data warehouse, no proprietary integration layer. Every derived metric the platform displays is "
    "computed from a Gold-layer query — there is no parallel data store to drift out of sync.")

add_heading(doc, "10.2 Validation Pipeline Architecture", 2)
add_para(doc, "The validation pipeline operates at three points in the data lifecycle:", space_after=4)
add_bullets(doc, [
    "On ingestion (Phase 0). The Pre-Publication Analytics Assessment validates the entire Gold-layer "
    "consumption end-to-end before public go-live. Signed off in writing by DLD.",
    "On display (runtime). Every KPI request goes through the cache layer; cache misses trigger a fresh Gold "
    "query whose result is reconciled against the prior snapshot before being served.",
    "On audit (continuous). The reconciliation sampler nightly job samples Gold-layer rows and compares to "
    "displayed data. Discrepancies > 0.5% raise an incident.",
])

add_heading(doc, "10.3 Bilingual Data Handling", 2)
add_para(doc,
    "Arabic field handling is first-class throughout the data layer. Zone names, developer names, owner "
    "names (for the Beneficial Ownership module), property categories, transaction types — all are stored "
    "and rendered in their native script. The Beneficial Ownership page already renders Arabic owner names "
    "under their English equivalents with proper RTL styling and lang=\"ar\".")

page_break(doc)

# =====================================================================
# SECTION 11 — TEAM
# =====================================================================
add_heading(doc, "11. Project Team & Qualifications", 1)

add_heading(doc, "11.1 Core Delivery Team", 2)
add_para(doc,
    "The IRETP delivery team uses standard professional-services profiles. Full per-phase cost breakdown is "
    "in Section 12.4.", space_after=4)
add_table(doc,
    ["Role (Rate Card)", "Location", "IRETP Remit"],
    [
        ["Project Manager",            "Dubai, UAE",            "Single contracting point of accountability. Daily standup, weekly DLD status, gate-sign-off custody."],
        ["Delivery / Technical Lead",  "Dubai, UAE",            "Architecture authority. Owns Clean-Architecture conformance, security posture, code-review gate."],
        ["Senior .NET Developer × 2",  "Dubai, UAE (onsite)",   "Phase 1–3 implementation lead. Owns WebAPI, AdminAPI, Hangfire jobs."],
        ["Mid .NET Developer × 2",     "Dubai, UAE / remote UAE timezone", "Feature implementation across phases. Owns Blazor surfaces and integration tests."],
        ["QA Engineer × 1",            "Dubai, UAE",            "Test-plan authorship, UAT facilitation, automated regression suite extension."],
        ["UI/UX Designer × 1",         "Dubai, UAE",            "Bilingual UX, RTL fidelity, design-system stewardship, ESG/GRETI display reviews."],
        ["DevOps / DESC Liaison × 1",  "Dubai, UAE",            "Bicep IaC, Azure landing zone, DESC-CSP filings, VAPT coordination."],
    ],
    col_widths_in=[2.2, 2.0, 3.0])

add_heading(doc, "11.2 UAE Market Presence", 2)
add_para(doc,
    "[Vendor Legal Entity Name] operates a UAE-registered legal entity. The Project Manager, Delivery / "
    "Technical Lead, onsite Senior Developers, QA Engineer, UI/UX Designer, and DevOps/DESC Liaison are all "
    "UAE-resident. Remote developer roles are filled from the nearest UAE-time-aligned office. DLD can "
    "request in-person meetings at any point during delivery.")

page_break(doc)

# =====================================================================
# SECTION 12 — COST
# =====================================================================
add_heading(doc, "12. Cost Structure", 1)

add_heading(doc, "12.1 Cost Summary", 2)
add_para(doc, "All amounts AED, exclusive of VAT. Option A (Azure UAE North) figures.", italic=True, color=COL_MUTED, size=9.5)
add_table(doc,
    ["Cost category", "Amount (AED)", "Type"],
    [
        ["Infrastructure CAPEX", "335,500", "One-time"],
        ["Software Licences (Year 1)", "235,200", "One-time"],
        ["Application Development (M1–M4)", "2,007,000", "One-time"],
        ["Total one-time CAPEX", "2,577,700", "—"],
        ["Annual OPEX (Year 1, includes 12-month warranty)", "629,200", "Recurring"],
    ],
    col_widths_in=[3.4, 1.8, 1.8])

add_heading(doc, "12.2 Infrastructure CAPEX (One-Time)", 2)
add_table(doc,
    ["Infrastructure item", "Option A — Azure UAE North", "Option B — On-Premises"],
    [
        ["Compute · App Service / VMs",                "48,000",  "320,000"],
        ["SQL Server (managed / licensed)",            "36,000",  "180,000"],
        ["Redis cache cluster",                        "18,000",  "85,000"],
        ["Azure AI Search / vector index",             "24,000",  "65,000"],
        ["Microsoft Fabric / OneLake lakehouse",       "42,000",  "200,000"],
        ["Blob Storage / on-prem NAS",                 "12,000",  "55,000"],
        ["WAF + Front Door / on-prem F5",              "30,000",  "140,000"],
        ["CDN (public static assets)",                  "8,000",  "25,000"],
        ["Private endpoints + site-to-site network",   "14,000",  "60,000"],
        ["DR · cross-region replication (≤ 4 h RPO)",  "28,000",  "180,000"],
        ["SSL certificates (wildcard, 3-year)",         "4,500",  "4,500"],
        ["Monitoring · telemetry (App Insights / Seq)","16,000",  "45,000"],
        ["Initial security hardening + pen-test",      "55,000",  "55,000"],
        ["Subtotal (AED)",                            "335,500", "1,414,500"],
    ],
    col_widths_in=[3.0, 1.9, 2.1])

add_heading(doc, "12.3 Software Licences (Year 1, Open-Source Excluded)", 2)
add_table(doc,
    ["Licence / Subscription", "Vendor", "Annual (AED)"],
    [
        [".NET 9 runtime + ASP.NET Core",       "Microsoft (open-source)",   "0"],
        ["Visual Studio Enterprise (5 seats)",  "Microsoft",                 "64,000"],
        ["Hangfire Pro (2 servers)",            "HangfireIO",                "14,400"],
        ["Azure OpenAI Service (tokens, Y1 est.)", "Microsoft",              "90,000"],
        ["AI fallback model (UAE-resident provider)", "Selected provider",   "40,000"],
        ["SMTP relay (SendGrid / Postmark)",    "Twilio / Postmark",         "2,400"],
        ["SMS gateway (UAE operator)",          "UAE operator",              "18,000"],
        ["PDF generation (QuestPDF)",           "QuestPDF",                  "2,800"],
        ["Seq log server (cloud)",              "Datalust",                  "3,600"],
        ["Subtotal (Year 1)",                   "—",                         "235,200"],
    ],
    col_widths_in=[3.0, 2.2, 1.8])

add_heading(doc, "12.4 Application Development (M1–M4)", 2)
add_table(doc,
    ["Deliverable", "Month", "Effort (PD)", "Cost (AED)"],
    [
        ["External Public Portal (CMS · Dashboard · Transactions · Map · Indices)", "M1 · P1", "120", "216,000"],
        ["Bilingual EN/AR · RTL · i18n pipeline",                  "M1 · P1", "40",  "72,000"],
        ["Slice & Dice Analytics engine",                          "M1 · P1", "60",  "108,000"],
        ["Real Estate AI Agent (Orchestrator · RAG · guardrails)", "M1 · P1", "100", "180,000"],
        ["Investor Notifications (email / SMS / in-platform)",     "M1 · P1", "60",  "108,000"],
        ["Open Data API · Developer Portal · API-key management",  "M1 · P1", "50",  "90,000"],
        ["Perf hardening · Redis cache · EF tuning",               "M1 · P1", "30",  "54,000"],
        ["Phase 1 QA · UAT support",                               "M1 · P1", "30",  "54,000"],
        ["Phase 1 subtotal",                                       "M1",      "490", "882,000"],
        ["Watchlist v2 · ESG · GRETI sub-index tracker",           "M2 · P2", "80",  "144,000"],
        ["Analytics enhancements · saved views · embed",           "M2 · P2", "40",  "72,000"],
        ["Performance hardening · OpenTelemetry · load-test sign-off","M2 · P2","30","54,000"],
        ["Phase 2 QA · UAT",                                       "M2 · P2", "20",  "36,000"],
        ["Phase 2 subtotal",                                       "M2",      "170", "306,000"],
        ["EWRS (risk engine · L1–L4 escalation · dashboard · playbooks)","M3 · P3","100","180,000"],
        ["Developer Performance & Rating + Escrow Monitoring",     "M3 · P3", "80",  "144,000"],
        ["Beneficial Ownership · Mortgage analytics · Name Validation","M3 · P3","40","72,000"],
        ["DESC ISR v3 hardening + VAPT (external consultant)",     "M3 · P3", "30",  "90,000"],
        ["Phase 3 QA · UAT · security acceptance",                 "M3 · P3", "25",  "45,000"],
        ["Phase 3 subtotal",                                       "M3",      "275", "531,000"],
        ["Multilingual extension (ZH · RU · UR · FR · HI · DE) + AI multi-language","M4 · P4","90","162,000"],
        ["Advanced analytics · AI fine-tuning pipeline",           "M4 · P4", "50",  "90,000"],
        ["Phase 4 QA · UAT · language acceptance",                 "M4 · P4", "20",  "36,000"],
        ["Phase 4 subtotal",                                       "M4",      "160", "288,000"],
        ["Total application development (M1–M4)",                  "M1–M4",   "1,095","2,007,000"],
    ],
    col_widths_in=[3.4, 0.9, 1.0, 1.3])
add_para(doc,
    "Blended day rates: Architect AED 2,200 · Senior Developer 1,800 · Mid Developer 1,500 · QA / DevOps 1,400.",
    italic=True, color=COL_MUTED, size=9.5)

add_heading(doc, "12.5 Annual OPEX (Year-1 Recurring)", 2)
add_table(doc,
    ["OPEX item", "Option A — Azure UAE North", "Option B — On-Premises"],
    [
        ["Compute · VM · App Service",            "48,000",  "85,000"],
        ["SQL Server (managed annual)",           "36,000",  "24,000"],
        ["Redis · caching annual",                "18,000",  "12,000"],
        ["AI model API consumption (annual est.)","130,000", "130,000"],
        ["OneLake · Fabric storage & compute",    "42,000",  "68,000"],
        ["Monitoring & SIEM",                     "16,000",  "28,000"],
        ["SMTP · SMS channels",                   "20,400",  "20,400"],
        ["Software licences (annual renewal)",    "84,800",  "84,800"],
        ["Warranty & support (12-mo post go-live)","180,000","180,000"],
        ["Security patching · monthly scans",     "36,000",  "36,000"],
        ["Data quality · reconciliation tooling", "18,000",  "18,000"],
        ["Annual OPEX (Year 1)",                  "629,200", "686,200"],
    ],
    col_widths_in=[3.0, 1.9, 2.1])

add_heading(doc, "12.6 Payment Milestones", 2)
add_para(doc, "Pay on UAT sign-off. 6 gates across 4 months.", italic=True, size=10, color=COL_MUTED)
add_table(doc,
    ["Milestone", "Trigger", "%", "Amount (AED)"],
    [
        ["M0", "Contract signature",                                                                "10%", "200,700"],
        ["M1", "Week 1 — mobilisation complete (infrastructure live · team onboard)",               "10%", "200,700"],
        ["M2", "End of Month 1 — Phase 1 UAT sign-off (External Portal · AI Agent · Open Data)",    "25%", "501,750"],
        ["M3", "End of Month 2 — Phase 2 UAT sign-off (ESG · GRETI · Analytics v2)",                "15%", "301,050"],
        ["M4", "End of Month 3 — Phase 3 UAT sign-off (EWRS · Dev Rating · Escrow · VAPT clean)",   "25%", "501,750"],
        ["M5", "End of Month 4 — Phase 4 UAT sign-off (Multilingual · Advanced Analytics)",         "15%", "301,050"],
    ],
    col_widths_in=[0.7, 4.0, 0.6, 1.8])
add_para(doc, "100% of the AED 2,007,000 development CAPEX is conditional on DLD-signed acceptance certificates.", italic=True, size=10, color=COL_MUTED)

add_heading(doc, "12.7 Indicative 5-Year Total Cost of Ownership", 2)
add_table(doc,
    ["Period", "Component", "Amount (AED)"],
    [
        ["Year 0", "One-time CAPEX (infrastructure + licences Year 1 + application dev)", "2,577,700"],
        ["Year 1", "Annual OPEX (includes 12-month warranty)",                            "629,200"],
        ["Year 2", "Annual OPEX",                                                          "629,200"],
        ["Year 3", "Annual OPEX",                                                          "629,200"],
        ["Year 4", "Annual OPEX",                                                          "629,200"],
        ["Year 5", "Annual OPEX",                                                          "629,200"],
        ["5-year Total", "Indicative TCO",                                                "5,723,700"],
    ],
    col_widths_in=[1.0, 4.2, 1.8])
add_para(doc,
    "TCO excludes inflation and the AI-token consumption volatility inherent to LLM usage. The token line will "
    "be re-tiered at contract sign based on agreed monthly volume. All licence renewals are subject to vendor "
    "list-price changes outside [Vendor Legal Entity Name]'s control.",
    italic=True, size=9.5, color=COL_MUTED)

page_break(doc)

# =====================================================================
# SECTION 13 — WARRANTY & SUPPORT
# =====================================================================
add_heading(doc, "13. Warranty & Technical Support — 12 Months Post Go-Live", 1)

add_heading(doc, "13.1 Support Philosophy", 2)
add_para(doc,
    "The 12-month warranty and technical support obligation is not a separate procurement or an optional "
    "add-on. It is a contractual extension of our delivery responsibility. We do not consider a project "
    "delivered until the platform has operated stably, accurately, and securely through its first full year "
    "of live operation.")
add_para(doc,
    "Our warranty and support scope is structured around five categories: defect resolution, data pipeline "
    "defect correction, security patch management, performance defect resolution, and DLD data structure "
    "change management. All five are covered within the warranty contract.")

add_heading(doc, "13.2 Support Incident SLA Commitments", 2)
add_table(doc,
    ["Priority", "Definition", "Response", "Resolution", "Coverage"],
    [
        ["P1", "Platform unavailable; investor-facing data wrong; security breach.",          "15 min", "4 h",  "24×7 incl. weekends/UAE holidays"],
        ["P2", "Major feature degraded; EWRS escalation impaired; AI Agent down on one tier.","30 min", "8 h",  "Business hours UAE (Mon–Fri 08:00–16:00)"],
        ["P3", "Minor functional defect; cosmetic; localisation issue.",                     "2 h",    "3 d",  "Business hours UAE"],
        ["P4", "Enhancement request; documentation correction; data-model expansion request.","1 d",   "By next planning cycle", "Business hours UAE"],
    ],
    col_widths_in=[0.6, 2.6, 0.8, 0.9, 2.1])
add_para(doc,
    "Business hours: Monday–Friday, 08:00–16:00 UAE Standard Time, excluding UAE public holidays. A named "
    "Support Lead is assigned as DLD's primary contact. An emergency out-of-hours escalation number is "
    "provided from Day 1 of the warranty period for P1 incidents.")

add_heading(doc, "13.3 Warranty & Support Pricing", 2)
add_table(doc,
    ["12-Month Warranty & Support Cost Item", "Annual Cost (AED)"],
    [
        ["S.1 Dedicated Support Team: Delivery/Technical Lead (25%) + Developer × 2 (50% each) + QA Engineer (25%)", "Included in OPEX"],
        ["S.2 120 person-hours minor changes pool. Additional hours: AED 950 per person-hour.", "Included"],
        ["S.3 Security vulnerability monitoring and patch management. Included in S.1 team allocation.", "Included"],
        ["S.4 Infrastructure operations. Covered by Azure managed services. No additional cost.", "0"],
        ["S.5 Help desk and incident management tooling. Enterprise ticketing with DLD shared view.", "18,260"],
        ["TOTAL 12-MONTH WARRANTY & SUPPORT COST (within Annual OPEX)", "180,000"],
    ],
    col_widths_in=[5.2, 1.8])

page_break(doc)

# =====================================================================
# SECTION 14 — AI KNOWLEDGE UPDATE
# =====================================================================
add_heading(doc, "14. AI Agent Knowledge Update Approach", 1)
add_para(doc,
    "The AI Agent's knowledge layer is grounded in the OneLake Gold layer. Knowledge updates are managed "
    "by the AdminAPI surface in IRETP.AdminAPI without vendor involvement after handover.")

add_heading(doc, "14.1 How Knowledge Updates Work", 2)
add_para(doc, "Each knowledge source is configured once in the AdminAPI:")
add_bullets(doc, [
    "Source — the Gold-layer table or view that backs the knowledge layer.",
    "Refresh schedule — scheduled (hourly, daily, weekly) via Hangfire cron, or event-triggered (on webhook from CMS publish, on Gold-layer freshness signal).",
    "Index strategy — full reindex or incremental delta.",
])
add_table(doc,
    ["Knowledge Layer", "Update Mechanism"],
    [
        ["Transactions / Projects / Developers", "Hangfire cron, hourly. Incremental delta from Gold."],
        ["RERA regulations / DLD procedures",    "CMS publish webhook → immediate reindex of the affected document."],
        ["Zone boundaries / GeoJSON",            "Manual trigger on Dubai Municipality updates."],
        ["Beneficial ownership",                  "Daily cron + on-demand re-index from AdminAPI."],
    ],
    col_widths_in=[2.4, 4.6])

add_heading(doc, "14.2 Index Management in the AdminAPI", 2)
add_bullets(doc, [
    "Index freshness dashboard — last update timestamp per knowledge layer, next scheduled update, record count.",
    "On-demand re-index — any knowledge layer can be re-indexed manually from the admin surface.",
    "Rollback — every index update creates a snapshot. SystemAdministrator can roll back to any prior version in two clicks; completes within 5 minutes.",
    "Accuracy benchmark — the accuracy harness runs automatically after each index update; deviation > threshold raises an alert before updated content reaches end users.",
])

add_heading(doc, "14.3 LLM Model Lifecycle (BYO LLM)", 2)
add_para(doc,
    "The orchestrator is LLM-agnostic. Switching the active model on any tier is a configuration update with "
    "no redeployment and no downtime. If DLD wants to re-evaluate models in Year 2 against newer releases, the "
    "accuracy harness runs again. Token spend per session is independent of the model identity from a billing "
    "structure perspective.")

page_break(doc)

# =====================================================================
# SECTION 15 — RISK REGISTER
# =====================================================================
add_heading(doc, "15. Risk Register", 1)
add_table(doc,
    ["Risk", "Probability", "Impact", "Mitigation"],
    [
        ["Gold-layer schema drift during Phase 0",
         "Medium",
         "Schedule",
         "Schema versioning agreed with DLD data liaison Day 2; integration tests run on every Gold-layer change; field-mapping document is contractual."],
        ["Pre-Publication Analytics Assessment fails reconciliation",
         "Low",
         "Schedule",
         "Dry-run assessment Days 21–24 surfaces issues 10 days before the formal gate. Failing items resolved in pipeline before the formal gate."],
        ["VAPT findings require remediation beyond Day 68",
         "Medium",
         "Schedule",
         "VAPT scheduled Days 56–68; 7-day remediation buffer baked into the Phase 3 plan; Critical/High findings block UAT sign-off."],
        ["AI residency drift",
         "Low",
         "Compliance",
         "SLA probe fails Unhealthy on any non-UAE inference; automated rollback to UAE-resident fallback tier."],
        ["Investor surge above 3× projected volume",
         "Medium",
         "Performance",
         "App Service Plan auto-scale on CPU/memory; Redis absorbs read pressure; load test rerun on every major release."],
        ["DESC certification timeline slips into warranty",
         "Medium",
         "Compliance",
         "DESC engagement begins Day 1; gap assessment by Day 78; full certification targeted within 6 months of go-live; documented roadmap delivered to DLD regardless."],
        ["Key personnel attrition",
         "Low",
         "Schedule",
         "Knowledge transfer is continuous (all artefacts in source repo); Delivery/Technical Lead has named backup; no single-person dependency on any phase."],
        ["DLD policy change mid-flight",
         "Medium",
         "Scope",
         "Configurable thresholds, weight matrix, and i18n strings mean most policy changes are configuration, not code. Material changes flow through change-control with cost/schedule impact agreed in writing."],
    ],
    col_widths_in=[2.0, 0.9, 0.9, 3.2])

page_break(doc)

# =====================================================================
# SECTION 16 — KT & EXIT
# =====================================================================
add_heading(doc, "16. Knowledge Transfer & Exit Plan", 1)

add_heading(doc, "16.1 Deliverables at Project Completion (Day 90)", 2)
add_bullets(doc, [
    "Source repositories — all .NET application code, Bicep IaC, CI/CD pipeline configurations, Hangfire job "
    "definitions, AI orchestration configuration, knowledge-source configuration files. Delivered as GitHub "
    "repositories transferred to DLD's GitHub organisation.",
    "Complete technical documentation — ARCHITECTURE.md, API_REFERENCE.md, RUNBOOK.md, DISASTER_RECOVERY.md, "
    "DATA_DICTIONARY.md (all five RFP §13 #53 pillars). Delivered with the source repository.",
    "Infrastructure control — Azure subscription access, Key Vault control, DNS management transferred to DLD's "
    "Azure tenant on go-live. No [Vendor Legal Entity Name]-controlled infrastructure persists after handover.",
    "Data export — full platform data export in machine-readable, open formats (JSON, CSV, Parquet) completed "
    "within 30 calendar days of contract expiry upon request.",
    "Data destruction certification — written certification of complete data destruction from all [Vendor "
    "Legal Entity Name]-controlled systems within 60 days of DLD's request.",
])

add_heading(doc, "16.2 Intellectual Property", 2)
add_para(doc,
    "All deliverables, results, works, documents, data, reports, designs, and materials prepared, developed, "
    "or submitted under this contract shall be the exclusive intellectual property of the Dubai Land "
    "Department, consistent with Section 4.4 of the RFP. [Vendor Legal Entity Name] retains no rights to use, "
    "reproduce, or publish any deliverable in whole or in part without prior written approval from DLD. There "
    "is no proprietary platform layer between DLD and the application source; everything in the deliverable "
    "stack is either standard Microsoft technology, open-source, or written specifically for DLD under this "
    "engagement.")

page_break(doc)

# =====================================================================
# SECTION 17 — GRETI
# =====================================================================
add_heading(doc, "17. Beyond the Brief — GRETI Score Maximisation", 1)
add_para(doc,
    "Section 20 of the RFP identifies additional features that maximise GRETI score improvement beyond core "
    "scope. Three of these are not roadmap items in this proposal — they are already shipping in the reference "
    "build today.")

add_heading(doc, "17.1 Beneficial Ownership Transparency (Already Shipping)", 2)
add_para(doc,
    "JLL's 2024 GRETI report identifies beneficial ownership transparency as a key gap for the UAE. "
    "International institutional investors, particularly those subject to FATF compliance, need visibility "
    "into the corporate ownership structures of developers they are considering.")
add_para(doc,
    "The reference build ships a Beneficial Ownership page that renders ownership data per registered "
    "developer: beneficial owner name (English + Arabic, RTL-styled when populated), ownership percentage, "
    "UAE commercial registration number. Where data is not available, a clear disclosure statement renders "
    "in place of a blank field — transparency about the absence of data is itself a transparency measure.")

add_heading(doc, "17.2 Mortgage & Debt Market Transparency (Already Shipping)", 2)
add_para(doc,
    "The GRETI 2024 report identified debt market transparency as a significant gap. The reference build "
    "ships a Mortgage page with five KPIs:")
add_bullets(doc, [
    "Total mortgage value (AED) per zone per period.",
    "Total mortgage count.",
    "Mortgage-to-transaction ratio.",
    "Average mortgage size — the fifth KPI surfaced explicitly per FR mapping.",
    "Year-on-year mortgage trend.",
])
add_para(doc,
    "All mortgage data is sourced exclusively from registered DLD transaction records — the same Gold-layer "
    "data as the rest of the portal. No additional source system is required.")

add_heading(doc, "17.3 ESG Indicators (Already Shipping)", 2)
add_para(doc,
    "The reference build ships an ESG dashboard with the LEED ↔ Estidama Pearl ↔ BREEAM ↔ investor-signal "
    "rubric. The GIS map exposes a fourth heatmap mode — ESG Coverage — that visualises green-certified "
    "coverage by zone with banded fill colour (>60% / 35–60% / 15–35% / <15%). Per-zone ESG Coverage % markers "
    "render alongside the standard analytical layers.")

add_heading(doc, "17.4 Implementation Roadmap for Beyond-Brief Features", 2)
add_para(doc,
    "We recommend DLD discuss the Beneficial Ownership, Mortgage Transparency, and ESG modules with the JLL "
    "GRETI assessment team before finalising the project scope. All three are cited directly in the GRETI 2024 "
    "report, and their inclusion creates measurable indicators of progress that the next GRETI assessment cycle "
    "can capture within this project's delivery timeline.")

page_break(doc)

# =====================================================================
# ANNEX A — ARCHITECTURE DIAGRAM
# =====================================================================
add_heading(doc, "Annex A — System Architecture Diagram", 1)
add_para(doc,
    "The diagram below illustrates the four-tier IRETP architecture: the User Access tier, the Application & "
    "AI tier, the Data & Integration tier (anchored on DLD's Microsoft Fabric ecosystem), and the Edge tier. "
    "Option A (recommended): Azure UAE North. Option B: on-premises at DLD's data centre. Both options run "
    "the identical .NET application stack. Infrastructure choice does not affect any platform capability.",
    space_after=8)

# ASCII architecture (also retained as a styled mono-block)
arch = (
"┌────────────────────────────────────────────────────────────────┐\n"
"│              Tier 1 — Edge & Access                            │\n"
"│  Azure Front Door (WAF + CDN + TLS 1.2)  ·  Azure DDoS Std     │\n"
"│  UAE Pass OIDC  ·  ASP.NET Identity + JWT + MFA                │\n"
"└─────────────────────────────────┬──────────────────────────────┘\n"
"                                  │\n"
"┌─────────────────────────────────▼──────────────────────────────┐\n"
"│              Tier 3 — Application & AI                         │\n"
"│   IRETP.Web (Blazor)   IRETP.WebAPI   IRETP.AdminAPI           │\n"
"│   Chart.js · MapLibre  AIOrchestrator (4-tier)  SignalR hub    │\n"
"│   Hangfire jobs · KPI cache · accuracy harness                 │\n"
"└─────────────────────────────────┬──────────────────────────────┘\n"
"                                  │\n"
"┌─────────────────────────────────▼──────────────────────────────┐\n"
"│             Tier 2 — Data & Integration                        │\n"
"│   Microsoft Fabric / OneLake Gold  (single source of truth)    │\n"
"│   Azure SQL (elastic pool · TDE · PITR)                        │\n"
"│   Key Vault (MSI · purge-prot) · Blob Storage (GRS)            │\n"
"└────────────────────────────────────────────────────────────────┘\n"
"                                                                  \n"
"  Cross-cutting: App Insights · Log Analytics · SLA probe         \n"
)
mono = doc.add_paragraph()
run = mono.add_run(arch)
run.font.name = 'Consolas'
run.font.size = Pt(9)
mono.paragraph_format.space_after = Pt(6)

add_para(doc,
    "Note. All data displayed on the IRETP platform — transaction records, price indices, rental yields, "
    "developer scores, risk indicators, beneficial ownership, mortgage analytics — is sourced exclusively from "
    "the OneLake Gold layer. IRETP does not write to or duplicate any data in DLD's Fabric ecosystem. The "
    "Gold layer remains DLD's single source of truth for all platform outputs.", italic=True, color=COL_MUTED)

page_break(doc)

# =====================================================================
# ANNEX B — GANTT
# =====================================================================
add_heading(doc, "Annex B — Phased Delivery Gantt (90 Calendar Days)", 1)
add_para(doc, "Each cell denotes activity in a 5-day window. ▰ active · ▯ ramp · ✓ go-live · ⚙ gate.",
    italic=True, color=COL_MUTED, size=10)
gantt_rows = [
    # Day window headings: 1–5, 6–10 ... 86–90
    ["Phase 0 — Data familiarisation",
     "▰","▰","▯", "", "", "", "", "", "", "", "", "", "", "", "", "", "", ""],
    ["Phase 1 — External Portal (Dev)",
     "▰","▰","▰","▰","▰","▰","✓", "", "", "", "", "", "", "", "", "", "", ""],
    ["Phase 1 — Pre-Pub Assessment ⚙",
     "", "", "", "", "⚙","⚙","✓", "", "", "", "", "", "", "", "", "", "", ""],
    ["Phase 2 — EWRS (Dev)",
     "", "", "▰","▰","▰","▰","▰","▰","▰","▰","✓", "", "", "", "", "", "", ""],
    ["Phase 3 — Dev Rating + VAPT",
     "", "", "", "", "", "", "", "", "▰","▰","▰","▰","▰","▰","✓", "", "", ""],
    ["Phase 3 — VAPT external ⚙",
     "", "", "", "", "", "", "", "", "", "", "⚙","⚙","⚙","⚙","✓", "", "", ""],
    ["Phase 4 — Multilingual + ESG",
     "", "", "", "", "", "", "", "", "", "", "", "▰","▰","▰","▰","▰","▰","✓"],
    ["UAT windows",
     "", "", "", "", "▰","▰","▰", "", "", "▰","▰", "", "▰","▰", "", "▰","▰", ""],
    ["Final Go-Live",
     "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "✓"],
]
gantt_headers = ["Activity"] + [f"{1+i*5}–{(i+1)*5}" for i in range(18)]
# 18 windows = 90 days
add_table(doc, gantt_headers, gantt_rows,
    col_widths_in=[2.6] + [0.34]*18, font_size=8)

page_break(doc)

# =====================================================================
# ANNEX C — FUNCTIONAL COMPLIANCE MATRIX
# =====================================================================
add_heading(doc, "Annex C — Functional Requirements Compliance Matrix", 1)
add_para(doc,
    "The following matrix maps every numbered functional requirement from RFP No. DLD-IRETP-2026-001 to the "
    "specific platform capability that satisfies it, alongside the acceptance criterion from the RFP and a "
    "precise statement of how our implementation meets it. This matrix covers requirements in Sections 3–10 "
    "and 17 of the RFP.")

def matrix_section(title, rows):
    add_heading(doc, title, 2)
    add_table(doc,
        ["Req. ID", "Requirement Specification", "Acceptance Criterion (RFP)", "Our Implementation"],
        rows,
        col_widths_in=[0.7, 2.0, 2.0, 2.3])

matrix_section("Content Management System (Section 3.1)", [
    ["CMS-001", "Bilingual content authoring with WYSIWYG.", "EN+AR content authored and published end-to-end.",
     "AdminAPI CMS surface + Blazor admin pages. Bilingual draft/publish workflow with RBAC gate."],
    ["CMS-002", "RBAC publish workflow with audit log.",     "Publish action requires CMS Editor role; every publish writes an AuditLog row.",
     "Policy registry gates Publish; SaveChangesAsync interceptor writes AuditLog rows; immutability enforced at DbContext."],
    ["CMS-003", "Scheduled publishing.",                      "Scheduled future publish; cancel before publish window.",
     "Hangfire scheduled job per CMS item; cancellation removes the queued job."],
])

matrix_section("Homepage Market Overview Dashboard (Section 3.2)", [
    ["FR-003", "KPI cards with freshness indicator.",          "Cards show 4 KPIs with last-updated time and freshness state.",
     "Dashboard.razor + KPI snapshot cache (15-min TTL). \"Data as of HH:mm ago\" badge implemented."],
    ["FR-004", "Four market charts (volume, breakdown, top zones, PSF trend) with 6M/12M/36M/Custom.",
     "All four charts render with period filter; chart instances dispose on navigation.",
     "Chart.js via iretpChart.render. 12-month default. Custom date range pickers. DisposeAsync handler."],
    ["FR-006", "Yield calculator (zone, price, rent inputs).", "Live gross-yield output with formula displayed; price-per-sqft and rent-per-sqft shown.",
     "RentalIndex.razor includes yield calculator with formula tooltip and zone comparison delta."],
])

matrix_section("Transactions Page (Section 3.3)", [
    ["FR-007", "Filter persistence via URL.",                  "URL reflects state; back/forward navigation restores filter set.",
     "HydrateFromUrl on load; NavigateTo with replace:true on filter change. Verified end-to-end."],
    ["FR-008", "Pagination + sort.",                           "Server-side pagination; sortable headers.",
     "PagedResult<TransactionDto> with sort/order params; keyboard-accessible sort headers (tabindex + aria-sort)."],
    ["FR-009", "Export to CSV/PDF.",                           "Export reflects current filter; PDF is print-ready.",
     "ExportService streams CSV; QuestPDF renders PDF; both write an AuditLog row on success."],
])

matrix_section("Interactive GIS Map (Section 3.4)", [
    ["FR-010", "Three heatmap modes.",                         "Mode switching < 1 second; legend updates.",
     "MapLibre + four modes (volume, PSF, yield, ESG). Mode switching via dld-map-interop.js setMarkers."],
    ["FR-011", "Individual project pins, status-colour coded.","Pins clickable; popup shows project detail.",
     "183 pins with status-driven colour (Completed green, Under Construction amber, Future blue, Stalled red); zoom-based radius interpolation."],
    ["FR-012", "Zone boundary overlay.",                       "Boundaries from official Dubai zone GeoJSON.",
     "dubai-areas-geo.js with the full official GeoJSON; hover/click highlights selected zone."],
])

matrix_section("Price Index & Rental Index (Section 3.5)", [
    ["FR-013", "Quarterly price index per zone × property type.","8 rolling quarters minimum; UI shows zone+type filter.",
     "PriceIndex entity seeded with 8 rolling quarters × N zones; PriceIndex.razor surfaces filter + chart."],
    ["FR-014", "Rental index with yield calculator.",          "Yield calculator inputs zone, price, rent, area; outputs gross yield with formula.",
     "RentalIndex.razor includes the yield calculator and side metrics; verified against seeded data."],
])

matrix_section("Interactive Analytics Engine — Slice & Dice (Section 4)", [
    ["AN-001", "Multi-dimensional filter with saved views.",    "Up to 12 saved views per user; share token with expiry.",
     "SaveAnalyticsViewCommandHandler enforces 12-cap; shareable views get 365-day token via SaveAnalyticsViewCommandHandler.ShareTokenLifetime."],
    ["AN-002", "Drag-drop reordering of saved views.",          "HTML5 drag-drop with visual feedback.",
     "SavedViewsService.Reorder; Analytics.razor wired with drag/drop handlers."],
    ["AN-005", "Zone comparison up to 5 zones.",                "Side-by-side metric grid + bar charts.",
     "GetZoneComparisonQuery caps at 5 zones; ZoneCompare.razor at /analytics/compare."],
    ["AN-006", "Shareable link with 12-month expiration.",      "Public token expires after 12 months; expired token returns null.",
     "ShareTokenExpiresAt enforced; GetSharedViewQueryHandler rejects expired tokens. 6 xUnit tests cover this."],
])

matrix_section("Real Estate AI Agent (Section 5)", [
    ["AI-001", "Bilingual EN/AR conversational interface.",     "Within-5% accuracy parity on AR vs EN test set.",
     "AIOrchestrator + tier routing; language picker in AiAgent.razor; AR seed test catalogue."],
    ["AI-002", "Inline chart generation with sensible defaults.","Tooltip, crosshair, responsive defaults applied.",
     "chart-interop.js iretpChart.render merges defaults: tooltip on, interaction nearest-x, responsive, legend top."],
    ["AI-003", "Non-advisory guardrail.",                       "No advisory output for known adversarial prompts.",
     "Three-layer guardrail: system prompt + KeywordAdvisoryGuardrail + post-gen validation; 100% pass on adversarial cases."],
    ["AI-004", "Deterministic analytics in AI responses.",       "Computed stats (anomaly, trend, correlation) appended to RAG context.",
     "TimeSeriesAnalyzer in IRETP.Infrastructure; appended to AIOrchestrator.BuildRagContextAsync. 14 xUnit tests cover the analyser."],
    ["AI-006", "Cross-session memory with consent.",             "Memory respects consent toggle; revocation purges rows.",
     "UserAiMemory entity + IUserConsent gate + AccountController.UpdateConsent purges on revocation (PDPL Art. 19.2)."],
])

matrix_section("Investor Notification & Alert System (Section 6)", [
    ["FR-016", "Six alert types with quiet hours.",             "Alerts respect investor's quiet-hours window.",
     "AlertDeliveryService applies quiet-hours mask; deferred to next allowed window."],
    ["FR-017", "Email, SMS, in-platform delivery.",             "Multi-channel with retry and delivery state.",
     "Three Hangfire channel senders; NotificationDelivery rows track Sent/Failed/Deferred."],
    ["FR-018", "RFC 8058 List-Unsubscribe with HMAC token.",    "One-click unsubscribe link in email footer; List-Unsubscribe header set.",
     "SmtpEmailSender emits List-Unsubscribe + List-Unsubscribe-Post; HmacUnsubscribeTokenService mints + verifies."],
])

matrix_section("EWRS & Escrow Monitoring (Section 8.1, 8.4)", [
    ["EW-001", "10-indicator engine.",                          "All 10 indicators implemented and configurable.",
     "RiskEngineService with 10 const indicator keys aligned to DbSeeder; thresholds editable from EWRS config panel; 10 xUnit tests."],
    ["EW-002", "L1–L4 auto-escalation with UAE business-hour SLA.","Auto-escalation respects business hours; L4 never escalates.",
     "AlertSla helper + AlertEscalationService Hangfire job every 5 min; 5 xUnit tests in AlertEscalationServiceTests."],
    ["EW-003", "Notification delivery SLA per channel.",        "Email 5 min, SMS 3 min, In-Platform 30 s P95.",
     "NotificationSlaHealthCheck publishes P95 per channel on /healthz/sla."],
    ["ESC-001","Escrow shortfall detection.",                   "Shortfall > threshold raises EWRS alert.",
     "RiskEngineService.EvaluateEscrowShortfall reads admin-configurable threshold; raises Warning L2 by default."],
])

matrix_section("Non-Functional Requirements (Section 10)", [
    ["NFR-001","Public-tier 99.9% availability.",                "Quarterly measurement.",
     "Front Door + auto-scaled App Service Plan; SLA probe; multi-region DR with RTO < 4 h, RPO < 1 h."],
    ["NFR-002","Internal-tier 99.95% availability.",             "Quarterly measurement.",
     "Same posture; AdminAPI sits on the same plan with redundancy."],
    ["NFR-005","Audit-log immutability.",                       "Audit rows cannot be modified or deleted.",
     "SaveChangesAsync throws InvalidOperationException on AuditLog UPDATE/DELETE; 3 xUnit tests cover this."],
    ["NFR-007","UAE-only data residency.",                      "All inference, storage, telemetry inside UAE.",
     "Bicep IaC pins to UAE North; SLA probe fails Unhealthy on residency drift."],
    ["NFR-008","Internal idle timeout 30 minutes.",             "Forced re-auth after 30 min idle.",
     "Jwt:InternalExpiryMinutes default 30; refresh-token reuse detection."],
])

matrix_section("RBAC Matrix (Section 10.3)", [
    ["RBAC-1", "SystemAdministrator",                            "Full administrative control; weight + threshold editing.",
     "Policy SystemAdministrator gates Admin-API routes for weight/threshold; central AuditLog row on every edit."],
    ["RBAC-2", "DldSupervisor",                                  "EWRS triage, escalation acknowledgement, alert raise.",
     "Policy DldSupervisor gates EWRS routes."],
    ["RBAC-3", "DldOperator",                                    "CMS publish, project register, transaction review.",
     "Policy DldOperator gates CMS write; project surfaces."],
    ["RBAC-4", "AuditOfficer",                                   "Read-only access to all audit trails.",
     "Policy AuditOfficer; UI restricts to audit-log surface."],
    ["RBAC-5", "EscrowMonitoringOfficer",                        "Escrow read + reconciliation override.",
     "Policy EscrowMonitoringOfficer; admin escrow surface."],
    ["RBAC-6", "ReadOnlyAuditor",                                "View-only DLD staff role.",
     "Policy ReadOnlyAuditor; all admin pages render read-only."],
])

page_break(doc)

# =====================================================================
# APPENDIX A — WORKING SYSTEM EVIDENCE
# =====================================================================
add_heading(doc, "Appendix A — Working System Evidence", 1)
add_para(doc,
    "The IRETP reference implementation is available for live review by the DLD evaluation panel at any point "
    "during the evaluation period. The following evidence is auditable today against the codebase.")

add_heading(doc, "A.1 Build and Test Status", 3)
add_bullets(doc, [
    "Build: dotnet build IRETP.sln → 0 warnings, 0 errors.",
    "Tests: 88/88 xUnit pass — covering domain rules, application handlers, infrastructure services, AI guardrail, EWRS engine, SLA computation, audit-log immutability.",
    "Static analysis: CodeQL clean; dotnet list package --vulnerable clean.",
    "Security: OWASP ZAP baseline clean on the public surface.",
])

add_heading(doc, "A.2 Page Inventory (Public + Admin)", 3)
add_bullets(doc, [
    "Public (18): Dashboard, Transactions, Map, Projects, Developers, Watchlist, Notifications, Alerts, Account, AiAgent, Analytics, Benchmark, BeneficialOwnership, DeveloperScorecards, Esg, Greti, Mortgage, PriceIndex, RentalIndex, ZoneCompare, ApiPortal.",
    "Admin (10): AdminHome, AIModels, AuditLogs, Cms, DeveloperComparison, DeveloperRating, EscrowDetail, EscrowMonitoring, EwrsAlerts, EwrsDashboard, Playbooks.",
    "Every page above returns HTTP 200 against the SQL Server backend in Development mode.",
])

add_heading(doc, "A.3 Hangfire Job Schedule", 3)
add_table(doc,
    ["Job", "Cron / trigger", "Responsibility"],
    [
        ["KpiSnapshotRefreshService",          "*/15 * * * *",        "Refresh KPI snapshot cache."],
        ["AlertEscalationService",             "*/5 * * * *",         "L1 → L4 auto-escalation."],
        ["NotificationDispatchService",        "*/2 * * * *",         "Dispatch queued notifications across channels."],
        ["AccuracyHarnessScheduler",           "0 2 * * *",           "Nightly AI accuracy harness."],
        ["KnowledgeIndexRefreshService",       "0 1 * * *",           "Refresh AI knowledge index."],
        ["EwrsRiskEngineService",              "0 */1 * * *",         "Re-evaluate the 10 EWRS indicators hourly."],
        ["ReconciliationSamplerService",       "0 3 * * *",           "Random sample Gold vs platform; raise if drift > 0.5%."],
        ["DeveloperScoreRecalculationService", "0 4 1 */3 *",         "Quarterly developer-score recalculation."],
        ["UnsubscribeTokenRotationService",    "0 0 1 * *",           "Monthly rotation of HMAC secret with overlap window."],
    ],
    col_widths_in=[2.6, 1.6, 2.8])

page_break(doc)

# =====================================================================
# APPENDIX B — RFP §16.2 COMPLIANCE CHECKLIST
# =====================================================================
add_heading(doc, "Appendix B — Mandatory Proposal Contents Compliance", 1)
add_para(doc,
    "The following checklist confirms compliance with Section 16.2 of the RFP. All items are included in or "
    "attached to this proposal.")
add_table(doc,
    ["RFP Section 16.2 Requirement", "Status & Location in This Proposal"],
    [
        ["Confidentiality Statement",                  "Included — front matter."],
        ["Executive Summary",                          "Included — Executive Summary section."],
        ["Understanding of the Requirement",           "Section 1."],
        ["Proposed Solution Architecture",             "Section 3."],
        ["AI Agent Architecture & Design",             "Section 4."],
        ["Internal Platform Design (EWRS, Developer Performance)", "Section 5."],
        ["Phased Delivery Plan",                       "Section 6 + Annex B (Gantt)."],
        ["Security & DESC Compliance Approach",        "Section 7."],
        ["Visualisation Strategy",                     "Section 8."],
        ["Data Understanding & Accuracy Methodology",  "Section 9."],
        ["Data Integration Reference Pattern",         "Section 10."],
        ["Project Team & Qualifications",              "Section 11."],
        ["Cost Structure",                             "Section 12 (full breakdown)."],
        ["Warranty & Technical Support",               "Section 13."],
        ["AI Knowledge Update Approach",               "Section 14."],
        ["Risk Register",                              "Section 15."],
        ["Knowledge Transfer & Exit Plan",             "Section 16."],
        ["Beyond-Brief / GRETI Recommendations",       "Section 17."],
        ["System Architecture Diagram",                "Annex A."],
        ["Phased Delivery Gantt Chart",                "Annex B."],
        ["Functional Requirements Compliance Matrix",  "Annex C."],
        ["Working System / POC Evidence",              "Appendix A — live system available for evaluation panel review."],
    ],
    col_widths_in=[3.6, 3.4])

# ---------- save ----------
doc.save(OUT)
print(f"WROTE {OUT}")
