"""Extract text from all four reports for cross-verification with implementation."""
import os
from pathlib import Path
import docx

REPORTS = Path(r"C:\Users\kalmi\IRETP\reports")
OUT = Path(r"C:\Users\kalmi\IRETP\tools\_extracted")
OUT.mkdir(parents=True, exist_ok=True)


def extract(path: Path) -> str:
    d = docx.Document(str(path))
    lines = []
    # Top-level paragraphs first
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = p.style.name if p.style is not None else ''
        if style.startswith('Heading'):
            lines.append(f"\n## [{style}] {t}\n")
        elif style.lower() == 'title':
            lines.append(f"\n# TITLE: {t}\n")
        else:
            lines.append(t)
    # Then all tables
    for ti, tbl in enumerate(d.tables, start=1):
        lines.append(f"\n--- TABLE {ti} ---")
        for r in tbl.rows:
            cells = [c.text.strip().replace('\n', ' / ') for c in r.cells]
            lines.append(" | ".join(cells))
        lines.append("--- END TABLE ---\n")
    return "\n".join(lines)


for f in sorted(REPORTS.glob("*.docx")):
    print(f"--- {f.name} ---")
    text = extract(f)
    out = OUT / (f.stem + ".txt")
    out.write_text(text, encoding="utf-8")
    print(f"  -> {out.name} ({len(text):,} chars, {text.count(chr(10)):,} lines)")
