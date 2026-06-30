"""
Generates two RFP response documents in the workspace root:
    - TechnicalProposal_IRETP.docx
    - ComplianceMatrix_IRETP.docx

Source RFPs:
    - requirements/2026 RFP طلب عرض سعر.pdf
    - requirements/Integrated_Real_Estate_Transparency_Platform_RFP_v1.3.pdf
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT_TECH = ROOT / "TechnicalProposal_IRETP.docx"
OUT_COMP = ROOT / "ComplianceMatrix_IRETP.docx"

BRAND_BLUE = RGBColor(0x0B, 0x3D, 0x91)
BRAND_GREY = RGBColor(0x44, 0x44, 0x44)
BRAND_LIGHT = RGBColor(0xE6, 0xEE, 0xF8)
ACCENT_GREEN = RGBColor(0x1E, 0x88, 0x3D)
ACCENT_AMBER = RGBColor(0xC8, 0x8A, 0x00)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def init_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = BRAND_BLUE
        s.font.bold = True
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dubai Land Department")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = BRAND_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RFP No. DLD-IRETP-2026-001")
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND_GREY

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = BRAND_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = BRAND_GREY

    for _ in range(8):
        doc.add_paragraph()

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    rows = [
        ("Programme",       "Integrated Real Estate Transparency Platform (IRETP)"),
        ("Submission Date", "May 2026"),
        ("Document Type",   title),
        ("Classification",  "Confidential — Vendor Proposal"),
    ]
    for i, (k, v) in enumerate(rows):
        c0, c1 = meta.rows[i].cells
        c0.width = Cm(5.5); c1.width = Cm(11.0)
        c0.text = k; c1.text = v
        for c in (c0, c1):
            _set_cell_borders(c)
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        c0.paragraphs[0].runs[0].bold = True
        _shade(c0, "E6EEF8")
    doc.add_page_break()


def h1(doc, text):  doc.add_heading(text, level=1)
def h2(doc, text):  doc.add_heading(text, level=2)
def h3(doc, text):  doc.add_heading(text, level=3)


def para(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def bullets(doc, items, style="List Bullet"):
    for it in items:
        doc.add_paragraph(it, style=style)



def kv_table(doc, rows, col_widths=(5.5, 11.0)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(col_widths[0]); c1.width = Cm(col_widths[1])
        c0.text = k; c1.text = v
        for c in (c0, c1):
            _set_cell_borders(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        c0.paragraphs[0].runs[0].bold = True
        _shade(c0, "E6EEF8")
    doc.add_paragraph()


def matrix_table(doc, headers, rows, widths_cm):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.width = Cm(widths_cm[j])
        c.text = h
        _shade(c, "0B3D91")
        _set_cell_borders(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.width = Cm(widths_cm[j])
            c.text = val
            _set_cell_borders(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
            if i % 2 == 0:
                _shade(c, "F5F8FD")
    doc.add_paragraph()


def code_ref(doc, label, path):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = BRAND_GREY
    r2 = p.add_run(path)
    r2.font.name = "Consolas"
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0x00, 0x33, 0x88)


def callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    c.width = Cm(16.5)
    _shade(c, "E6EEF8")
    _set_cell_borders(c)
    p = c.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = BRAND_BLUE
    r.font.size = Pt(11)
    p2 = c.add_paragraph(body)
    for r in p2.runs:
        r.font.size = Pt(10)
    doc.add_paragraph()



# ---------------------------------------------------------------------------
# Technical Proposal builder
# ---------------------------------------------------------------------------
def build_technical_proposal() -> None:
    from content_technical import EXEC_SUMMARY, ARCH_OVERVIEW, SECTIONS
    doc = Document()
    init_styles(doc)

    add_cover(
        doc,
        title="Technical Solution Proposal",
        subtitle="Integrated Real Estate Transparency Platform (IRETP) — "
                 "RFP DLD-IRETP-2026-001 Response, Volume 1 of 2",
    )

    h1(doc, "1. Executive Summary")
    para(doc, EXEC_SUMMARY)
    callout(
        doc,
        "Why this proposal is differentiated",
        "The IRETP platform exists today as a Clean-Architecture .NET 8 "
        "codebase with feature folders, identity, AI orchestration, EWRS, "
        "Escrow monitoring, Developer Rating, CMS and Open Data API already "
        "in place. The vendor is not proposing greenfield development — the "
        "proposal answers each RFP clause with a live module, file path and "
        "deployment plan, plus an explicit list of value-add engineering "
        "improvements (Section 18) committed at no extra cost."
    )

    h1(doc, "2. Solution Architecture Overview")
    para(doc, ARCH_OVERVIEW)
    kv_table(doc, [
        ("Web Host (Blazor Server)", "src/IRETP.Web — public + admin UI"),
        ("Public/Investor API",      "src/IRETP.WebAPI — JWT + OIDC, rate-limited"),
        ("Internal/Admin API",       "src/IRETP.AdminAPI — DLD operational endpoints"),
        ("Application (CQRS)",       "src/IRETP.Application/Features/* (MediatR)"),
        ("Domain (DDD)",             "src/IRETP.Domain/Entities + Enums"),
        ("Infrastructure",           "src/IRETP.Infrastructure — EF Core, OneLake, Hangfire, AI, Identity"),
    ])

    # All RFP-aligned sections.
    for heading, intro, kv_rows, bullet_list, refs, callout_pair in SECTIONS:
        h1(doc, heading)
        para(doc, intro)
        if kv_rows:
            h3(doc, "Implementation Mapping")
            kv_table(doc, kv_rows)
        if bullet_list:
            h3(doc, "Highlights")
            bullets(doc, bullet_list)
        if refs:
            h3(doc, "Code References")
            for label, path in refs:
                code_ref(doc, label, path)
        if callout_pair:
            doc.add_paragraph()
            callout(doc, callout_pair[0], callout_pair[1])

    # Closing section
    h1(doc, "19. Closing Statement")
    para(doc,
         "The vendor confirms full understanding of RFP DLD-IRETP-2026-001 "
         "and commits to deliver the Integrated Real Estate Transparency "
         "Platform on the phasing, performance, security and data-quality "
         "terms described above. The accompanying Compliance Matrix "
         "(Volume 2) provides a clause-by-clause traceability statement.")
    para(doc, "Submitted with respect to Dubai Land Department.", italic=True)

    doc.save(OUT_TECH)
    print(f"Wrote {OUT_TECH}")


# ---------------------------------------------------------------------------
# Compliance Matrix builder
# ---------------------------------------------------------------------------
def build_compliance_matrix() -> None:
    from content_compliance import GROUPED
    doc = Document()
    init_styles(doc)

    add_cover(
        doc,
        title="RFP Compliance Matrix",
        subtitle="Integrated Real Estate Transparency Platform (IRETP) — "
                 "RFP DLD-IRETP-2026-001 Response, Volume 2 of 2",
    )

    h1(doc, "1. How to Read this Matrix")
    para(doc,
         "Every numbered or named requirement in the RFP is reproduced in "
         "this matrix together with a compliance status, the implementing "
         "component in the IRETP codebase, the delivery phase, and any "
         "clarifying notes.")
    kv_table(doc, [
        ("C  — Compliant",
         "Capability is implemented in the codebase today. The Implementing "
         "Component column points to the live file or service."),
        ("CD — Compliant by Delivery",
         "Capability is in the phase plan; required foundations exist and "
         "the work is scoped, estimated and committed."),
        ("E  — Enhanced (vendor commitment)",
         "Capability goes beyond the literal requirement. Listed in the "
         "Vendor Commitments group as a contractual obligation."),
    ])

    h1(doc, "2. Compliance Summary")
    summary_rows = []
    for group_title, rows in GROUPED:
        c  = sum(1 for r in rows if r[2] == "C")
        cd = sum(1 for r in rows if r[2] == "CD")
        e  = sum(1 for r in rows if r[2] == "E")
        summary_rows.append([group_title, str(c), str(cd), str(e), str(len(rows))])
    matrix_table(
        doc,
        ["RFP Section", "C", "CD", "E", "Total"],
        summary_rows,
        widths_cm=[8.5, 1.6, 1.6, 1.6, 2.0],
    )

    h1(doc, "3. Detailed Compliance Matrix")
    for group_title, rows in GROUPED:
        h2(doc, group_title)
        matrix_table(
            doc,
            ["RFP Ref", "Requirement", "Status", "Implementing Component", "Phase", "Notes"],
            [list(r) for r in rows],
            widths_cm=[1.6, 5.2, 1.4, 4.6, 1.2, 3.5],
        )

    h1(doc, "4. Acceptance Statement")
    para(doc,
         "The vendor confirms that every requirement of RFP "
         "DLD-IRETP-2026-001 is addressed by this Compliance Matrix. Items "
         "marked C are demonstrable today against the IRETP codebase. Items "
         "marked CD are committed as part of the phased delivery plan with "
         "the foundations already in place. Items marked E are engineering "
         "commitments offered at no additional cost to strengthen the "
         "platform's operability and DLD's NFR posture.")

    doc.save(OUT_COMP)
    print(f"Wrote {OUT_COMP}")


if __name__ == "__main__":
    build_technical_proposal()
    build_compliance_matrix()
