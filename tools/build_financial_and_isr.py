"""
Generates two additional RFP response documents in the workspace root:
    - FinancialProposal_IRETP.docx  (RFP §14 cost schedules)
    - AppendixA_DESC_ISR_IRETP.docx (DESC ISR v3 control mapping)
"""
from __future__ import annotations
import sys, os
sys.path.insert(0, os.path.dirname(__file__))

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT     = Path(__file__).resolve().parents[1]
OUT_FIN  = ROOT / "FinancialProposal_IRETP.docx"
OUT_ISR  = ROOT / "AppendixA_DESC_ISR_IRETP.docx"

BRAND_BLUE  = RGBColor(0x0B, 0x3D, 0x91)
BRAND_GREY  = RGBColor(0x44, 0x44, 0x44)
GREEN_DARK  = RGBColor(0x1E, 0x7E, 0x34)
AMBER_DARK  = RGBColor(0x9A, 0x6B, 0x00)
RED_DARK    = RGBColor(0x8B, 0x00, 0x00)

STATUS_COLORS = {
    "I":  ("D4EDDA", GREEN_DARK),
    "PI": ("FFF3CD", AMBER_DARK),
    "P":  ("F8D7DA", RED_DARK),
    "C":  ("D4EDDA", GREEN_DARK),
    "CD": ("D1ECF1", RGBColor(0x0C, 0x5C, 0x60)),
    "E":  ("E6EEF8", BRAND_BLUE),
}


def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    tc_pr.append(borders)


def init_styles(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    for lvl, sz in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(sz)
        s.font.color.rgb = BRAND_BLUE
        s.font.bold = True
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)


def add_cover(doc, title, subtitle):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dubai Land Department")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = BRAND_BLUE

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RFP No. DLD-IRETP-2026-001")
    r.font.size = Pt(11); r.font.color.rgb = BRAND_GREY

    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = BRAND_BLUE

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True; r.font.size = Pt(13); r.font.color.rgb = BRAND_GREY

    for _ in range(6): doc.add_paragraph()
    meta = doc.add_table(rows=3, cols=2)
    meta.autofit = False
    for i, (k, v) in enumerate([
        ("Programme",      "Integrated Real Estate Transparency Platform (IRETP)"),
        ("Document Date",  "May 2026"),
        ("Classification", "Confidential — Vendor Proposal"),
    ]):
        c0, c1 = meta.rows[i].cells
        c0.width = Cm(5.5); c1.width = Cm(11.0)
        c0.text = k; c1.text = v
        for c in (c0, c1):
            _borders(c)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
        c0.paragraphs[0].runs[0].bold = True
        _shade(c0, "E6EEF8")
    doc.add_page_break()


def h1(doc, t): doc.add_heading(t, level=1)
def h2(doc, t): doc.add_heading(t, level=2)
def h3(doc, t): doc.add_heading(t, level=3)
def para(doc, t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(10.5)


def callout_box(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    c.width = Cm(16.5); _shade(c, "E6EEF8"); _borders(c)
    p = c.paragraphs[0]
    r = p.add_run(title); r.bold = True; r.font.color.rgb = BRAND_BLUE; r.font.size = Pt(11)
    p2 = c.add_paragraph(body)
    for r in p2.runs: r.font.size = Pt(10)
    doc.add_paragraph()


def cost_table(doc, headers, rows, widths_cm, status_col=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.width = Cm(widths_cm[j]); c.text = h
        _shade(c, "0B3D91"); _borders(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)
    for i, row in enumerate(rows, start=1):
        is_total = any(str(v).upper().startswith("TOTAL") or str(v).upper().startswith("PHASE") for v in row[:2])
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.width = Cm(widths_cm[j]); c.text = str(val)
            _borders(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if is_total:
                        r.bold = True
            if is_total:
                _shade(c, "E6EEF8")
            elif i % 2 == 0:
                _shade(c, "F5F8FD")
            if status_col is not None and j == status_col:
                code = str(val).strip()
                bg, fg = STATUS_COLORS.get(code, ("FFFFFF", BRAND_GREY))
                _shade(c, bg)
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = fg
                        r.bold = True
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Financial Proposal builder
# ---------------------------------------------------------------------------
def build_financial_proposal():
    from content_financial import (
        INFRA_CAPEX, SW_LICENCES, DEV_COSTS, OPEX_ANNUAL, MILESTONES,
        COST_SUMMARY,
    )
    doc = Document()
    init_styles(doc)
    add_cover(
        doc,
        "Financial Proposal",
        "Integrated Real Estate Transparency Platform (IRETP) — "
        "RFP DLD-IRETP-2026-001 Response, Volume 3 of 3",
    )

    h1(doc, "1. Preamble & Disclaimer")
    para(doc,
         "All costs in this document are indicative ranges based on current "
         "market rates (UAE region, May 2026). Final binding figures will be "
         "provided in the vendor's official bid submission. All amounts are "
         "in UAE Dirhams (AED) exclusive of VAT. VAT at 5 % will be added to "
         "the final invoices. The cost structure follows the format mandated "
         "by RFP §14.")

    h1(doc, "2. Executive Cost Summary")
    cost_table(
        doc,
        ["Cost Category", "Amount (AED)"],
        [[k, v] for k, v in COST_SUMMARY],
        widths_cm=[10.0, 6.5],
    )
    callout_box(
        doc,
        "Hosting model selection",
        "Infrastructure CAPEX and annual OPEX differ by hosting option. "
        "Option A (Azure UAE North) is recommended for fastest time-to-production, "
        "built-in BCDR and lowest CAPEX. Option B (On-Premises) has higher CAPEX "
        "but may be preferred for data-sovereignty assurance. Option C (Hybrid) "
        "balances both objectives. DLD selects the model at contract signature "
        "and the financial schedules below are adjusted accordingly.",
    )

    h1(doc, "3. Infrastructure & Hosting Setup Costs (RFP §14.1 — One-Time CAPEX)")
    para(doc,
         "The table below covers all hosting components required to support "
         "the three application hosts (IRETP.WebAPI, IRETP.AdminAPI, IRETP.Web), "
         "the OneLake lakehouse, Redis cache, WAF, DR replication and security "
         "monitoring — across the three hosting options.")
    cost_table(
        doc,
        INFRA_CAPEX["headers"],
        INFRA_CAPEX["rows"],
        widths_cm=INFRA_CAPEX["widths_cm"],
    )

    h1(doc, "4. Software Licences & Subscriptions (RFP §14.2 — First Year)")
    para(doc,
         "Licences required for the platform's first year of operation. "
         "Open-source components (MIT/Apache 2) carry no licence cost. "
         "AI model API costs are estimated at current token pricing and "
         "will be adjusted at contract based on agreed usage tiers.")
    cost_table(
        doc,
        SW_LICENCES["headers"],
        SW_LICENCES["rows"],
        widths_cm=SW_LICENCES["widths_cm"],
    )

    h1(doc, "5. Application Development & Implementation (RFP §14.3 — By Phase)")
    para(doc,
         "Development costs are broken down by phase gate deliverable. "
         "Rates are based on UAE-market blended day rates: "
         "Architect AED 2,200/day, Senior Developer AED 1,800/day, "
         "Mid Developer AED 1,500/day, QA/DevOps AED 1,400/day.")
    cost_table(
        doc,
        DEV_COSTS["headers"],
        DEV_COSTS["rows"],
        widths_cm=DEV_COSTS["widths_cm"],
    )

    h1(doc, "6. Annual Recurring Costs (RFP §14.4 — OPEX)")
    para(doc,
         "Annual operating costs including hosting, AI model consumption, "
         "software renewals, notification channels, monitoring, security "
         "scanning and the 12-month post-go-live warranty & support package.")
    cost_table(
        doc,
        OPEX_ANNUAL["headers"],
        OPEX_ANNUAL["rows"],
        widths_cm=OPEX_ANNUAL["widths_cm"],
    )

    h1(doc, "7. Payment Milestone Schedule")
    para(doc,
         "Payments are tied to DLD-signed acceptance certificates at each "
         "phase gate, ensuring the vendor is financially motivated to deliver "
         "on time and to the acceptance criteria defined in RFP §17.")
    cost_table(
        doc,
        ["Milestone", "Trigger Event", "% of Dev CAPEX", "Amount (AED)"],
        MILESTONES,
        widths_cm=[1.5, 9.0, 2.8, 3.2],
    )

    h1(doc, "8. Value-Add Items (Included at No Extra Charge)")
    para(doc,
         "The 15 engineering improvements listed in the Technical Proposal "
         "(§18) — MediatR validation behavior, global exception handler, "
         "Redis caching, domain events, outbox pattern, SignalR hub, AI "
         "streaming, Polly resilience, Hangfire idempotency, OpenTelemetry, "
         "health checks, security headers, architecture tests, CI/CD pipelines, "
         "resource-based authorization — are included within the Phase 1 "
         "development budget and carry no additional line cost.",
         bold=False)
    callout_box(
        doc,
        "Value delivered",
        "At the indicative total CAPEX of AED 2,577,700 (Option A), DLD "
        "receives a production-ready, enterprise-grade IRETP platform that "
        "satisfies every clause of the RFP, with 15 additional engineering "
        "commitments that protect the platform's long-term operability — "
        "all within a phase-gated, pay-on-acceptance payment model.",
    )

    doc.save(OUT_FIN)
    print(f"Wrote {OUT_FIN}")


# ---------------------------------------------------------------------------
# DESC ISR v3 Appendix A builder
# ---------------------------------------------------------------------------
def build_isr_appendix():
    from content_isr import CONTROLS
    doc = Document()
    init_styles(doc)
    add_cover(
        doc,
        "Appendix A — DESC ISR v3 Control Mapping",
        "Integrated Real Estate Transparency Platform (IRETP) — "
        "RFP DLD-IRETP-2026-001 Security Compliance Evidence",
    )

    h1(doc, "1. Purpose & Scope")
    para(doc,
         "This document provides a clause-by-clause mapping of the Dubai "
         "Electronic Security Centre (DESC) Information Security Regulation "
         "version 3 (ISR v3) controls to the IRETP platform's implementation. "
         "It serves as the vendor's pre-VAPT self-assessment and supports the "
         "DESC compliance submission required before Phase 3 acceptance "
         "(RFP §10.2.1).")
    para(doc,
         "Status definitions:",
         bold=True)
    for code, label in [
        ("I",  "Implemented — control is active in the codebase / hosting layer today."),
        ("PI", "Partially Implemented — core control is in place; a process or "
               "configuration item is documented for completion in the operational "
               "runbook (Phase 1)."),
        ("P",  "Planned — control will be fully in place before Phase 3 VAPT. "
               "Owner and due date tracked in the project risk register."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{code}: ")
        r.bold = True; r.font.size = Pt(10.5)
        p.add_run(label).font.size = Pt(10.5)

    h1(doc, "2. Control Summary")
    domains = sorted({c[1] for c in CONTROLS})
    summary_rows = []
    for d in domains:
        dc = [c for c in CONTROLS if c[1] == d]
        i  = sum(1 for c in dc if c[3] == "I")
        pi = sum(1 for c in dc if c[3] == "PI")
        p  = sum(1 for c in dc if c[3] == "P")
        summary_rows.append([d, str(i), str(pi), str(p), str(len(dc))])
    cost_table(
        doc,
        ["Control Domain", "Implemented", "Partial", "Planned", "Total"],
        summary_rows,
        widths_cm=[5.0, 2.5, 2.0, 2.0, 2.0],
    )

    h1(doc, "3. Detailed Control Mapping")
    current_domain = None
    domain_rows = []
    all_domain_rows = []

    for ctrl in CONTROLS:
        ctrl_id, domain, title, status, impl, evidence = ctrl
        if current_domain is None:
            current_domain = domain
        if domain != current_domain:
            all_domain_rows.append((current_domain, domain_rows[:]))
            domain_rows = []
            current_domain = domain
        domain_rows.append([ctrl_id, title, status, impl, evidence])
    all_domain_rows.append((current_domain, domain_rows))

    for domain_name, rows in all_domain_rows:
        h2(doc, f"Domain: {domain_name}")
        cost_table(
            doc,
            ["Control ID", "Control Title", "Status", "IRETP Implementation", "Evidence"],
            rows,
            widths_cm=[1.6, 3.5, 1.4, 6.0, 4.0],
            status_col=2,
        )

    h1(doc, "4. Residual Risks & Remediation Plan")
    risk_rows = [
        ["R-1", "Partially Implemented controls",
         "3 controls (1.4, 2.3, 6.7) marked PI",
         "Medium", "Operational runbook and formal process completed by Phase 1 M1"],
        ["R-2", "Planned controls pre-VAPT",
         "3 controls (3.1, 3.2, 9.2) marked P",
         "Low", "Completed before Phase 3 VAPT; tracked in risk register"],
        ["R-3", "AI model provider UAE residency",
         "Depends on provider UAE endpoint availability",
         "Medium", "Provider contractually bound to UAE inference; UAE-hosted "
                   "open-source fallback (Llama/Mistral) available"],
    ]
    cost_table(
        doc,
        ["Risk ID", "Risk Description", "Current State", "Likelihood", "Remediation"],
        risk_rows,
        widths_cm=[1.5, 3.5, 3.2, 2.0, 6.3],
    )

    h1(doc, "5. Acceptance Statement")
    para(doc,
         "The vendor certifies that this ISR v3 self-assessment is accurate "
         "as of the proposal submission date and commits to resolving all "
         "Partial and Planned controls before the Phase 3 DESC-authorised "
         "VAPT engagement. The vendor will engage a DESC-approved penetration "
         "testing firm and submit the final VAPT report as a Phase 3 "
         "acceptance deliverable.")

    doc.save(OUT_ISR)
    print(f"Wrote {OUT_ISR}")


if __name__ == "__main__":
    build_financial_proposal()
    build_isr_appendix()

