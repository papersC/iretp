"""
DLD-IRETP-2026-001 Technical & Commercial Proposal v2.

Enhanced from v1 with:
  - Embedded high-fidelity screenshots of the live IRETP system
  - Full RFP §3-§20 Compliance Matrix with C/CD/E status (22 sub-tables)
  - DESC ISR v3 35-control mapping (10 domains)
  - 15 Vendor Commitments / Value-Add list
  - RACI Responsibility matrix (DLD · Vendor · Partner)
  - Three hosting options (Azure UAE / On-Prem / Hybrid)
  - Hardware sizing per option
  - Per-section "Code References" callouts
  - Hangfire job schedule table
  - Test strategy table
  - Roll-up Compliance List Acknowledgement (CLA)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r"C:\Users\kalmi\IRETP\DLD-IRETP-2026-TechnicalCommercialProposal_v2.docx"
SCREENS = r"C:\Users\kalmi\IRETP\proposal_screens"

# ---------- palette ----------
COL_INK   = RGBColor(0x0F, 0x1B, 0x2D)
COL_DLD   = RGBColor(0x06, 0x67, 0x35)
COL_TEXT  = RGBColor(0x1A, 0x1A, 0x1A)
COL_MUTED = RGBColor(0x6A, 0x72, 0x80)
COL_RULE  = RGBColor(0xD9, 0xDD, 0xD7)
COL_HEAD  = RGBColor(0xFF, 0xFF, 0xFF)
COL_BAND  = RGBColor(0xF1, 0xF5, 0xF1)
COL_GOLD  = RGBColor(0xC9, 0x9A, 0x2E)
COL_WARN  = RGBColor(0xB8, 0x50, 0x42)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D9DDD7", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),size); b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        borders.append(b)
    tcPr.append(borders)

def add_para(doc, text, bold=False, italic=False, size=10.5, color=COL_TEXT, align=None, space_after=4, space_before=0, font="Calibri"):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text); run.font.name = font; run.font.size = Pt(size); run.font.color.rgb = color
    run.bold = bold; run.italic = italic
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}; colors = {1: COL_DLD, 2: COL_INK, 3: COL_INK}
    space_before = {1: 14, 2: 10, 3: 6}; space_after = {1: 6, 2: 4, 3: 3}
    size = sizes.get(level, 11); color = colors.get(level, COL_INK)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before.get(level, 8))
    p.paragraph_format.space_after = Pt(space_after.get(level, 4)); p.paragraph_format.keep_with_next = True
    p.style = doc.styles[f'Heading {min(level,3)}']
    run = p.add_run(text); run.font.name = 'Calibri'; run.font.size = Pt(size); run.font.color.rgb = color; run.bold = True
    return p

def add_bullets(doc, items, indent_cm=0.5, size=10.5):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(indent_cm); p.paragraph_format.space_after = Pt(2)
        p.text = ''
        run = p.add_run(it); run.font.name = 'Calibri'; run.font.size = Pt(size); run.font.color.rgb = COL_TEXT

def add_table(doc, headers, rows, col_widths_in=None, header_fill="066735", band=True, font_size=9.5):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if col_widths_in:
        for i,w in enumerate(col_widths_in):
            for c in table.columns[i].cells:
                c.width = Inches(w)
    for j,h in enumerate(headers):
        c = table.rows[0].cells[j]; c.text = ''
        shade(c, header_fill); set_cell_borders(c, color="066735")
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h); r.font.name='Calibri'; r.font.size=Pt(font_size+0.5); r.bold=True; r.font.color.rgb=COL_HEAD
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = table.rows[1+i].cells[j]; c.text=''
            if band and i % 2 == 1:
                shade(c, "F1F5F1")
            set_cell_borders(c)
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
            r = p.add_run(str(val)); r.font.name='Calibri'; r.font.size=Pt(font_size); r.font.color.rgb=COL_TEXT
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def page_break(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

def add_figure(doc, image_path, caption, width_in=6.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(image_path, width=Inches(width_in))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption); r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.font.color.rgb = COL_MUTED; r.italic = True

def add_code_refs(doc, refs):
    """Render a tinted callout box listing code file paths."""
    add_para(doc, "Code references", bold=True, size=10.5, color=COL_DLD, space_before=4, space_after=2)
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run("• " + r)
        run.font.name = 'Consolas'; run.font.size = Pt(9.5); run.font.color.rgb = COL_INK

# =====================================================================
# build
# =====================================================================
doc = Document()
for s in doc.sections:
    s.page_height=Cm(29.7); s.page_width=Cm(21.0)
    s.left_margin=Cm(2.0); s.right_margin=Cm(2.0); s.top_margin=Cm(2.0); s.bottom_margin=Cm(2.0)

style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10.5)
for lvl,size,color in [(1,18,COL_DLD),(2,14,COL_INK),(3,12,COL_INK)]:
    s = doc.styles[f'Heading {lvl}']; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=color

# COVER
for _ in range(2): doc.add_paragraph()
add_para(doc, "TECHNICAL & COMMERCIAL PROPOSAL", bold=True, size=24, color=COL_DLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc, "In Response to RFP No. DLD-IRETP-2026-001", size=14, color=COL_INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "Integrated Real Estate Transparency Platform (IRETP)", italic=True, size=14, color=COL_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para(doc, "Volume 1 — Technical Solution · Compliance · DESC ISR · Commercial", italic=True, size=12, color=COL_DLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
add_para(doc, "Submitted to:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Dubai Land Department", size=14, color=COL_DLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Government of Dubai", size=11, color=COL_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_para(doc, "Submitted by:", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "[Vendor Legal Entity Name]", size=14, color=COL_INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "UAE Registered | Dubai, United Arab Emirates", italic=True, size=11, color=COL_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
add_para(doc, "Submission Date: 18 May 2026", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Proposal Validity: 120 days from submission date", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Document Version: 2.0", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
add_para(doc, "Classification: CONFIDENTIAL", bold=True, size=11, color=COL_DLD, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break(doc)

# CONFIDENTIALITY
add_heading(doc, "Confidentiality Statement", 1)
add_para(doc, "This proposal and all information contained herein are submitted in strict confidence by [Vendor Legal Entity Name] in response to RFP No. DLD-IRETP-2026-001 issued by the Dubai Land Department. The contents of this document constitute confidential and proprietary information and are intended solely for the use of the Dubai Land Department's evaluation team. Disclosure of any part of this document to any third party without the prior written consent of [Vendor Legal Entity Name] is strictly prohibited.")
add_para(doc, "[Vendor Legal Entity Name] acknowledges and accepts the confidentiality obligations set out in Section 4.3 of the RFP. All information provided by DLD in connection with this procurement shall be treated as strictly confidential and shall not be disclosed or used for any purpose other than the preparation and submission of this proposal.")

# TOC
add_heading(doc, "Table of Contents", 1)
toc = [
    "Executive Summary",
    "1. Understanding of the Requirement",
    "2. Why This Proposal",
    "3. Proposed Solution Architecture",
    "4. Real Estate AI Agent — Architecture & Design",
    "5. Internal Platform — EWRS & Developer Performance",
    "6. Phased Delivery Plan",
    "7. Security & DESC Compliance",
    "8. Visualisation Strategy — Public and Internal",
    "9. Data Understanding & Accuracy Methodology",
    "10. Data Integration Reference Architecture",
    "11. Project Team & Responsibility Matrix",
    "12. Cost Structure",
    "13. Warranty & Technical Support — 12 Months Post Go-Live",
    "14. AI Agent Knowledge Update Approach",
    "15. Risk Register",
    "16. Knowledge Transfer & Exit Plan",
    "17. Beyond the Brief — GRETI Score Maximisation",
    "18. Vendor Commitments — 15 Engineering Improvements at No Extra Cost",
    "19. Quality Posture & Test Strategy",
    "Annex A — System Architecture Diagram",
    "Annex B — Phased Delivery Gantt (90 Calendar Days)",
    "Annex C — Functional Requirements Compliance Matrix",
    "Annex D — DESC ISR v3 Control Mapping (35 controls, 10 domains)",
    "Annex E — Hardware Sizing & Hosting Options",
    "Appendix A — Working System Evidence (with screenshots)",
    "Appendix B — Mandatory Proposal Contents Compliance",
    "Appendix C — Compliance List Acknowledgement (CLA)",
]
for entry in toc:
    add_para(doc, entry, size=11, space_after=3)
page_break(doc)

# EXECUTIVE SUMMARY
add_heading(doc, "Executive Summary", 1)
add_para(doc, "[Vendor Legal Entity Name] is pleased to present this comprehensive proposal to the Dubai Land Department (DLD) for the design, development, deployment, and post-go-live support of the Integrated Real Estate Transparency Platform (IRETP). We have read RFP No. DLD-IRETP-2026-001 v1.3 in full — every functional, non-functional, hosting, security, and compliance requirement — and we submit this proposal in the belief that [Vendor Legal Entity Name] is uniquely positioned to deliver this platform on time, to specification, and to a standard that will measurably advance Dubai's standing on the JLL Global Real Estate Transparency Index (GRETI).")
add_para(doc, "Our response is grounded in a working reference implementation already running today against the full RFP scope. The platform is a Clean-Architecture .NET 9 solution: Blazor Server portal, public WebAPI, internal AdminAPI, Hangfire orchestration, EF Core on SQL Server, Microsoft Fabric / OneLake as the analytical source of truth, MapLibre for GIS, and an Arabic-first AI agent with a UAE-resident tiered LLM orchestrator. The build is verified by 88 of 88 xUnit tests passing with zero warnings on a clean solution build, and the solution exposes 28 controllers (20 in WebAPI, 8 in AdminAPI), 23 public pages, 11 admin pages, and 9 Hangfire recurring jobs that respond 200 / run green today.")
add_para(doc, "What DLD will receive is the productionisation, the DLD-specific data integration, and the formal compliance certification of a working foundation — not a build from scratch.", space_after=8)

add_heading(doc, "Headline Commitments", 3)
add_table(doc, ["Commitment", "Target"], [
    ["Functional coverage of RFP mandatory requirements", "100% at Go-Live"],
    ["Bilingual EN/AR parity", "Full UI, Open Data, AI Agent, exports"],
    ["Data residency", "100% UAE (Azure UAE North primary, UAE Central DR)"],
    ["Public-tier availability SLA", "99.9% measured quarterly"],
    ["Internal-tier availability SLA", "99.95% measured quarterly"],
    ["AI Agent P95 response latency", "< 3.5 s text, < 8 s chart, < 15 s end-to-end"],
    ["Chart render (KPI dashboard)", "P95 ≤ 2 s via KpiSnapshotCache"],
    ["KPI dashboard freshness", "≤ 15 minutes (cache TTL, freshness badge in UI)"],
    ["EWRS Level-1 alert delivery", "< 5 minutes — auto-escalation L1→L4 every 5 min"],
    ["Open Data API availability", "99.9% with partitioned rate limits (60/240/600 req/min)"],
    ["Final Go-Live", "Day 90 from contract signing (30 days under RFP maximum)"],
    ["Source code & IP ownership", "100% transferred to DLD on contract signature"],
    ["5-year Total Cost of Ownership (Option A)", "AED ≈ 5.72M (excluding VAT)"],
    ["Vendor Commitments (value-add)", "15 engineering improvements, no extra cost"],
    ["DESC ISR v3 controls implemented today", "27 of 35 (3 partial, 5 planned pre-VAPT)"],
], col_widths_in=[3.0, 4.0])

add_heading(doc, "Investment Snapshot", 3)
add_table(doc, ["Cost category", "Amount (AED)", "Type"], [
    ["Infrastructure CAPEX (Option A — Azure UAE North)", "335,500", "One-time"],
    ["Infrastructure CAPEX (Option B — On-Premises)",      "1,414,500", "One-time"],
    ["Infrastructure CAPEX (Option C — Hybrid)",           "335,500",  "One-time"],
    ["Software Licences (Year 1)",                          "235,200",  "One-time"],
    ["Application Development (M1–M4, 1,095 person-days)", "2,007,000", "One-time"],
    ["Total one-time CAPEX (Option A)",                     "2,577,700", "—"],
    ["Annual OPEX (Option A, includes 12-month warranty)", "629,200",   "Recurring"],
    ["Annual OPEX (Option B)",                              "686,200",   "Recurring"],
    ["Annual OPEX (Option C)",                              "629,200",   "Recurring"],
    ["Indicative 5-year TCO (Option A)",                    "≈ 5,723,500", "Reference"],
], col_widths_in=[3.4, 1.8, 1.8])
add_para(doc, "All figures AED, exclusive of VAT. DLD selects the hosting option at contract signature; financial schedules adjust accordingly. Detailed breakdowns are in Section 12. The 15 value-add engineering commitments listed in Section 18 are included at no extra cost.", italic=True, size=9, color=COL_MUTED)

page_break(doc)

# =====================================================================
# SECTION 1 — UNDERSTANDING
# =====================================================================
add_heading(doc, "1. Understanding of the Requirement", 1)

add_heading(doc, "1.1 Strategic Context", 2)
add_para(doc, "DLD's objective extends well beyond building a website. IRETP is a strategic instrument designed to advance Dubai's position on the JLL GRETI — a biennial index that directly influences where over 80% of global direct commercial real estate investment flows. Every functional and technical decision in this proposal has been made with that strategic objective in mind.")
add_para(doc, "We have mapped every identified GRETI gap against the platform components we propose. We are confident that a correct implementation of IRETP will create measurable, auditable improvements across all six GRETI sub-indices within the first assessment cycle following Go-Live.", space_after=6)

add_table(doc, ["GRETI Sub-Index", "Identified Gap", "IRETP Component Addressing It"], [
    ["Performance Measurement",
     "Inconsistent quarterly KPIs, no public price/rental index, no developer accountability data.",
     "Public/Dashboard.razor + KpiSnapshotCache (15-min cache); Public/PriceIndex.razor + Public/RentalIndex.razor with yield calculator; Public/DeveloperScorecards.razor with 8-criterion composite."],
    ["Market Fundamentals",
     "Property-type breakdowns and transaction-type splits not surfaced publicly.",
     "Public/Transactions.razor (FR-007 URL-persisted filters); Dashboard market charts (FR-004 — monthly volume, type breakdown, top-5 zones, PSF dual-axis)."],
    ["Governance",
     "Beneficial ownership opacity; absence of public developer rating.",
     "Public/BeneficialOwnership.razor (Arabic owner names rendered RTL with lang=ar); Public/DeveloperScorecards.razor with weight transparency."],
    ["Sustainability",
     "ESG indicators not surfaced; no LEED/Estidama Pearl/BREEAM mapping.",
     "Public/Esg.razor with rubric explainer; ESG Coverage as 4th heatmap mode on the GIS map (dld-map-interop.js)."],
    ["Regulatory & Legal Framework",
     "Escrow, AML, and dispute-resolution processes not transparent to investors.",
     "Admin/EscrowMonitoring.razor + Admin/EscrowDetail.razor; CmsContent-served regulatory content."],
    ["Transaction Process",
     "Real-time pricing, debt-market intelligence, and yield analytics not available.",
     "Public/PriceIndex + Public/RentalIndex live; Public/Mortgage.razor with 5 KPIs (avg mortgage size, M:Tx ratio, LTV distribution)."],
], col_widths_in=[1.8, 2.4, 2.8])

add_heading(doc, "1.2 Platform Architecture Summary", 2)
add_para(doc, "IRETP comprises two integrated pillars, delivered across four phases within a 90 calendar day window. We have read and understood every section of RFP DLD-IRETP-2026-001 v1.3 in full, including all 55 deliverables in Section 13, all 10 UAT test categories in Section 17.3, all 14 cost-table line items in Section 14, and all architectural requirements in Section 11.", space_after=6)
add_table(doc, ["Pillar", "Description"], [
    ["External Public Portal",
     "Public investor surface: Market Overview, Transactions, Interactive GIS Map, Price/Rental Indices, Slice-and-Dice Analytics, AI Agent, Watchlist & Notifications, Open Data API, ESG, Public Developer Scorecard, Beneficial Ownership, Mortgage."],
    ["Internal Management Platform",
     "DLD-staff surface: EWRS 10-indicator engine with L1–L4 auto-escalation, Escrow Monitoring, Developer Performance & Rating, AI Model Performance, Audit Logs, CMS, AI orchestration admin, RBAC for 6 internal roles."],
    ["Open Data Layer",
     "Open Data API, OpenAPI 3.0 dual specs (public-filtered + full-v1), interactive console at /api-docs, ApiKey middleware with tiered rate limits (Free 60 · Plus 240 · Partner 600 req/min)."],
], col_widths_in=[2.0, 5.0])

add_heading(doc, "1.3 Critical Constraints We Have Fully Absorbed", 2)
add_bullets(doc, [
    "Pre-Publication Analytics Assessment gate. No data is published until the assessment is signed off by DLD's Data Liaison Officer. Hard-coded into our Phase 1 plan as a go/no-go gate (Days 28–32).",
    "AI non-advisory constraint. Enforced at three layers: system-prompt, runtime input filter (keyword + regex), post-generation validation with bilingual refusal copy. Exercised by the accuracy harness as mandatory UAT cases.",
    "Microsoft Fabric / OneLake. DLD's existing data architecture is the foundation, not an afterthought. We do not duplicate data outside Fabric. All platform visualisations and APIs are served from the Gold layer.",
    "UAE data residency. All AI inference, data storage, processing, backups, and DR replicas remain within UAE sovereign cloud regions regardless of which AI model is active. UAE-residency metadata is stamped on every AI snapshot for audit.",
    "DESC ISR v3. The security plan is built to this standard from Day 1, not retrofitted. Audit-log immutability is enforced at the IretpDbContext SaveChangesAsync layer; central audit-trail rows are written on every scoring-weight or risk-threshold change.",
])
page_break(doc)

# =====================================================================
# SECTION 2 — WHY THIS PROPOSAL
# =====================================================================
add_heading(doc, "2. Why This Proposal", 1)

add_heading(doc, "2.1 A Working Reference Implementation, Not a Concept", 2)
add_para(doc, "[Vendor Legal Entity Name] does not approach this RFP with a concept paper. We have already built a functional reference implementation that covers approximately 88% of the RFP scope today. The platform builds clean (0 warnings, 0 errors), passes 88 of 88 xUnit tests, and renders 23 public pages and 11 admin pages with HTTP 200 against a real SQL Server backend.")

# Insert dashboard screenshot
add_figure(doc, os.path.join(SCREENS, '01_dashboard.png'),
    "Figure 2.1 — Live Market Overview Dashboard. KPI cards, four FR-004 market charts (monthly volume, type breakdown, top-5 zones, PSF dual-axis), 15-minute KpiSnapshotCache, freshness badge.")

add_para(doc, "What is already built and working in the reference codebase:", space_after=4)
add_bullets(doc, [
    "23 public pages — Dashboard, Transactions, Map, Projects, Developers, Watchlist, Notifications, Alerts, Account, AiAgent, Analytics, Benchmark, BeneficialOwnership, DeveloperScorecards, Esg, Greti, Mortgage, PriceIndex, RentalIndex, ZoneCompare, ApiPortal, CmsPreview, AnalyticsEmbed.",
    "11 admin pages — AdminHome, AIModels, AuditLogs, Cms, DeveloperComparison, DeveloperRating, EscrowDetail, EscrowMonitoring, EwrsAlerts, EwrsDashboard, Playbooks.",
    "28 API controllers — 20 in WebAPI (Auth, AI Agent, Alerts, Analytics, Beneficial Ownership, CMS preview, Currency, Dashboard, Developer, ESG, Export, GRETI, Map, Mortgage, NameValidation, Notifications, OpenData, PriceIndex, RentalIndex, Transactions, Watchlist), 8 in AdminAPI (Audit, CMS, Developer Rating, EWRS, Escrow, Escrow Ingestion, AI Models, User Admin).",
    "9 Hangfire recurring jobs — KpiSnapshotRefreshService, AlertEscalationService, AlertDeliveryService, InvestorAlertEvaluator, AccuracyHarnessScheduler, EwrsRiskEngineService, ReconciliationSamplerService, BenchmarkRefreshService, CurrencyRatesRefreshService.",
    "88/88 xUnit tests passing — domain rules, application handlers, infrastructure services, AI guardrail, EWRS engine, SLA computation, audit-log immutability, share-token expiry, notification SLA probe.",
])
add_code_refs(doc, [
    "src/IRETP.Web/Components/Pages/Public/*.razor",
    "src/IRETP.Web/Components/Pages/Admin/*.razor",
    "src/IRETP.WebAPI/Controllers/*.cs",
    "src/IRETP.AdminAPI/Controllers/*.cs",
    "src/IRETP.Application/Features/*",
    "src/IRETP.Infrastructure/Services/*",
    "tests/IRETP.Tests/",
])

add_heading(doc, "2.2 Clean Architecture on a Microsoft-Native Stack", 2)
add_para(doc, "The platform is architected as a Clean-Architecture .NET 9 solution. Outer layers depend inward; the Domain layer has zero framework references. Strict inward dependency rules are enforced by NetArchTest architecture tests. This is the architecture the Microsoft developer ecosystem treats as the long-term default for enterprise-grade business systems.", space_after=6)
add_table(doc, ["Project", "Responsibility"], [
    ["IRETP.Domain", "Framework-free entities, value objects, enums, domain helpers (~30 entities including Transaction, Project, EscrowAccount, RiskAlert, ScoringWeight, BeneficialOwner, UserAiMemory, SavedAnalyticsView). Depends on nothing."],
    ["IRETP.Application", "MediatR commands, queries, DTOs, feature folders (AIAgent, Alerts, Analytics, Audit, Auth, Benchmark, CMS, Dashboard, DeveloperRating, EWRS, Escrow, Esg, Export, Greti, Map, Mortgage, NameValidation, Ownership, PriceIndex, RentalIndex, Transactions)."],
    ["IRETP.Infrastructure", "EF Core IretpDbContext + migrations, repository impls, ASP.NET Identity, Hangfire services, AIOrchestrator + RAG retriever + guardrails + accuracy harness, notification senders, audit-log service."],
    ["IRETP.WebAPI", "Public + investor HTTP surface (port 5000). Hosts SignalR for real-time alerts and the Hangfire dashboard."],
    ["IRETP.AdminAPI", "DLD-internal endpoints (port 5002). Threshold editing, scoring-weight configuration, AI model status, audit, escrow ingestion. Bound to private VNET."],
    ["IRETP.Web", "Blazor Server frontend (port 5010). Public portal + admin pages. Routes API calls through typed HttpClient services."],
    ["IRETP.Tests", "xUnit suite — 88 tests passing today covering domain, application, infrastructure, AI, EWRS."],
], col_widths_in=[2.0, 5.0])

add_heading(doc, "2.3 Source Ownership From Day One", 2)
add_para(doc, "Every line of code, every EF Core migration, every Bicep template, every Hangfire job definition, every test, and every architectural document becomes the exclusive intellectual property of the Dubai Land Department on the date of contract signature — not on Day 90, not at warranty expiry.")
add_para(doc, "There is no proprietary platform layer hidden inside the stack. The components are all standard Microsoft and open-source technologies — ASP.NET Core, Blazor Server, EF Core, SQL Server, Hangfire, SignalR, MapLibre, Chart.js, OpenAPI. Any qualified .NET team in the UAE can maintain this codebase without vendor involvement. DLD is not paying a platform subscription, a per-user licence, a per-tenant fee, or a knowledge-base index licence at any point during the platform's operational lifetime.")

add_heading(doc, "2.4 De-Risked Delivery Pathway", 2)
add_bullets(doc, [
    "The reference implementation removes construction risk on every Phase 1 feature. Phase 1 work is integrating the working portal with DLD's OneLake Gold layer, running the Pre-Publication Analytics Assessment, completing UAT, and going live.",
    "Microsoft Fabric / OneLake is consumed via the standard Fabric SQL endpoint and semantic model APIs. There is no proprietary connector to license, configure, or wait for. The data layer is a SQL connection string and a credential — both configured on Day 1.",
    "The 9 Hangfire jobs that drive KPI refresh, EWRS escalation, notification delivery, knowledge re-indexing, and the accuracy harness are already operating in the reference build. Phase 1 work on these is schedule configuration, not implementation.",
])

add_heading(doc, "2.5 UAE Government Operating Posture", 2)
add_para(doc, "All AI inference, data storage, telemetry, backups, and DR replicas are pinned to UAE sovereign cloud regions. The Bicep IaC under /infra/bicep/ pins every resource to Azure UAE North; the DR plan replicates to a physically separate UAE region. UAE Pass OIDC federation is wired into the authentication stack as an opt-in scheme. MFA is mandatory for every internal DLD user via the policy registry. PDPL Article 19.2 (right to erasure) is implemented end-to-end: opting out of AI memory purges UserAiMemory rows in the same transaction.")
add_code_refs(doc, [
    "infra/bicep/main.bicep — subscription scope, pinned UAE North",
    "infra/bicep/platform.bicep — platform components, MSI-enabled, Key Vault purge-protected",
    "src/IRETP.Infrastructure/Identity/UserConsentService.cs — PDPL §19.2 consent + erasure",
    "src/IRETP.Infrastructure/HealthChecks/SlaHealthCheck.cs — UAE-residency drift probe",
])

page_break(doc)

# =====================================================================
# SECTION 3 — SOLUTION ARCHITECTURE
# =====================================================================
add_heading(doc, "3. Proposed Solution Architecture", 1)

add_heading(doc, "3.1 Architecture Overview", 2)
add_para(doc, "The IRETP architecture is structured across four tiers with clearly defined responsibilities, data flows, and security boundaries. The architecture fully leverages DLD's existing Microsoft Fabric ecosystem: the External Public Portal's visualisation layer is a Blazor Server frontend consuming data from the OneLake Gold layer, consistent with Section 11.4.2 of the RFP.", space_after=4)
add_table(doc, ["Tier", "Responsibility"], [
    ["Tier 1 — Edge & Access", "Azure Front Door Premium (WAF, CDN, TLS 1.2 min), Azure DDoS Standard, geo-routing pinned to UAE North/Central. UAE Pass OIDC federation for investors; ASP.NET Identity + JWT + TOTP MFA for DLD staff."],
    ["Tier 2 — Data & Integration", "Microsoft Fabric / OneLake Gold layer as the analytical source of truth, consumed via the Fabric SQL endpoint. SQL Server (Azure SQL elastic pool) for OLTP. Hangfire orchestration for scheduled refresh, EWRS escalation, accuracy harness, knowledge re-indexing."],
    ["Tier 3 — Application & AI", "ASP.NET Core 9 on three hosts: public WebAPI (:5000), AdminAPI (:5002), Blazor Server Web (:5010). MediatR CQRS. AIOrchestrator with four-tier model routing. SignalR for real-time alerts. 9 Hangfire jobs across WebAPI + AdminAPI."],
    ["Tier 4 — Presentation & Surfaces", "Public users and registered investors access the External Portal via browser. DLD staff access the Internal Platform via browser with MFA. The Open Data API surfaces machine-readable data for researchers, PropTech developers, and financial institutions."],
], col_widths_in=[1.7, 5.3])

add_heading(doc, "3.2 Microsoft Fabric / OneLake Alignment", 2)
add_para(doc, "DLD's existing investment in the Microsoft Fabric ecosystem is the data foundation of this platform. The IRETP service tier connects directly to OneLake Gold via the Fabric SQL endpoint. From the IRETP service tier, OneLake is one entry in the connection-string configuration — no proprietary connector, no platform layer, no third-party intermediary.")
add_para(doc, "The medallion architecture (bronze → silver → gold) is honoured. IRETP reads only from the Gold layer.", space_after=4)
add_table(doc, ["Medallion Layer", "Purpose in IRETP", "What We Add"], [
    ["Bronze (raw)", "Not consumed. DLD source-of-truth registries, escrow bank feeds, RERA regulatory feeds, ESG certifications.", "—"],
    ["Silver (conformed)", "Not consumed. DLD's data quality rules (Great Expectations / dbt tests). Anomalies pause publication.", "—"],
    ["Gold (business-ready)", "EF Core reads the Gold layer for transactional UI. Hot KPIs synchronised into the OLTP store every 15 min by Hangfire.", "KpiSnapshotRefreshService · BenchmarkRefreshService · CurrencyRatesRefreshService — three Hangfire workers keep the Gold layer warm."],
], col_widths_in=[1.8, 2.6, 2.6])

add_heading(doc, "3.3 External Public Portal — Visualisation Layer", 2)
add_para(doc, "We agree with DLD's assessment in Section 11.4.2 that Power BI is not suitable as the primary visualisation engine for the External Public Portal. The licensing model does not scale cost-effectively to 10,000+ concurrent anonymous users; embed-token management adds operational overhead; and the UX customisation surface is constrained.")
add_para(doc, "Our External Portal visualisation strategy is a Blazor Server frontend with Chart.js for analytics and MapLibre for GIS:", space_after=4)

add_figure(doc, os.path.join(SCREENS, '02_map.png'),
    "Figure 3.1 — Interactive GIS Map. FR-010 four heatmap modes (volume, PSF, yield, ESG); FR-011 individual project pins colour-coded by status; official Dubai zone GeoJSON; MapLibre GL JS.")

add_table(doc, ["Dimension", "Power BI Embedded (A-series)", "IRETP: Blazor Server + Chart.js + MapLibre"], [
    ["Year-1 licence cost (public users)", "AED 260K – 1M+ (A4–A6 capacity)", "AED 0 — MIT / open-source"],
    ["Anonymous public access", "App-owns-data token broker + service principal", "Native — anonymous by default"],
    ["UI / UX customisation depth", "Canvas-bound; theme JSON only; custom visuals need TypeScript", "Full HTML/CSS/JS; same design tokens as rest of IRETP"],
    ["RTL (Arabic) parity", "Incomplete on several visual types", "Layout-level; every visual inherits page direction"],
    ["Concurrent users per AED 100K", "~50–100 (A5 capacity share)", "~500–1,500 (App Service Plan share)"],
    ["P95 chart render", "5–15 s cold without pre-warming", "≤ 2 s via KpiSnapshotCache + no-tracking projections"],
    ["Source control · test automation", "Binary .pbix; weak git diffs; minimal test tooling", "Razor + JS in .NET solution; xUnit + Playwright + CI"],
    ["DLD IP ownership (RFP §19.1)", "Microsoft-owned service surface", "Code vests in DLD on payment"],
], col_widths_in=[2.0, 2.5, 2.5])
add_para(doc, "Power BI is retained for Internal Platform components where named-user licensing is cost-appropriate: EWRS dashboards, Developer Rating dashboards, internal management reporting. This is a deliberate split — right tool for each audience.")

add_heading(doc, "3.4 Technology Stack Proposal", 2)
add_table(doc, ["Layer", "Technology", "Justification / NFR Alignment"], [
    ["Frontend",            ".NET 9 + Blazor Server, MapLibre GL JS, Chart.js / ApexCharts", "Native RTL/LTR; WCAG 2.1 AA; SSR posture for low-bandwidth public; zero charting licence cost."],
    ["Backend API",         "ASP.NET Core 9 controllers + MediatR + FluentValidation + Swashbuckle OpenAPI 3.0", "P95 API < 500 ms; partitioned rate limiting; standard .NET tooling."],
    ["AI / LLM",            "AIOrchestrator (4 tiers) + AIModelMetrics + KeywordAdvisoryGuardrail + AiAccuracyHarness; BYO-LLM by configuration", "Multi-model orchestration; configuration-only model swap; UAE residency metadata."],
    ["OLTP",                "Azure SQL Database (elastic pool); EF Core 9 migrations", "TDE; Always Encrypted for PDPL columns; GRS backup; PITR."],
    ["Lakehouse",           "Microsoft Fabric OneLake Gold via Fabric SQL endpoint", "DLD's single source of truth; non-duplicated."],
    ["GIS",                 "MapLibre GL JS + official Dubai zone GeoJSON", "Map initial load < 5 s; 4 heatmap modes; 511 individual project pins."],
    ["CMS",                 "CmsContent + CmsContentVersion entities; Admin/Cms.razor authoring; Public/CmsPreview.razor rendering", "Eliminates CMS licence cost; bilingual content; RBAC publish."],
    ["Auth",                "ASP.NET Identity + JWT (HS256) + TOTP MFA; UAE Pass / Azure AD / Keycloak OIDC dual scheme; refresh-token rotation; 30-min internal idle timeout", "Centralised identity; MFA mandatory; reuse detection."],
    ["Job Orchestration",   "Hangfire (open-source) — 9 jobs", "Persisted in SQL; dashboard at /hangfire; durable retries."],
    ["Realtime",            "SignalR over WebSocket (HTTPS fallback)", "EWRS alert push; investor notification ping."],
    ["IaC",                 "Bicep — main.bicep + platform.bicep, all UAE North-pinned, MSI-enabled, Key Vault purge-protected", "Reproducible deployments; CI-tested; delivered to DLD on Day 1."],
    ["Observability",       "App Insights + Log Analytics; health endpoints /healthz/live, /healthz/ready, /healthz/sla", "P95 latency, error rate, AI residency, KPI freshness — all measured."],
], col_widths_in=[1.4, 2.5, 3.1])

add_heading(doc, "3.5 Scalability & Performance Targets", 2)
add_bullets(doc, [
    "Public-tier availability: 99.9% measured quarterly.",
    "Internal-tier availability: 99.95% measured quarterly.",
    "P95 API response: < 500 ms under 10,000 concurrent sessions.",
    "P95 chart render: ≤ 2 s via KpiSnapshotCache + EF AsNoTracking projections + Blazor pre-rendering.",
    "P90 AI text response: < 8 seconds; P95 AI chart response: < 12 seconds.",
    "KPI dashboard data freshness: ≤ 15 minutes (cache TTL, signalled via freshness badge).",
    "EWRS Level-1 alert delivery: < 5 minutes; L2 within 4 business hours; L3 within 1 business hour; L4 immediate.",
])
add_code_refs(doc, [
    "src/IRETP.Infrastructure/Services/KpiSnapshotCache.cs · 15-min TTL",
    "src/IRETP.Infrastructure/Services/KpiSnapshotRefreshService.cs · Hangfire every 15 min",
    "src/IRETP.Infrastructure/HealthChecks/SlaHealthCheck.cs · aggregated SLO probe",
    "src/IRETP.Infrastructure/HealthChecks/NotificationSlaHealthCheck.cs · P95 per channel",
])

page_break(doc)

# =====================================================================
# SECTION 4 — AI AGENT
# =====================================================================
add_heading(doc, "4. Real Estate AI Agent — Architecture & Design", 1)

add_heading(doc, "4.1 Our Position on AI in This Platform", 2)
add_para(doc, "Artificial intelligence is among the most over-used terms in technology proposals. The distinction between a credible AI proposal and a marketing document rests not on the presence of AI, but on whether each AI capability has been placed where it creates demonstrable value — and whether the temptation to apply AI indiscriminately has been resisted.")
add_para(doc, "We make one commitment we consider the most important in this proposal: the non-advisory constraint is non-negotiable, technically enforceable, and UAT-verified. The AI Agent will retrieve, present, and analyse DLD data. It will never recommend properties, forecast prices as facts, or provide investment guidance. This is a hard architectural constraint enforced at three independent layers and exercised by an automated accuracy harness on every release.")

add_figure(doc, os.path.join(SCREENS, '04_ai_agent.png'),
    "Figure 4.1 — Real Estate AI Agent. RAG-grounded responses with citation chips, deterministic statistics appended on analytics queries, three-layer non-advisory guardrail, EN/AR language picker, tier-routed model selection.")

add_heading(doc, "4.2 Multi-Model Orchestration with Task-Aware Tier Routing", 2)
add_table(doc, ["Tier", "Use Case", "Default Model Posture", "Latency Budget"], [
    ["nav",       "DLD service navigation",                                   "Lightweight UAE-resident model",            "P95 < 2 s"],
    ["analytics", "Numerical / data-heavy retrieval, RAG over Gold",         "Higher-capability UAE-resident + stats",    "P95 < 5 s"],
    ["primary",   "Open-ended bilingual conversation",                        "Top-tier UAE-resident (Azure OpenAI)",      "P90 < 8 s"],
    ["secondary", "Fallback if primary fails",                                "Alternate UAE-resident provider",           "P90 < 10 s"],
], col_widths_in=[0.9, 2.6, 2.5, 1.0])
add_para(doc, "Tier selection and model identity are configurable from the AI Model admin surface. Changing the active model on any tier is a configuration update — no redeployment, no code change. New providers are added via configuration entry + Key Vault credential.")

add_heading(doc, "4.3 RAG with Deterministic Analytics", 2)
add_bullets(doc, [
    "Step 1 — Query parsing. The natural-language input is classified (data, navigation, analytical, comparative) and routed to the appropriate tier.",
    "Step 2 — Retrieval. The retrieval layer queries OneLake Gold for relevant transactions, projects, developer profiles, zone boundaries, RERA regulations, and DLD service procedures. Hybrid keyword + vector retrieval.",
    "Step 3 — Deterministic statistics. TimeSeriesAnalyzer computes anomalies (z-score), trends (least-squares with R² and flat-threshold), and correlations (Pearson with no-variance guard). Appended to RAG context.",
    "Step 4 — Augmentation. Retrieved data + stats injected into the system prompt with the non-advisory constraint and citation requirement.",
    "Step 5 — Generation. The model produces a response grounded exclusively in retrieved data. Every data point is cited with source identifier and timestamp.",
    "Step 6 — Validation. Post-generation scan rejects investment-advisory language. Flagged responses are reformulated or declined with bilingual refusal copy.",
])

add_heading(doc, "4.4 Non-Advisory Guardrail — Three Independent Layers", 2)
add_bullets(doc, [
    "Layer 1 — System prompt. Orchestrator's system prompt explicitly forbids advisory language; reinforced by example.",
    "Layer 2 — Input filter (KeywordAdvisoryGuardrail). Bilingual keyword catalogue + regex catalogue blocks queries that request recommendations or forecasts. Tightened in Pass 19 to catch \"by 20XX prices will\", \"prices are expected to rise/climb/jump/grow\", \"yields are projected\", \"forecast of AED N\".",
    "Layer 3 — Output validation. Post-generation scan inspects the response even when the input passed. Triggered responses replaced with bilingual refusal message and logged for audit.",
])

add_heading(doc, "4.5 Arabic NLP Capability", 2)
add_para(doc, "Phase 1 ships full conversational capability in both English and Arabic. Arabic NLP accuracy is required within 5% of English on the same standardised test set. Arabic zone, project, and developer names are validated against DLD's official records before go-live. All AI chat surfaces render RTL when the user's active language is Arabic.")

add_heading(doc, "4.6 Accuracy Harness — Continuous Quality Measurement", 2)
add_para(doc, "The platform ships an automated accuracy harness (IAiAccuracyHarness + AiAccuracyHarness in IRETP.Infrastructure). The harness runs a seed catalogue of 14 questions on every release — data retrieval, service navigation, adversarial guardrail probes, bilingual variants. The catalogue is extensible to the full 100-question DLD UAT catalogue. Each run produces a structured report with expected-keyword matching plus a manual review queue. Endpoint: POST /api/admin/ai-models/accuracy-test (SystemAdministrator only).")

add_heading(doc, "4.7 Cross-Session Memory with Consent (AI-006)", 2)
add_para(doc, "AI-006 requires the agent to remember a registered investor's preferences across sessions when they opt in. UserAiMemory (UserId + Kind [zone|topic] + Key + Frequency + LastUsedAt) with explicit consent at registration or any time after. Revocation purges every UserAiMemory row for the user in the same transaction (PDPL Article 19.2).")

add_code_refs(doc, [
    "src/IRETP.Infrastructure/Services/AIOrchestrator.cs",
    "src/IRETP.Infrastructure/Services/AIModelMetrics.cs",
    "src/IRETP.Infrastructure/Services/KeywordAdvisoryGuardrail.cs",
    "src/IRETP.Infrastructure/Services/AiAccuracyHarness.cs",
    "src/IRETP.Infrastructure/Services/TimeSeriesAnalyzer.cs",
    "src/IRETP.Infrastructure/Identity/UserConsentService.cs",
    "src/IRETP.Web/Components/Pages/Admin/AIModels.razor",
    "src/IRETP.Web/Components/Pages/Public/AiAgent.razor",
])

page_break(doc)

# =====================================================================
# SECTION 5 — INTERNAL PLATFORM
# =====================================================================
add_heading(doc, "5. Internal Platform — EWRS & Developer Performance", 1)

add_heading(doc, "5.1 EWRS — 10-Indicator Risk Engine", 2)
add_para(doc, "The Early Warning Risk System is the intelligence core of the Internal Management Platform. The engine evaluates ten configurable indicators across three scopes (project, zone, developer), each with admin-configurable thresholds and SLA-bound severity bands.", space_after=4)

add_figure(doc, os.path.join(SCREENS, '05_ewrs.png'),
    "Figure 5.1 — EWRS Alert Inbox & SLA Escalation Ladder. 10-indicator engine, L1–L4 ladder with UAE business-hour SLAs, AlertEscalationService Hangfire job every 5 min, immutable audit log on every threshold edit.")

add_table(doc, ["Indicator", "Default Threshold", "Risk Level", "Implementation"], [
    ["Project Delivery Delay",         "90 days vs original",                                   "Critical → L3", "RiskEngineService.EvaluateProjectDeliveryDelay"],
    ["Escrow Shortfall",                "20% below required balance",                            "Warning → L2",  "RiskEngineService.EvaluateEscrowShortfall"],
    ["Construction Suspension",         "30 days inactivity",                                    "L2 / L3",       "RiskEngineService.EvaluateConstructionSuspension"],
    ["Transaction Volume Decline",      "30-day count vs 12-mo rolling avg, −40%",              "L2",            "RiskEngineService.EvaluateTransactionVolumeDecline"],
    ["Price Decline (PSF QoQ)",         "QoQ −15% with ≥5 tx per bucket",                        "L2 / L3",       "RiskEngineService.EvaluatePriceDecline"],
    ["Developer Score Deterioration",   "QoQ composite drop > 20 pts",                           "L2; L3 if <40", "RiskEngineService.EvaluateDeveloperScore"],
    ["High-Risk Concentration",         "≥3 active projects w/ open High or Stalled",            "L2",            "RiskEngineService.EvaluateHighRiskConcentration"],
    ["Severe Regulatory Violation",     "Critical RERA violation within 90 days",                "L3",            "RiskEngineService.EvaluateSevereRegulatoryViolation"],
    ["AI Anomaly Flag",                 "TimeSeriesAnalyzer z-score threshold",                  "L1 / L2",       "RiskEngineService.EvaluateAiAnomaly"],
    ["Data-Feed Anomaly",               "Publication held until ack",                            "L1",            "ReconciliationSamplerService"],
], col_widths_in=[2.0, 1.5, 1.0, 2.5])

add_heading(doc, "5.2 SLA-Driven Auto-Escalation (L1 → L4)", 2)
add_bullets(doc, [
    "L1 — Project Officer. Email + Platform. SLA 4 business hours.",
    "L2 — Section Manager + Director. Email + SMS + Platform. SLA 1 business hour.",
    "L3 — DG + Deputies. All channels. SLA immediate.",
    "L4 — DG + Regulator. All channels + Regulator notice. Direct.",
])
add_para(doc, "AlertEscalationService Hangfire job runs every 5 minutes; AlertSla helper computes deadlines in UAE business hours; acknowledged alerts are not touched; L4 never escalates further. Five xUnit tests in AlertEscalationServiceTests cover the regression cases.")

add_heading(doc, "5.3 Developer Performance & Rating Engine", 2)
add_para(doc, "The Developer Performance & Rating engine computes a composite score across eight criteria. All weights are configurable from Admin/DeveloperRating.razor with full audit logging.", space_after=4)

add_figure(doc, os.path.join(SCREENS, '06_scorecards.png'),
    "Figure 5.2 — Public Developer Scorecards. 8-criterion composite, public-facing simplified view, weight changes audited end-to-end, Distinguished/Excellent/Good tier banding.")

add_table(doc, ["Scoring Criterion", "Source"], [
    ["On-Time Project Delivery Rate", "DeveloperScoringService + Project entity"],
    ["Unit Sales Completion Rate",    "DeveloperScoringService + Transaction"],
    ["Escrow Account Health Score",   "DeveloperScoringService + EscrowAccount"],
    ["Regulatory Compliance Record",  "DeveloperScoringService + RegulatoryViolation"],
    ["Customer Complaints",           "ScoringWeight + Complaint entity"],
    ["Construction Quality",          "ScoringWeight + inspection rows"],
    ["Financial Stability",           "ScoringWeight + audited finances"],
    ["Innovation & Sustainability",   "ScoringWeight + ESG signals"],
], col_widths_in=[3.0, 4.0])

add_heading(doc, "5.4 Escrow Monitoring (§8.4)", 2)
add_bullets(doc, [
    "ESC-001 — Real-time dashboard. Admin/EscrowMonitoring.razor + Admin/EscrowDetail.razor show current vs required balance, adequacy ratio, Adequate / Warning / Critical badges.",
    "ESC-002 — Immutable escrow audit log. EscrowTransaction is append-only at the entity level. PDF export via InvestorScorecardPdfService pattern. Unauthorised discrepancy raises a Level 3 alert.",
    "ESC-003 — Monthly Escrow Health Report. EscrowHealthReportService is a Hangfire job; generates a per-project PDF month-end; emails the assigned officer within 24 h.",
])

add_code_refs(doc, [
    "src/IRETP.Infrastructure/Services/RiskEngineService.cs",
    "src/IRETP.Infrastructure/Services/AlertEscalationService.cs",
    "src/IRETP.Infrastructure/Services/AuditLogService.cs",
    "src/IRETP.Infrastructure/Services/DeveloperScoringService.cs",
    "src/IRETP.Infrastructure/Services/EscrowHealthReportService.cs",
    "src/IRETP.Web/Components/Pages/Admin/EwrsDashboard.razor",
    "src/IRETP.Web/Components/Pages/Admin/EwrsAlerts.razor",
    "src/IRETP.Web/Components/Pages/Admin/EscrowMonitoring.razor",
    "src/IRETP.Web/Components/Pages/Admin/DeveloperRating.razor",
    "src/IRETP.Web/Components/Pages/Public/DeveloperScorecards.razor",
])

page_break(doc)

# =====================================================================
# SECTION 6 — DELIVERY
# =====================================================================
add_heading(doc, "6. Phased Delivery Plan", 1)
add_heading(doc, "6.1 Timeline Overview — 90 Calendar Days", 2)
add_table(doc, ["Phase", "Deliverable", "Dev / Go-Live", "Key Activities"], [
    ["Phase 0", "Data Familiarisation & Architecture Alignment", "Days 1–14 (parallel with P1)",
     "Connect OneLake Gold; schema mapping with DLD data liaison; derived-metric pipelines; LLM evaluation report Day 14."],
    ["Phase 1", "External Public Portal (EN + AR)", "Dev 1–28 · UAT 28–35 · Go-Live Day 35",
     "Portal hardening + DLD data integration + i18n. Pre-Publication Analytics Assessment Days 28–32. Load-test sign-off required."],
    ["Phase 2", "Early Warning Risk System (EWRS)", "Dev 14–48 · UAT 48–55 · Go-Live Day 56",
     "EWRS engine integrated against DLD data; RBAC for 6 internal roles; Internal AI Agent; escrow monitoring; audit-log immutability sign-off."],
    ["Phase 3", "Developer Rating + VAPT", "Dev 42–68 · VAPT 56–68 · UAT 65–75 · Go-Live Day 78",
     "Developer Rating Engine integration; public scorecard data sign-off; VAPT external assessor (pre-booked Day 1 for Day 56 slot)."],
    ["Phase 4", "Extended Languages + Supplementary", "Dev 60–85 · UAT 80–88 · Go-Live Day 90",
     "Multilingual (ZH/RU/UR/FR/HI/DE); ESG enhancements; benchmarking; PDF generator; documentation; training; handover."],
], col_widths_in=[0.8, 1.8, 1.7, 2.7])

add_heading(doc, "6.2 Why 90 Days is Achievable", 2)
add_bullets(doc, [
    "Phase 1 is integration, not construction. The External Portal is in production-build state today. P1 work is connecting OneLake Gold, running Pre-Publication Analytics Assessment, completing UAT, and going live.",
    "Phase 2 EWRS is hardening, not building. The 10-indicator engine, L1–L4 auto-escalation, and central audit log are live (verified by 10 + 5 xUnit tests).",
    "Parallel team tracks. Dedicated sub-teams own each phase. P2 development starts Day 14 regardless of P1 status. P3 Day 42. P4 Day 60.",
    "Phase 4 is configuration-dominant. Extended language support is primarily an i18n resource-file exercise. Combined with supplementary modules, P4 is completable in three weeks.",
])

add_heading(doc, "6.3 Pre-Publication Analytics Assessment — Hard Go-Live Gate", 2)
add_bullets(doc, [
    "Raw data reconciliation — random sample of transactions from Gold compared against platform display. Required match: 99.5%+.",
    "Derived metric validation — every computed metric independently recalculated from raw Gold data. Zero tolerance for calculation errors.",
    "Edge case review — outliers identified, classified, handled by agreement with DLD before go-live.",
])
add_para(doc, "Results documented in a Pre-Publication Analytics Assessment Report. DLD provides written sign-off before the platform goes live. If the assessment fails any threshold, Phase 1 go-live moves — this gate does not bend.")

page_break(doc)

# =====================================================================
# SECTION 7 — SECURITY
# =====================================================================
add_heading(doc, "7. Security & DESC Compliance", 1)
add_heading(doc, "7.1 DESC ISR v3 Compliance Framework", 2)
add_para(doc, "The complete 35-control DESC ISR v3 mapping is in Annex D. The current posture is 27 Implemented, 3 Partial, 5 Planned (pre-VAPT). The summary roll-up by domain:", space_after=4)
add_table(doc, ["Control Domain", "Implemented", "Partial", "Planned", "Total"], [
    ["Access Control",       "6", "1", "0", "7"],
    ["Asset Management",     "2", "1", "0", "3"],
    ["Business Continuity",  "1", "0", "1", "2"],
    ["Compliance",           "3", "0", "1", "4"],
    ["Cryptography",         "2", "0", "0", "2"],
    ["Governance",           "4", "1", "0", "5"],
    ["HR Security",          "1", "0", "2", "3"],
    ["Incident Management",  "3", "0", "0", "3"],
    ["Operations",           "6", "0", "0", "6"],
    ["Physical",             "2", "0", "0", "2"],
    ["Total",                "30", "3", "4", "37"],
], col_widths_in=[2.4, 1.2, 1.0, 1.0, 1.0])

add_heading(doc, "7.2 Security Testing Plan", 2)
add_table(doc, ["Test Type", "Schedule & Approach", "Remediation Commitment"], [
    ["Static analysis (SAST)", "Every commit. CodeQL + .NET analysers (TreatWarningsAsErrors).", "Critical/High blocks the build."],
    ["Dependency scan",        "Every commit. dotnet list package --vulnerable in CI.",          "Critical/High blocks the build."],
    ["DAST baseline",          "Every push to main. OWASP ZAP baseline scan.",                  "High triggers immediate remediation."],
    ["External VAPT",          "Phase 3 (Days 56–68). DESC-approved assessor pre-engaged Day 1.","Clean report required before P3 UAT sign-off."],
    ["Re-VAPT in warranty",    "Annually plus on any major release.",                            "Clean report required before production deploy."],
], col_widths_in=[1.4, 3.0, 2.6])

add_heading(doc, "7.3 ISO 27001 Roadmap", 2)
add_para(doc, "[Vendor Legal Entity Name] commits to completing an ISO 27001 gap assessment against the IRETP-specific platform configuration by Phase 3 (Day 78) and presenting a documented certification roadmap to DLD. Target: full ISO 27001 certification for the IRETP environment within six months of final go-live.")

add_heading(doc, "7.4 Hosting Options — Three Costed Models", 2)
add_table(doc, ["Criterion", "Option A — Azure UAE", "Option B — On-Prem", "Option C — Hybrid"], [
    ["Data residency",       "100% UAE (Azure UAE North)",                       "100% UAE (DLD DC)",      "100% UAE (compute on-prem + lakehouse in Azure UAE)"],
    ["Infrastructure",       "App Service Plan + Azure SQL + Redis + Front Door","K8s + SQL Server + Redis + open-source obs", "Hybrid · compute on-prem · lakehouse in Azure UAE"],
    ["DESC-CSP",             "Pre-certified",                                    "Full ISR v3 on-prem assessment",   "Mixed — admin docs + on-prem assessment"],
    ["Time to live",         "1 week via Bicep IaC",                            "4–8 weeks for hardware",  "3–5 weeks"],
    ["5-year CAPEX",         "AED 335,500",                                      "AED 1,414,500",           "AED 793,500"],
    ["5-year OPEX (annual)", "AED 629,200",                                      "AED 686,200",             "AED 654,200"],
    ["Recommended",          "Yes (fastest TTL, lowest CAPEX)",                  "If sovereignty mandates",  "Balanced posture"],
], col_widths_in=[1.6, 1.9, 1.8, 1.8])
add_para(doc, "DLD selects the hosting model at contract signature; financial schedules adjust accordingly. All AI inference and personal data remain in UAE on any option.")

add_heading(doc, "7.5 Data Residency, IAM, Audit", 2)
add_bullets(doc, [
    "Data residency. Bicep templates pin every Azure resource to UAE North; DR replicates to UAE Central. /healthz/sla fails Unhealthy on residency drift.",
    "IAM. ASP.NET Identity backed by SQL. MFA mandatory for internal users via TOTP. UAE Pass OIDC federation. Refresh-token rotation with reuse detection. 30-min internal idle timeout.",
    "Audit. Every threshold change, weight change, and AI model swap writes an append-only AuditLog row. The table is immutable at the DbContext layer.",
    "Backups. Daily automated encrypted backups, 30-day retention; PITR enabled on Azure SQL. Quarterly restoration drill reported to DLD. RTO < 4 h; RPO < 1 h.",
])
add_code_refs(doc, [
    "src/IRETP.Infrastructure/Identity/AuthorizationPolicies.cs · 6-role policy registry",
    "src/IRETP.Infrastructure/Services/AuditLogService.cs · append-only enforcement",
    "src/IRETP.Infrastructure/Data/IretpDbContext.cs · SaveChangesAsync throws on AuditLog UPDATE/DELETE",
    ".github/workflows/security.yml · vuln package + OWASP ZAP + Bicep lint",
])

page_break(doc)

# =====================================================================
# SECTION 8 — VISUALISATION
# =====================================================================
add_heading(doc, "8. Visualisation Strategy — Public and Internal", 1)
add_heading(doc, "8.1 Public Portal — Blazor Server + Chart.js + MapLibre", 2)
add_para(doc, "For the External Public Portal we use a Blazor Server frontend rendering Chart.js charts and MapLibre GL JS maps. The visualisation choice rationale is in Section 3.3 (full comparison vs Power BI Embedded). Public-portal pages share a consistent loading-state / error-state experience via ModuleBanner, Skeleton, NetworkError, EmptyState components.")

add_figure(doc, os.path.join(SCREENS, '03_transactions.png'),
    "Figure 8.1 — Transactions Explorer. FR-007 URL-persisted filters, server-side pagination, sortable headers (keyboard-accessible with aria-sort), CSV / PDF export via Features/Export.")

add_heading(doc, "8.2 Internal Platform — Power BI for Named-User Dashboards", 2)
add_para(doc, "Power BI is retained for Internal Platform components where named-user licensing is cost-appropriate: EWRS dashboards, Developer Rating dashboards, internal management reporting. Power BI tiles are embedded via the standard report-server SDK; the Internal AI Agent can answer questions over the same semantic models.")

page_break(doc)

# =====================================================================
# SECTION 9 — DATA UNDERSTANDING
# =====================================================================
add_heading(doc, "9. Data Understanding & Accuracy Methodology", 1)

add_heading(doc, "9.1 Connecting to DLD's Data Ecosystem", 2)
add_bullets(doc, [
    "Reading DLD's existing OneLake Gold-layer schema with DLD's designated data liaison.",
    "Configuring IRETP read-only access via the Fabric SQL endpoint — no writes, no modifications.",
    "Agreeing the calculation logic for all derived metrics in writing with DLD before any pipeline is wired.",
    "Documenting the field mapping between Gold-layer fields and IRETP display fields, formally approved by DLD.",
])

add_heading(doc, "9.2 Pre-Publication Analytics Assessment Methodology", 2)
add_para(doc, "Covered in Section 6.3. Three-layer validation (99.5% reconciliation, zero-tolerance derived-metric, edge-case agreement) signed off in writing by DLD before any data is publicly visible. Automated via Great Expectations / dbt-tests harness.")

add_heading(doc, "9.3 Ongoing Data Accuracy Monitoring", 2)
add_bullets(doc, [
    "Freshness probe — /healthz/sla reports KPI snapshot age vs 15-min TTL; drift beyond 30 min raises an alert.",
    "Reconciliation sampler — nightly Hangfire job samples 100 random transactions from Gold; discrepancies > 0.5% raise a P2 incident.",
    "Notification SLA probe — last hour of dispatched alerts, P95 latency per channel vs §6.2 budgets.",
    "AI residency probe — every AI snapshot logs the inference region; trailing-hour residency % computed.",
])

page_break(doc)

# =====================================================================
# SECTION 10 — DATA INTEGRATION
# =====================================================================
add_heading(doc, "10. Data Integration Reference Architecture", 1)
add_heading(doc, "10.1 Gold Layer as Single Source of Analytical Truth", 2)
add_para(doc, "The IRETP service tier connects directly to OneLake Gold via the Fabric SQL endpoint. No intermediate data warehouse, no proprietary integration layer. Every derived metric the platform displays is computed from a Gold-layer query — there is no parallel data store to drift out of sync.")

add_heading(doc, "10.2 Validation Pipeline Architecture", 2)
add_bullets(doc, [
    "On ingestion (Phase 0). Pre-Publication Analytics Assessment validates Gold-layer consumption end-to-end before public go-live.",
    "On display (runtime). KPI requests flow through KpiSnapshotCache; cache misses trigger a fresh Gold query whose result is reconciled before serving.",
    "On audit (continuous). Reconciliation sampler nightly job samples Gold rows and compares to displayed data.",
])

add_heading(doc, "10.3 Bilingual Data Handling", 2)
add_para(doc, "Arabic field handling is first-class throughout the data layer. Zone names, developer names, owner names (Beneficial Ownership), property categories, transaction types — all stored and rendered in their native script.")

add_figure(doc, os.path.join(SCREENS, '07_beneficial_ownership.png'),
    "Figure 10.1 — Beneficial Ownership page. Arabic owner names rendered under their English equivalents with proper RTL styling and lang=ar. Where data is unavailable, an explicit disclosure statement renders.")

page_break(doc)

# =====================================================================
# SECTION 11 — TEAM + RACI
# =====================================================================
add_heading(doc, "11. Project Team & Responsibility Matrix", 1)

add_heading(doc, "11.1 Core Delivery Team", 2)
add_table(doc, ["Role", "Location", "IRETP Remit"], [
    ["Project Manager",            "Dubai, UAE",            "Single point of accountability. Daily standup, weekly DLD status, gate-sign-off custody."],
    ["Delivery / Technical Lead",  "Dubai, UAE",            "Architecture authority. Clean-Architecture conformance, security posture, code-review gate."],
    ["Senior .NET Developer × 2",  "Dubai, UAE (onsite)",   "P1–P3 implementation lead. WebAPI, AdminAPI, Hangfire jobs."],
    ["Mid .NET Developer × 2",     "Dubai, UAE / UAE TZ",   "Feature implementation. Blazor surfaces, integration tests."],
    ["QA Engineer",                "Dubai, UAE",            "Test plan authorship, UAT facilitation, automated regression."],
    ["UI/UX Designer",             "Dubai, UAE",            "Bilingual UX, RTL fidelity, design-system stewardship."],
    ["DevOps / DESC Liaison",      "Dubai, UAE",            "Bicep IaC, Azure landing zone, DESC-CSP filings, VAPT coordination."],
], col_widths_in=[2.2, 1.9, 3.0])

add_heading(doc, "11.2 Responsibility Matrix (DLD · Vendor · Partner)", 2)
add_para(doc, "Legend: L = Lead · S = Support · A = Approver · I = Informed", italic=True, size=10, color=COL_MUTED)
add_table(doc, ["Activity", "DLD", "Vendor", "Partner"], [
    ["Programme governance & steering",                       "A", "L", "S"],
    ["Solution architecture & design",                         "I", "L", "S"],
    ["Application development (.NET 9 · Blazor · CQRS)",      "I", "L", "—"],
    ["Infrastructure provisioning & hosting",                  "A", "S", "L"],
    ["Data discovery, migration & lakehouse build",            "S", "L", "S"],
    ["AI / RAG pipeline & guardrails",                         "A", "L", "I"],
    ["Security hardening · DESC ISR v3 · VAPT",                "A", "L", "S"],
    ["QA · test execution · UAT facilitation",                 "S", "L", "S"],
    ["Acceptance sign-off",                                    "A", "I", "I"],
    ["Post-go-live operations, support & SLA",                "I", "L", "S"],
], col_widths_in=[3.4, 1.2, 1.2, 1.2])
add_para(doc, "VAPT is executed by a DESC-approved third-party firm engaged by the Vendor; results reviewed jointly with DLD before Phase 3 sign-off.", italic=True, size=10, color=COL_MUTED)

add_heading(doc, "11.3 UAE Market Presence", 2)
add_para(doc, "[Vendor Legal Entity Name] operates a UAE-registered legal entity. The Project Manager, Delivery / Technical Lead, onsite Senior Developers, QA Engineer, UI/UX Designer, and DevOps/DESC Liaison are all UAE-resident. Remote developer roles are filled from the nearest UAE-time-aligned office. DLD can request in-person meetings at any point during delivery.")

page_break(doc)

# =====================================================================
# SECTION 12 — COST
# =====================================================================
add_heading(doc, "12. Cost Structure", 1)
add_heading(doc, "12.1 Cost Summary (Option A — Azure UAE North recommended)", 2)
add_table(doc, ["Cost category", "Amount (AED)", "Type"], [
    ["Infrastructure CAPEX", "335,500", "One-time"],
    ["Software Licences (Year 1)", "235,200", "One-time"],
    ["Application Development (M1–M4)", "2,007,000", "One-time"],
    ["Total one-time CAPEX", "2,577,700", "—"],
    ["Annual OPEX (Year 1, includes 12-month warranty)", "629,200", "Recurring"],
], col_widths_in=[3.4, 1.8, 1.8])

add_heading(doc, "12.2 Infrastructure CAPEX — Three Hosting Options", 2)
add_table(doc, ["Infrastructure Item", "Option A — Azure", "Option B — On-Prem", "Option C — Hybrid"], [
    ["Compute · App Service / VMs",                "48,000",  "320,000",   "48,000"],
    ["SQL Server (managed / licensed)",            "36,000",  "180,000",   "36,000"],
    ["Redis cache cluster",                        "18,000",  "85,000",    "18,000"],
    ["Azure AI Search / vector index",             "24,000",  "65,000",    "24,000"],
    ["Microsoft Fabric / OneLake lakehouse",       "42,000",  "200,000",   "42,000"],
    ["Blob Storage / on-prem NAS",                 "12,000",  "55,000",    "12,000"],
    ["WAF + Front Door / on-prem F5",              "30,000",  "140,000",   "30,000"],
    ["CDN (public static assets)",                  "8,000",  "25,000",     "8,000"],
    ["Private endpoints + S2S network",            "14,000",  "60,000",    "14,000"],
    ["DR · cross-region replication (≤4 h RPO)",   "28,000",  "180,000",   "28,000"],
    ["SSL certificates (wildcard, 3-year)",         "4,500",  "4,500",      "4,500"],
    ["Monitoring / telemetry",                     "16,000",  "45,000",    "16,000"],
    ["Initial security hardening + pen-test",      "55,000",  "55,000",    "55,000"],
    ["TOTAL INFRA CAPEX",                         "335,500", "1,414,500", "335,500"],
], col_widths_in=[2.6, 1.6, 1.6, 1.6])

add_heading(doc, "12.3 Software Licences (Year 1)", 2)
add_table(doc, ["Licence / Subscription", "Vendor", "Annual (AED)"], [
    [".NET 9 runtime + ASP.NET Core",            "Microsoft (open-source)", "0"],
    ["Visual Studio Enterprise (5 seats)",       "Microsoft",               "64,000"],
    ["Hangfire Pro (2 servers)",                 "HangfireIO",              "14,400"],
    ["Azure OpenAI Service (tokens, Y1 est.)",   "Microsoft",               "90,000"],
    ["AI fallback model (UAE-resident provider)","Selected provider",       "40,000"],
    ["SMTP relay (SendGrid / Postmark)",         "Twilio / Postmark",       "2,400"],
    ["SMS gateway (UAE operator)",               "UAE operator",            "18,000"],
    ["PDF generation (QuestPDF)",                "QuestPDF",                "2,800"],
    ["Seq log server (cloud)",                   "Datalust",                "3,600"],
    ["Subtotal (Year 1)",                        "—",                       "235,200"],
], col_widths_in=[3.0, 2.2, 1.8])

add_heading(doc, "12.4 Application Development (M1–M4, 1,095 person-days)", 2)
add_table(doc, ["Deliverable", "Month", "PD", "Cost (AED)"], [
    ["External Public Portal (CMS · Dashboard · Transactions · Map · Indices)", "M1 · P1", "120", "216,000"],
    ["Bilingual EN/AR · RTL · i18n pipeline",                  "M1 · P1", "40",  "72,000"],
    ["Slice & Dice Analytics engine",                          "M1 · P1", "60",  "108,000"],
    ["Real Estate AI Agent (Orchestrator · RAG · guardrails)", "M1 · P1", "100", "180,000"],
    ["Investor Notifications (email / SMS / in-platform)",     "M1 · P1", "60",  "108,000"],
    ["Open Data API · Developer Portal · API-key management",  "M1 · P1", "50",  "90,000"],
    ["Perf hardening · Redis cache · EF tuning",               "M1 · P1", "30",  "54,000"],
    ["Phase 1 QA · UAT support",                               "M1 · P1", "30",  "54,000"],
    ["Phase 1 subtotal",                                       "M1",      "490", "882,000"],
    ["Watchlist v2 · ESG · GRETI sub-index tracker",           "M2 · P2", "80",  "144,000"],
    ["Analytics enhancements · saved views · embed",           "M2 · P2", "40",  "72,000"],
    ["Performance hardening · OpenTelemetry · load test",      "M2 · P2", "30",  "54,000"],
    ["Phase 2 QA · UAT",                                       "M2 · P2", "20",  "36,000"],
    ["Phase 2 subtotal",                                       "M2",      "170", "306,000"],
    ["EWRS (engine · L1–L4 · dashboard · playbooks)",          "M3 · P3", "100", "180,000"],
    ["Developer Performance & Rating + Escrow Monitoring",     "M3 · P3", "80",  "144,000"],
    ["Beneficial Ownership · Mortgage analytics · Name Validation","M3 · P3","40","72,000"],
    ["DESC ISR v3 hardening + VAPT (external consultant)",     "M3 · P3", "30",  "90,000"],
    ["Phase 3 QA · UAT · security acceptance",                 "M3 · P3", "25",  "45,000"],
    ["Phase 3 subtotal",                                       "M3",      "275", "531,000"],
    ["Multilingual extension (ZH · RU · UR · FR · HI · DE)",   "M4 · P4", "90",  "162,000"],
    ["Advanced analytics · AI fine-tuning pipeline",           "M4 · P4", "50",  "90,000"],
    ["Phase 4 QA · UAT · language acceptance",                 "M4 · P4", "20",  "36,000"],
    ["Phase 4 subtotal",                                       "M4",      "160", "288,000"],
    ["TOTAL development CAPEX",                                "M1–M4",   "1,095", "2,007,000"],
], col_widths_in=[3.4, 0.9, 0.9, 1.4])
add_para(doc, "Blended day rates: Architect AED 2,200 · Senior Developer 1,800 · Mid Developer 1,500 · QA / DevOps 1,400.", italic=True, color=COL_MUTED, size=9.5)

add_heading(doc, "12.5 Annual OPEX — Three Hosting Options", 2)
add_table(doc, ["OPEX Item", "Option A", "Option B", "Option C"], [
    ["Compute / VM / App Service",            "48,000",  "85,000",  "75,000"],
    ["SQL Server managed annual",             "36,000",  "24,000",  "28,000"],
    ["Redis / caching annual",                "18,000",  "12,000",  "18,000"],
    ["AI model API consumption (annual est.)","130,000", "130,000", "130,000"],
    ["OneLake / Fabric storage & compute",    "42,000",  "68,000",  "42,000"],
    ["Monitoring & SIEM",                     "16,000",  "28,000",  "22,000"],
    ["SMTP / SMS channels",                   "20,400",  "20,400",  "20,400"],
    ["Software licences (annual renewal)",    "84,800",  "84,800",  "84,800"],
    ["Warranty & support (12-mo post go-live)","180,000","180,000", "180,000"],
    ["Security patching · monthly scans",     "36,000",  "36,000",  "36,000"],
    ["Data quality · reconciliation tooling", "18,000",  "18,000",  "18,000"],
    ["TOTAL Annual OPEX",                     "629,200", "686,200", "654,200"],
], col_widths_in=[3.0, 1.4, 1.4, 1.4])

add_heading(doc, "12.6 Payment Milestones", 2)
add_table(doc, ["Milestone", "Trigger", "%", "Amount (AED)"], [
    ["M0", "Contract signature",                                                              "10%", "200,700"],
    ["M1", "Phase 1 mobilisation complete (infrastructure live, team onboard)",               "10%", "200,700"],
    ["M2", "Phase 1 UAT sign-off (External Portal · AI Agent · Open Data)",                   "25%", "501,750"],
    ["M3", "Phase 2 UAT sign-off (ESG · GRETI · Analytics v2)",                               "15%", "301,050"],
    ["M4", "Phase 3 UAT sign-off (EWRS · Developer Rating · Escrow · VAPT clean)",            "25%", "501,750"],
    ["M5", "Phase 4 UAT sign-off (Multilingual · Advanced Analytics)",                        "15%", "301,050"],
    ["Total", "—",                                                                              "100%", "2,007,000"],
], col_widths_in=[0.7, 3.8, 0.6, 2.0])
add_para(doc, "100% of development CAPEX is conditional on DLD-signed acceptance certificates. No payment is released without an executed acceptance package.", italic=True, color=COL_MUTED, size=9.5)

add_heading(doc, "12.7 Indicative 5-Year Total Cost of Ownership (Option A)", 2)
add_table(doc, ["Period", "Component", "Amount (AED)"], [
    ["Year 0", "One-time CAPEX",     "2,577,700"],
    ["Year 1", "Annual OPEX",         "629,200"],
    ["Year 2", "Annual OPEX",         "629,200"],
    ["Year 3", "Annual OPEX",         "629,200"],
    ["Year 4", "Annual OPEX",         "629,200"],
    ["Year 5", "Annual OPEX",         "629,200"],
    ["5-year Total", "Indicative TCO","5,723,700"],
], col_widths_in=[1.0, 4.2, 1.8])

page_break(doc)

# =====================================================================
# SECTION 13 — WARRANTY
# =====================================================================
add_heading(doc, "13. Warranty & Technical Support — 12 Months Post Go-Live", 1)

add_heading(doc, "13.1 Support Philosophy", 2)
add_para(doc, "The 12-month warranty and technical support obligation is not a separate procurement. It is a contractual extension of our delivery responsibility. We do not consider a project delivered until the platform has operated stably, accurately, and securely through its first full year of live operation. Five categories: defect resolution, data pipeline correction, security patches, performance fixes, DLD data-structure change management.")

add_heading(doc, "13.2 Support Incident SLA Commitments", 2)
add_table(doc, ["Severity", "Definition", "Response", "Resolution", "Coverage"], [
    ["Sev 1", "Platform unavailable; investor-facing data wrong; security breach.", "1 h", "4 h", "24×7 incl. weekends/UAE holidays"],
    ["Sev 2", "Major feature degraded; EWRS escalation impaired; AI down on one tier.","2 h","1 day","Business hours UAE"],
    ["Sev 3", "Minor functional defect; cosmetic; localisation.",                     "1 d", "1 week","Business hours UAE"],
    ["Sev 4", "Enhancement request; doc correction; data-model expansion.",          "1 d", "Next planning cycle","Business hours UAE"],
], col_widths_in=[0.6, 2.6, 0.8, 0.9, 2.1])
add_para(doc, "Business hours: Mon–Fri 08:00–16:00 UAE Standard Time, excluding UAE public holidays. Named Support Lead assigned as DLD's primary contact. Emergency out-of-hours escalation number for Sev 1.")

add_heading(doc, "13.3 Warranty & Support Pricing", 2)
add_table(doc, ["Cost Item", "Annual (AED)"], [
    ["S.1 Dedicated Support Team (Lead 25% + 2 Devs 50% each + QA 25%)", "Included in OPEX"],
    ["S.2 120 PH minor changes pool. Additional: AED 950 / PH.",         "Included"],
    ["S.3 Security vulnerability monitoring + patch management.",        "Included"],
    ["S.4 Infrastructure operations. Covered by Azure managed services.","0"],
    ["S.5 Help desk + incident management tooling.",                     "18,260"],
    ["TOTAL warranty & support (within Annual OPEX)",                    "180,000"],
], col_widths_in=[5.2, 1.8])

page_break(doc)

# =====================================================================
# SECTION 14 — AI KNOWLEDGE UPDATE
# =====================================================================
add_heading(doc, "14. AI Agent Knowledge Update Approach", 1)
add_para(doc, "The AI Agent's knowledge layer is grounded in the OneLake Gold layer. Knowledge updates are managed by the AdminAPI surface without vendor involvement after handover.")
add_table(doc, ["Knowledge Layer", "Update Mechanism"], [
    ["Transactions / Projects / Developers", "Hangfire cron, hourly. Incremental delta from Gold."],
    ["RERA regulations / DLD procedures",    "CMS publish webhook → immediate reindex of the affected document."],
    ["Zone boundaries / GeoJSON",            "Manual trigger on Dubai Municipality updates."],
    ["Beneficial ownership",                  "Daily cron + on-demand re-index from AdminAPI."],
], col_widths_in=[2.4, 4.6])

add_heading(doc, "14.1 Index Management in AdminAPI", 2)
add_bullets(doc, [
    "Index freshness dashboard — last update timestamp, next scheduled update, record count.",
    "On-demand re-index — any layer can be re-indexed manually from the admin surface.",
    "Rollback — every update creates a snapshot; SystemAdministrator can roll back in 2 clicks within 5 min.",
    "Accuracy benchmark — the accuracy harness runs automatically after each update; deviation raises an alert before updated content reaches end users.",
])

add_heading(doc, "14.2 LLM Model Lifecycle (BYO LLM)", 2)
add_para(doc, "Orchestrator is LLM-agnostic. Switching the active model on any tier is a configuration change — no redeployment, no downtime. DLD can re-evaluate models in Year 2 against newer releases by running the accuracy harness again.")

page_break(doc)

# =====================================================================
# SECTION 15 — RISK REGISTER
# =====================================================================
add_heading(doc, "15. Risk Register", 1)
add_table(doc, ["Risk", "Probability", "Impact", "Mitigation"], [
    ["Gold-layer schema drift during Phase 0", "Medium", "Schedule",
     "Schema versioning agreed Day 2; integration tests on every Gold change; field-mapping doc is contractual."],
    ["Pre-Publication Analytics Assessment fails", "Low", "Schedule",
     "Dry-run assessment Days 21–24 surfaces issues 10 days early; failing items resolved before formal gate."],
    ["VAPT findings exceed Day 68", "Medium", "Schedule",
     "VAPT scheduled Days 56–68; 7-day remediation buffer; Critical/High blocks UAT sign-off."],
    ["AI residency drift", "Low", "Compliance",
     "SLA probe fails Unhealthy on any non-UAE inference; automated rollback to UAE-resident fallback."],
    ["Investor surge above 3× projected volume", "Medium", "Performance",
     "Auto-scale on CPU/memory; Redis absorbs read pressure; load test rerun on major releases."],
    ["DESC certification slips into warranty", "Medium", "Compliance",
     "DESC engagement begins Day 1; gap assessment Day 78; documented roadmap delivered to DLD regardless."],
    ["Key personnel attrition", "Low", "Schedule",
     "Continuous KT in source; Lead has named backup; no single-person dependency per phase."],
    ["DLD policy change mid-flight", "Medium", "Scope",
     "Configurable thresholds + i18n strings mean most policy changes are configuration. Material changes flow through change-control."],
    ["AI model provider UAE residency unavailable", "Medium", "Compliance",
     "Provider contractually bound to UAE inference; UAE-hosted OSS fallback (Llama/Mistral) ready."],
], col_widths_in=[2.0, 0.9, 0.9, 3.2])

page_break(doc)

# =====================================================================
# SECTION 16 — KT & EXIT
# =====================================================================
add_heading(doc, "16. Knowledge Transfer & Exit Plan", 1)

add_heading(doc, "16.1 Deliverables at Project Completion (Day 90)", 2)
add_bullets(doc, [
    "Source repositories — all .NET application code, Bicep IaC, CI/CD pipeline configurations, Hangfire job definitions, AI orchestration configuration, knowledge-source configuration. Delivered as GitHub repositories transferred to DLD's GitHub organisation.",
    "Complete technical documentation — ARCHITECTURE.md, API_REFERENCE.md, RUNBOOK.md, DISASTER_RECOVERY.md, DATA_DICTIONARY.md (all five RFP §13 #53 pillars).",
    "Infrastructure control — Azure subscription, Key Vault, DNS transferred to DLD's tenant on go-live. No vendor-controlled infrastructure persists.",
    "Data export — full platform export in JSON, CSV, Parquet within 30 calendar days of contract expiry upon request.",
    "Data destruction certification — within 60 days of DLD's request.",
])

add_heading(doc, "16.2 Intellectual Property", 2)
add_para(doc, "All deliverables, results, works, documents, data, reports, designs, and materials prepared, developed, or submitted under this contract shall be the exclusive intellectual property of the Dubai Land Department, consistent with Section 4.4 of the RFP. There is no proprietary platform layer between DLD and the application source; everything in the deliverable stack is either standard Microsoft technology, open-source, or written specifically for DLD.")

page_break(doc)

# =====================================================================
# SECTION 17 — GRETI BEYOND BRIEF
# =====================================================================
add_heading(doc, "17. Beyond the Brief — GRETI Score Maximisation", 1)
add_para(doc, "RFP §20 identifies additional features that maximise GRETI score improvement beyond core scope. Three are not roadmap items — they are already shipping in the reference build today.")

add_heading(doc, "17.1 Beneficial Ownership Transparency (Already Shipping)", 2)
add_para(doc, "Per JLL's 2024 GRETI report, beneficial ownership transparency is a key gap for the UAE. Public/BeneficialOwnership.razor renders ownership data per registered developer: beneficial owner name (English + Arabic, RTL-styled), ownership percentage, UAE commercial registration number. See Figure 10.1.")

add_heading(doc, "17.2 Mortgage & Debt Market Transparency (Already Shipping)", 2)

add_figure(doc, os.path.join(SCREENS, '08_mortgage.png'),
    "Figure 17.1 — Mortgage page with 5 KPIs (Total Value, Mortgage Count, M:Tx Ratio, Avg Mortgage Size, Avg LTV) plus 4 supporting charts. All data sourced exclusively from registered DLD transaction records.")

add_heading(doc, "17.3 ESG Indicators (Already Shipping)", 2)
add_para(doc, "Public/Esg.razor surfaces the LEED ↔ Estidama Pearl ↔ BREEAM ↔ investor-signal rubric. GIS map exposes a 4th heatmap mode — ESG Coverage — that visualises green-certified coverage by zone with banded fill colour and per-zone percentage markers.")

add_heading(doc, "17.4 Implementation Roadmap", 2)
add_para(doc, "We recommend DLD discuss Beneficial Ownership, Mortgage Transparency, and ESG modules with the JLL GRETI assessment team before finalising the project scope. All three are cited in the GRETI 2024 report; their inclusion creates measurable indicators of progress that the next GRETI assessment cycle can capture within this project's delivery timeline.")

page_break(doc)

# =====================================================================
# SECTION 18 — VENDOR COMMITMENTS
# =====================================================================
add_heading(doc, "18. Vendor Commitments — 15 Engineering Improvements at No Extra Cost", 1)
add_para(doc, "Beyond the RFP-mandated scope, the vendor commits to deliver the following 15 engineering improvements during Phase 1 hardening at no additional charge. These items lift the codebase to enterprise-grade operability and directly strengthen DLD's NFR posture. Two items below (VC-6 SignalR hub, VC-14 Dockerfiles + CI) are already shipped in the reference build; the remaining 13 are committed within the Phase 1 hardening scope.")
add_table(doc, ["ID", "Commitment", "Implementation"], [
    ["VC-1",  "MediatR ValidationBehavior",     "Auto-runs every FluentValidation validator already registered in IRETP.Application."],
    ["VC-2",  "Global IExceptionHandler",        "Maps domain/validation/auth exceptions to RFC 7807 ProblemDetails."],
    ["VC-3",  "Distributed cache via Redis",     "ICacheableQuery marker + CachingBehavior backed by Redis. Serves KPI/GRETI/ESG queries to meet P95 ≤ 2 s headroom."],
    ["VC-4",  "Domain event dispatcher",         "IDomainEventDispatcher invoked from UnitOfWork.SaveChangesAsync. Decouples notifications/audit/SignalR push."],
    ["VC-5",  "Notification outbox pattern",     "NotificationOutbox written in same DB transaction; dispatcher drains to SMTP/SMS. Guarantees §6.2 SLAs on transient failure."],
    ["VC-6",  "SignalR notification hub",        "Already shipped — replaces NoOpNotificationBroadcaster. Bell icon updates instantaneously without polling."],
    ["VC-7",  "AI streaming over SSE",           "AiAgent.razor streams tokens as IAsyncEnumerable<string>. Sub-15-second perceived response."],
    ["VC-8",  "Polly resilience",                "AddStandardResilienceHandler() on AIOrchestrator, currency, SMTP, SMS HttpClients. Retry + circuit breaker — implements §5.3 fallback declaratively."],
    ["VC-9",  "Hangfire idempotency",            "Job fingerprints prevent double-execution on retries. Protects monthly Escrow Health Reports and KPI snapshots."],
    ["VC-10", "OpenTelemetry",                   "Traces / metrics / logs exported via OTLP, tagged with user, role, AI-model dimensions. Powers SLA dashboard."],
    ["VC-11", "/healthz + /readyz",              "Liveness + readiness probes (DB + Redis + Hangfire + AI provider)."],
    ["VC-12", "Tighter CORS + CSP + HSTS",       "Replaces AllowAnyOrigin with allow-list. Blazor CSP nonces for inline scripts."],
    ["VC-13", "Architecture tests",              "NetArchTest assertions: Domain has no Infrastructure refs; Application has no EF Core refs; Controllers don't depend on Repository<>."],
    ["VC-14", "Dockerfiles + CI/CD",             "Already shipped — multi-stage Dockerfiles + docker-compose + GitHub Actions / Azure Pipelines."],
    ["VC-15", "Resource-based authorization",    "Ownership IAuthorizationHandler — investors see own data only. Complements role-level policies."],
], col_widths_in=[0.7, 2.2, 4.1])

page_break(doc)

# =====================================================================
# SECTION 19 — QUALITY POSTURE
# =====================================================================
add_heading(doc, "19. Quality Posture & Test Strategy", 1)

add_heading(doc, "19.1 Hangfire Job Schedule", 2)
add_table(doc, ["Job", "Cadence", "Responsibility"], [
    ["RiskEngineService",              "every hour",  "Re-evaluate the 10 EWRS indicators."],
    ["DeveloperScoringService",        "nightly",     "Recompute developer composite scores."],
    ["AlertDeliveryService",           "every minute","Dispatch queued notifications across channels."],
    ["AlertEscalationService",         "every 5 min", "L1 → L4 auto-escalation."],
    ["KpiSnapshotRefreshService",      "every 15 min","Refresh the KPI snapshot cache."],
    ["InvestorAlertEvaluator",         "every 10 min","Evaluate investor alert conditions."],
    ["EscrowHealthReportService",      "month-end",   "Per-project PDF Escrow Health Report."],
    ["BenchmarkRefreshService",        "hourly",      "Refresh market benchmark data."],
    ["CurrencyRatesRefreshService",    "hourly",      "Refresh UAECB FX rates."],
], col_widths_in=[2.8, 1.2, 3.0])

add_heading(doc, "19.2 Test Strategy", 2)
add_table(doc, ["Layer", "Tool", "Scope"], [
    ["Unit",            "xUnit",                       "Per CQRS handler and service."],
    ["Integration",     "xUnit + EF Core",             "Per controller against real EF Core."],
    ["End-to-end UI",   "Playwright",                  "Critical investor and DLD-staff journeys."],
    ["Load / Performance","k6 / NBomber",              "P95 sign-off per release."],
    ["Security / VAPT", "OWASP ZAP + DESC-authorised VAPT", "Baseline scan in CI + external Phase 3 VAPT."],
    ["Accessibility",   "axe-core",                    "Automated scan per release; zero Critical violations."],
    ["Data accuracy",   "Great Expectations / dbt-tests", "Pre-Publication Analytics Assessment + ongoing."],
    ["Language & RTL",  "Native-speaker review",       "Per locale, every release."],
], col_widths_in=[1.5, 2.2, 3.3])

add_heading(doc, "19.3 Observability", 2)
add_bullets(doc, [
    "Structured logging via Serilog sinks (Seq / Azure Monitor).",
    "OpenTelemetry traces + metrics + logs to OTLP, tagged with user, role, AI-model.",
    "Health endpoints: /healthz/live, /healthz/ready, /healthz/sla (aggregated SLO health — AI P95, KPI freshness, notification P95).",
    "Audit log (Admin/AuditLogs.razor) queryable per user, per entity, per action with timestamp.",
])

page_break(doc)

# =====================================================================
# ANNEX A — ARCHITECTURE DIAGRAM
# =====================================================================
add_heading(doc, "Annex A — System Architecture Diagram", 1)
add_para(doc, "Four-tier IRETP architecture: User Access (Tier 1), Application & AI (Tier 3), Data & Integration (Tier 2, anchored on DLD's Microsoft Fabric ecosystem), and the Edge perimeter. Option A (recommended): Azure UAE North. Option B: on-premises at DLD's data centre. Option C: hybrid. All three run the identical .NET application stack.", space_after=8)

arch = (
"┌────────────────────────────────────────────────────────────────┐\n"
"│              Tier 1 — Edge & Access                            │\n"
"│  Azure Front Door (WAF + CDN + TLS 1.2)  ·  Azure DDoS Std     │\n"
"│  UAE Pass OIDC  ·  ASP.NET Identity + JWT + TOTP MFA           │\n"
"└─────────────────────────────────┬──────────────────────────────┘\n"
"                                  │\n"
"┌─────────────────────────────────▼──────────────────────────────┐\n"
"│              Tier 3 — Application & AI                         │\n"
"│   IRETP.Web (Blazor)   IRETP.WebAPI   IRETP.AdminAPI           │\n"
"│   Chart.js · MapLibre  AIOrchestrator (4-tier)  SignalR hub    │\n"
"│   Hangfire (9 jobs) · KPI cache · accuracy harness             │\n"
"└─────────────────────────────────┬──────────────────────────────┘\n"
"                                  │\n"
"┌─────────────────────────────────▼──────────────────────────────┐\n"
"│             Tier 2 — Data & Integration                        │\n"
"│   Microsoft Fabric / OneLake Gold  (single source of truth)    │\n"
"│   Azure SQL (elastic pool · TDE · PITR)                        │\n"
"│   Key Vault (MSI · purge-prot) · Blob Storage (GRS)            │\n"
"└────────────────────────────────────────────────────────────────┘\n"
"                                                                  \n"
"  Cross-cutting: App Insights · Log Analytics · SLA probe         \n"
)
mono = doc.add_paragraph()
r = mono.add_run(arch); r.font.name = 'Consolas'; r.font.size = Pt(9)
mono.paragraph_format.space_after = Pt(6)

add_para(doc, "All platform outputs are sourced exclusively from the OneLake Gold layer. IRETP does not write to or duplicate any data in DLD's Fabric ecosystem.", italic=True, color=COL_MUTED)

page_break(doc)

# =====================================================================
# ANNEX B — GANTT
# =====================================================================
add_heading(doc, "Annex B — Phased Delivery Gantt (90 Calendar Days)", 1)
add_para(doc, "Each cell denotes activity in a 5-day window. ▰ active · ▯ ramp · ✓ go-live · ⚙ gate.", italic=True, color=COL_MUTED, size=10)
gantt_rows = [
    ["Phase 0 — Data familiarisation",
     "▰","▰","▯", "", "", "", "", "", "", "", "", "", "", "", "", "", "", ""],
    ["Phase 1 — External Portal (Dev)",
     "▰","▰","▰","▰","▰","▰","✓", "", "", "", "", "", "", "", "", "", "", ""],
    ["Phase 1 — Pre-Pub Assessment ⚙",
     "", "", "", "", "⚙","⚙","✓", "", "", "", "", "", "", "", "", "", "", ""],
    ["Phase 2 — EWRS (Dev)",
     "", "", "▰","▰","▰","▰","▰","▰","▰","▰","✓", "", "", "", "", "", "", ""],
    ["Phase 3 — Dev Rating + VAPT",
     "", "", "", "", "", "", "", "", "▰","▰","▰","▰","▰","▰","✓", "", "", ""],
    ["Phase 3 — VAPT external ⚙",
     "", "", "", "", "", "", "", "", "", "", "⚙","⚙","⚙","⚙","✓", "", "", ""],
    ["Phase 4 — Multilingual + ESG",
     "", "", "", "", "", "", "", "", "", "", "", "▰","▰","▰","▰","▰","▰","✓"],
    ["UAT windows",
     "", "", "", "", "▰","▰","▰", "", "", "▰","▰", "", "▰","▰", "", "▰","▰", ""],
    ["Final Go-Live",
     "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "✓"],
]
gantt_headers = ["Activity"] + [f"{1+i*5}–{(i+1)*5}" for i in range(18)]
add_table(doc, gantt_headers, gantt_rows, col_widths_in=[2.6] + [0.34]*18, font_size=8)

page_break(doc)

# =====================================================================
# ANNEX C — FULL COMPLIANCE MATRIX (with C/CD/E status)
# =====================================================================
add_heading(doc, "Annex C — Functional Requirements Compliance Matrix", 1)
add_para(doc, "Every numbered or named requirement in the RFP is reproduced in this matrix together with a compliance status, the implementing component in the IRETP codebase, the delivery phase, and any clarifying notes.")
add_para(doc, "Status legend:", bold=True, space_after=2)
add_table(doc, ["Status", "Meaning"], [
    ["C",  "Compliant today — implementing component is live in the codebase. Implementing Component column points to the live file or service."],
    ["CD", "Compliant by Delivery — capability is in the phase plan; required foundations exist and the work is scoped, estimated and committed."],
    ["E",  "Enhanced (vendor commitment) — capability goes beyond the literal requirement. Listed in the Vendor Commitments group (Section 18)."],
], col_widths_in=[0.7, 6.3])

add_heading(doc, "C.1 Compliance Summary Roll-Up", 3)
add_table(doc, ["RFP Section", "C", "CD", "E", "Total"], [
    ["§3 — External Public Portal",              "5",  "0",  "0", "5"],
    ["§4 — Interactive Analytics",                "3",  "0",  "0", "3"],
    ["§5 — Real Estate AI Agent",                 "9",  "2",  "1", "12"],
    ["§6 — Investor Notifications",               "9",  "0",  "1", "10"],
    ["§7 — Multilingual Requirements",            "7",  "2",  "0", "9"],
    ["§8 — EWRS & Escrow Monitoring",             "16", "0",  "0", "16"],
    ["§9 — Developer Performance & Rating",       "7",  "0",  "0", "7"],
    ["§10 — Non-Functional, Security & RBAC",     "9",  "1",  "1", "11"],
    ["§11 — Technology & Hosting",                 "2",  "4",  "0", "6"],
    ["§12 — Data Familiarisation",                 "0",  "3",  "0", "3"],
    ["§13 — Phased Delivery Plan",                 "1",  "3",  "0", "4"],
    ["§14 — Cost Structure",                       "0",  "4",  "0", "4"],
    ["§15 — Warranty & Support",                   "0",  "4",  "0", "4"],
    ["§16 — Vendor Qualifications",                "0",  "2",  "0", "2"],
    ["§17 — Acceptance Criteria",                  "0",  "3",  "0", "3"],
    ["§18 — Proposal Evaluation",                  "0",  "3",  "0", "3"],
    ["§19 — Data Governance & Privacy",            "2",  "2",  "0", "4"],
    ["§20 — Additional Recommended Features",      "0",  "0",  "8", "8"],
    ["Vendor Commitments (Value-Add)",             "0",  "0",  "15", "15"],
    ["TOTAL",                                       "70", "33", "26", "129"],
], col_widths_in=[3.0, 0.7, 0.7, 0.7, 0.7])

def matrix_section(title, rows):
    add_heading(doc, title, 3)
    add_table(doc,
        ["RFP Ref", "Requirement", "Status", "Implementing Component", "Phase"],
        rows, col_widths_in=[0.7, 2.4, 0.6, 2.6, 0.7])

matrix_section("C.2 §3 — External Public Portal", [
    ["3.1", "Headless CMS for non-technical content updates", "C", "CmsContent + CmsContentVersion · Admin/Cms.razor · Public/CmsPreview.razor", "P1"],
    ["3.2", "Homepage Market Overview Dashboard with KPI tiles, charts and CTAs", "C", "Public/Dashboard.razor + DashboardController + KpiSnapshotCache", "P1"],
    ["3.3", "Transactions page with filters, search, export", "C", "Public/Transactions.razor + TransactionsController + Features/Transactions", "P1"],
    ["3.4", "Interactive GIS Map — zones & projects", "C", "Public/Map.razor + MapController + Features/Map; Zone entity polygons; Project pins", "P1"],
    ["3.5", "Price Index & Rental Index pages", "C", "Public/PriceIndex.razor + Public/RentalIndex.razor + BenchmarkRefreshService", "P1"],
])

matrix_section("C.3 §4 — Interactive Analytics", [
    ["4", "Slice-and-dice analytics engine", "C", "Public/Analytics.razor + AnalyticsController + SavedAnalyticsView", "P1"],
    ["4", "Embed mode for partner sites", "C", "Public/AnalyticsEmbed.razor + EmbedLayout", "P1"],
    ["4", "Export to CSV / PDF", "C", "Features/Export Commands + InvestorScorecardPdfService", "P1"],
])

matrix_section("C.4 §5 — Real Estate AI Agent", [
    ["5.1",   "AI Agent scope (data, navigation, valuation, regulations)", "C", "AIAgentController + Public/AiAgent.razor + Features/AIAgent", "P1"],
    ["AI-001","Natural-language data queries", "C", "Features/AIAgent/Queries handlers", "P1"],
    ["AI-002","Service navigation + deep links to DLD pages", "C", "Features/AIAgent + AIOrchestrator response templates", "P1"],
    ["AI-003","Valuation procedure explanations linked to DLD service pages", "C", "AIOrchestrator content sources + RAG index", "P1"],
    ["AI-006","Session context memory with opt-in cross-session persistence", "C", "UserAiMemory + UserConsentService + Account.razor toggle", "P1"],
    ["AI-007","Multilingual EN/AR P1; ZH/RU/UR/FR/HI/DE P4", "CD", "AIOrchestrator locale-aware prompts + per-language model selection", "P1/P4"],
    ["5.3",   "Multi-model architecture with orchestration layer", "C", "AIOrchestrator + AIModelMetrics + Admin/AIModels.razor", "P1"],
    ["5.3",   "Mandatory model-agnostic RAG layer", "C", "RAG retriever component (Infrastructure)", "P1"],
    ["5.3",   "Model fallback & redundancy without user-visible disruption", "E", "AIOrchestrator warm-fallback + Polly resilience (vendor commitment)", "P1"],
    ["5.3",   "UAE data residency for all AI inference", "C", "Hosting model + provider selection (Azure UAE North or UAE-hosted OSS)", "P1"],
    ["5.3",   "Custom fine-tuning option", "CD", "Fine-tuning pipeline (planned in Phase 2)", "P2"],
    ["5.3",   "Model performance transparency in admin UI", "C", "Admin/AIModels.razor + AIModelMetrics", "P1"],
])

matrix_section("C.5 §6 — Investor Notifications", [
    ["6.1","Price Movement alert (user threshold 1–25%)", "C", "InvestorAlert + InvestorAlertEvaluator", "P1"],
    ["6.1","New Project Launch alert", "C", "InvestorAlertEvaluator + Project entity feed", "P1"],
    ["6.1","Watchlist Project Status Change alert", "C", "WatchlistItem + InvestorAlertEvaluator", "P1"],
    ["6.1","Rental Yield Threshold alert", "C", "InvestorAlertEvaluator + RentalIndex", "P1"],
    ["6.1","Periodic Market Digest (weekly/monthly)", "C", "Hangfire scheduled job + AlertDeliveryService", "P1"],
    ["6.1","Regulation/Policy Update alert", "C", "RegulatoryViolation + CmsContent triggers", "P1"],
    ["6.2","Email branded HTML + embedded chart + unsubscribe", "C", "SmtpEmailSender + NotificationTemplates + HmacUnsubscribeTokenService", "P1"],
    ["6.2","SMS ≤160 chars with shortened link", "C", "HttpSmsSender + NotificationTemplates", "P1"],
    ["6.2","In-platform notification centre with bell + unread badge", "C", "Notification entity + HeaderWidgets.razor + Public/Notifications.razor", "P1"],
    ["6.2","Delivery SLA: email ≤5 min, SMS ≤3 min, in-platform instant", "E", "Hangfire scheduling + outbox pattern (vendor commitment)", "P1"],
])

matrix_section("C.6 §7 — Multilingual Requirements", [
    ["7",   "Internal platform full EN/AR (UI, content, AI, alerts)", "C", "MainLayout + i18n resources + Admin pages", "P1"],
    ["7",   "External portal core EN/AR with full NLP + export labels", "C", "Public pages + AIOrchestrator locale prompts + Features/Export", "P1"],
    ["7",   "External portal extended: ZH, RU, UR, FR, HI, DE", "CD", "i18n resource expansion + AI multi-language", "P4"],
    ["7.1", "RTL layout mirroring for AR and UR", "C", "MainLayout.razor + per-page CSS", "P1"],
    ["7.1", "All text in Unicode/UTF-8; no hardcoded strings", "C", "Resource files + CMS locale fields", "P1"],
    ["7.1", "Locale-aware chart labels, tooltips, exports, PDFs", "C", "Chart components + Export commands", "P1"],
    ["7.1", "AR validation of zone/project/developer names against DLD records", "CD", "Data discovery sprint (§12.1)", "P1"],
    ["7.1", "Language preference per account + browser; preserved across nav", "C", "ApplicationUser locale + browser storage", "P1"],
    ["3/7", "Currency switcher in header", "C", "CurrencyController + CurrencyRate + CurrencyRatesRefreshService", "P1"],
])

matrix_section("C.7 §8 — EWRS & Escrow Monitoring", [
    ["8.1","Project Delivery Delay (Warning / Critical)", "C", "RiskEngineService + RiskThreshold rows", "P3"],
    ["8.1","Escrow Shortfall (Warning / Critical)", "C", "RiskEngineService + EscrowAccount", "P3"],
    ["8.1","Construction Activity Suspension (30/60 day)", "C", "RiskEngineService", "P3"],
    ["8.1","Sharp Transaction Volume Decline", "C", "RiskEngineService + TimeSeriesAnalyzer", "P3"],
    ["8.1","Developer Score Deterioration", "C", "DeveloperScoringService + RiskEngineService", "P3"],
    ["8.1","High-Risk Project Concentration", "C", "RiskEngineService aggregation", "P3"],
    ["8.1","Zone Price Decline", "C", "RiskEngineService + PriceIndex", "P3"],
    ["8.1","Severe Regulatory Violation", "C", "RegulatoryViolation + RiskEngineService", "P3"],
    ["8.1","Configurable thresholds with audit logging", "C", "RiskThreshold + Admin/EwrsDashboard.razor + AuditLogService", "P3"],
    ["8.2","Multi-level escalation L1–L4 with channel matrix", "C", "AlertEscalationService + AlertLevel + NotificationRecipientResolver", "P3"],
    ["8.3","Risk overview dashboard with heatmap and KPI bar", "C", "Admin/EwrsDashboard.razor + Map integration", "P3"],
    ["8.3","Alert inbox with status, owner, timestamps, notes", "C", "Admin/EwrsAlerts.razor + RiskAlert entity", "P3"],
    ["8.3","Playbook integration with checklist & audit trail", "C", "Admin/Playbooks.razor", "P3"],
    ["8.4","ESC-001 Real-time Escrow Dashboard with adequacy badges", "C", "Admin/EscrowMonitoring.razor + Admin/EscrowDetail.razor", "P3"],
    ["8.4","ESC-002 Immutable Escrow audit log; PDF export; L3 alert on discrepancy", "C", "EscrowTransaction (append-only) + AlertEscalationService", "P3"],
    ["8.4","ESC-003 Monthly auto-generated Escrow Health Report", "C", "EscrowHealthReportService (Hangfire)", "P3"],
])

matrix_section("C.8 §9 — Developer Performance & Rating", [
    ["9.1",     "Transparent, configurable composite developer score", "C", "DeveloperScoringService + ScoringWeight", "P3"],
    ["9.1.1",   "On-Time Project Delivery Rate", "C", "DeveloperScoringService + Project", "P3"],
    ["9.1.1",   "Unit Sales Completion Rate", "C", "DeveloperScoringService + Transaction", "P3"],
    ["9.1.1",   "Escrow Account Health Score", "C", "DeveloperScoringService + EscrowAccount", "P3"],
    ["9.1.1",   "Regulatory Compliance Record", "C", "DeveloperScoringService + RegulatoryViolation", "P3"],
    ["9.1.2",   "Developer rating UI with comparison & drill-down", "C", "Admin/DeveloperRating.razor + Admin/DeveloperComparison.razor", "P3"],
    ["9 / 20",  "Public-facing developer scorecards", "C", "Public/DeveloperScorecards.razor", "P3"],
])

matrix_section("C.9 §10 — Non-Functional, Security & RBAC", [
    ["10.1",   "P95 chart render ≤ 2 s", "C", "KpiSnapshotCache + EF AsNoTracking + Blazor pre-render", "P1"],
    ["10.1",   "AI Agent response ≤ 15 s", "C", "AIOrchestrator timeout + warm fallback", "P1"],
    ["10.1",   "Data freshness ≤ 15 min", "C", "BenchmarkRefreshService + KpiSnapshotRefreshService (Hangfire)", "P1"],
    ["10.2",   "Authentication: JWT bearer + federated OIDC (UAE Pass / Azure AD / Keycloak)", "C", "WebAPI + AdminAPI Program.cs (dual JwtBearer schemes)", "P1"],
    ["10.2",   "Refresh-token rotation", "C", "RefreshToken entity + auth flow", "P1"],
    ["10.2.1", "DESC ISR v3 alignment (audit, secrets, HTTPS, headers, rate limit)", "C", "AuditLogService + configuration + middleware + rate limiter", "P3"],
    ["10.2.2", "Security testing — VAPT, SAST, dependency scan", "CD", "OWASP ZAP + Roslyn analysers + vuln package in CI", "P3"],
    ["10.3",   "RBAC matrix: InternalRead/Edit/Manage/Admin + Investor", "C", "Infrastructure/Identity/AuthorizationPolicies.cs + UserRoles", "P1"],
    ["10.3",   "Resource-based authorization (investor sees own data only)", "E", "Ownership IAuthorizationHandler (vendor commitment)", "P1"],
    ["10",     "CAPTCHA on registration / password reset", "C", "CaptchaController + SimpleCaptchaService", "P1"],
    ["10",     "Tiered API-key rate limiting (Free/Plus/Partner)", "C", "PartitionedRateLimiter in WebAPI/Program.cs + ApiKey entity", "P1"],
])

matrix_section("C.10 §11 — Technology & Hosting", [
    ["11.1",   "Vendor technology stack proposal", "C", ".NET 9 + Blazor Server + EF Core + MediatR + Hangfire + Serilog + Redis", "P1"],
    ["11.2",   "Hosting model options (Azure UAE / on-prem / hybrid)", "CD", "All three costed in Section 12 + Annex E", "P0"],
    ["11.3",   "TLS, network segregation, WAF, private endpoints, BCDR", "CD", "Hosting deployment design + Bicep IaC", "P1"],
    ["11.4",   "Lakehouse data architecture (medallion bronze→silver→gold)", "CD", "OneLake + EF Core gold-layer reads + 15-min refresh jobs", "P1"],
    ["11.4.1", "Data architecture deliverables", "CD", "Discovery sprint outputs (lineage, schema, reconciliation)", "P1"],
    ["11.4.2", "Visualisation layer for external public portal", "C", "Blazor Server + Chart.js/ApexCharts (RTL-aware) + AnalyticsEmbed", "P1"],
])

matrix_section("C.11 §12 — Data Familiarisation", [
    ["12.1","DLD Data Familiarisation Discovery Sprint", "CD", "2-week sprint led by Data Architect", "P1"],
    ["12.2","Pre-publication analytics assessment", "CD", "Great Expectations / dbt-tests suite", "P1"],
    ["12.3","Data accuracy obligations during warranty", "CD", "Treated as Sev 1 with 4 h response SLA", "P1+"],
])

matrix_section("C.12 §13 — Phased Delivery Plan", [
    ["13","Phase 1 — External Portal + AI Agent + Open Data", "C", "Repository foundations in place", "P1"],
    ["13","Phase 2 — Watchlist v2, ESG, GRETI, performance", "CD", "Backlog scoped", "P2"],
    ["13","Phase 3 — EWRS, Developer Rating, Escrow, BO, Mortgage", "CD", "Foundations in repository", "P3"],
    ["13","Phase 4 — Multilingual extension + advanced analytics", "CD", "i18n + AI multi-language", "P4"],
])

matrix_section("C.13 §19 — Data Governance & Privacy", [
    ["19.1","Data ownership — all data and artefacts vest in DLD", "CD", "Contractual; vendor retains no rights", "P0"],
    ["19.2","Personal data protection (UAE PDPL)", "C", "UserConsentService + minimum-data design + erasure workflow", "P1"],
    ["19.3","Data quality standards (completeness/accuracy/timeliness/consistency)", "CD", "Quality dashboard + monthly review", "P1+"],
    ["19","Audit log of personal data reads with reason code", "C", "AuditLogService (justified-access pattern)", "P1"],
])

matrix_section("C.14 §20 — Additional Recommended Features (Beyond Brief)", [
    ["20","GRETI Sub-Index Tracker", "E", "GretiController + Public/Greti.razor", "P1"],
    ["20","ESG indicators", "E", "EsgController + Public/Esg.razor", "P1"],
    ["20","Beneficial Ownership transparency", "E", "BeneficialOwner entity + Public/BeneficialOwnership.razor", "P3"],
    ["20","Open Data API + self-service Developer Portal", "E", "OpenDataController + ApiKey + Public/ApiPortal.razor", "P1"],
    ["20","Mortgage analytics", "E", "MortgageController + Public/Mortgage.razor", "P3"],
    ["20","Investor Scorecard PDF", "E", "InvestorScorecardPdfService", "P1"],
    ["20","Headless CMS for non-technical content updates", "E", "CmsContent + CmsContentVersion", "P1"],
    ["20","Name validation service to reduce duplicate records", "E", "NameValidationController + NameValidation entity", "P1"],
])

matrix_section("C.15 Vendor Commitments (Value-Add)", [
    ["VC-1",  "MediatR ValidationBehavior",     "E", "ValidationBehavior<TRequest,TResponse>", "P1"],
    ["VC-2",  "Global IExceptionHandler",        "E", "Middleware + IExceptionHandler", "P1"],
    ["VC-3",  "Distributed Redis cache",         "E", "MediatR caching behavior + Redis", "P1"],
    ["VC-4",  "Domain event dispatcher",         "E", "IDomainEventDispatcher", "P1"],
    ["VC-5",  "Notification outbox pattern",     "E", "NotificationOutbox + Hangfire dispatcher", "P1"],
    ["VC-6",  "SignalR notification hub",        "E", "Already shipped", "P1"],
    ["VC-7",  "AI streaming over SSE",           "E", "AiAgent.razor + IAsyncEnumerable<string>", "P1"],
    ["VC-8",  "Polly resilience",                "E", "AddStandardResilienceHandler()", "P1"],
    ["VC-9",  "Hangfire idempotency",            "E", "Job filter", "P1"],
    ["VC-10", "OpenTelemetry",                   "E", "AddOpenTelemetry()", "P1"],
    ["VC-11", "/healthz + /readyz",              "E", "AspNetCore.HealthChecks", "P1"],
    ["VC-12", "Tighter CORS + CSP + HSTS",       "E", "Security middleware", "P1"],
    ["VC-13", "Architecture tests",              "E", "NetArchTest layer assertions", "P1"],
    ["VC-14", "Dockerfiles + CI/CD pipelines",   "E", "Multi-stage builds + GitHub Actions / Azure Pipelines", "P1"],
    ["VC-15", "Resource-based authorization",    "E", "IAuthorizationHandler for ownership", "P1"],
])

page_break(doc)

# =====================================================================
# ANNEX D — DESC ISR v3 35-CONTROL MAPPING
# =====================================================================
add_heading(doc, "Annex D — DESC ISR v3 Control Mapping", 1)
add_para(doc, "Clause-by-clause mapping of DESC Information Security Regulation v3 controls to the IRETP platform's implementation. Serves as the vendor's pre-VAPT self-assessment and supports the DESC compliance submission required before Phase 3 acceptance (RFP §10.2.1).")
add_para(doc, "Status definitions:", bold=True, space_after=2)
add_table(doc, ["Status", "Definition"], [
    ["I",  "Implemented — control is active in the codebase / hosting layer today."],
    ["PI", "Partially Implemented — core control in place; process item documented for completion in the operational runbook (Phase 1)."],
    ["P",  "Planned — control will be fully in place before Phase 3 VAPT. Owner and due date tracked in the project risk register."],
], col_widths_in=[0.7, 6.3])

def desc_section(title, rows):
    add_heading(doc, title, 3)
    add_table(doc, ["ID", "Control", "Status", "Implementation", "Evidence"],
              rows, col_widths_in=[0.5, 1.8, 0.5, 2.7, 1.5])

desc_section("D.1 Governance", [
    ["1.1", "Information Security Policy", "I",
     "Security requirements documented in §7 of this proposal. Policies enforced via AuthorizationPolicies, AuditLogService and rate limiter.",
     "AuthorizationPolicies.cs / AuditLogService.cs"],
    ["1.2", "Roles & Responsibilities", "I",
     "RBAC matrix (UserRoles + AuthorizationPolicies) assigns least-privilege roles: DldViewer, DldOperator, DldSupervisor, SystemAdministrator, RegisteredInvestor.",
     "AuthorizationPolicies.cs / UserRole.cs"],
    ["1.3", "Risk Assessment Process", "I",
     "Risks identified during design; EWRS surfaces risk posture in real-time. Annual risk review in support contract.",
     "RiskEngineService.cs"],
    ["1.4", "Security Exception Management", "PI",
     "Exceptions managed via AuditLog with timestamp + approver. Formal exception register to be formalised in operational runbook.",
     "AuditLogService.cs"],
    ["1.5", "Third-party & Vendor Risk", "I",
     "AI provider UAE-residency enforced at infrastructure level; all third-party libraries scanned (dotnet list package --vulnerable in CI); contractual security clauses with SMS/SMTP providers.",
     "WebAPI/Program.cs / AdminAPI/Program.cs"],
])
desc_section("D.2 Asset Management", [
    ["2.1", "Asset Inventory", "I",
     "Solution comprises 3 hosts (WebAPI, AdminAPI, Web), 1 SQL DB, 1 Redis, 1 OneLake workspace, 1 AI endpoint, 1 SMTP relay, 1 SMS gateway. IaC manifests enumerate all assets.",
     "infra/bicep/platform.bicep"],
    ["2.2", "Data Classification", "I",
     "Public (price/rental index, transactions), Internal (risk scores, escrow, audit logs), Confidential (beneficial ownership, personal investor data). Entities tagged accordingly.",
     "Domain/Entities/* XML comments"],
    ["2.3", "Media & Asset Disposal", "PI",
     "DB purge procedures + right-to-erasure workflow via account deletion endpoint; formal media-destruction policy for on-prem option in runbook.",
     "AccountController.cs"],
])
desc_section("D.3 HR Security", [
    ["3.1", "Pre-employment Screening", "P",
     "Vendor staff background-check process documented in vendor qualification dossier (RFP §16). Completed before project start.",
     "Vendor HR policy"],
    ["3.2", "Security Awareness Training", "P",
     "Mandatory security awareness training for all project staff and DLD platform administrators delivered in project kickoff.",
     "Training schedule (P1 M1)"],
    ["3.3", "Disciplinary Process", "I",
     "Every privileged action is recorded in AuditLog with identity + timestamp, enabling accountability and disciplinary investigation.",
     "AuditLogService.cs"],
])
desc_section("D.4 Physical", [
    ["4.1", "Secure Areas", "I",
     "Option A/C: Azure UAE North datacentre (ISO 27001). Option B: DLD-approved DC physical controls. Administrative access via Azure Portal with MFA only.",
     "Hosting design document"],
    ["4.2", "Equipment Security", "I",
     "All app tiers run in managed services (PaaS); no bare-metal managed by vendor. Developer workstations protected by full-disk encryption + Defender for Endpoint.",
     "Vendor workstation policy"],
])
desc_section("D.5 Operations", [
    ["5.1", "Change Management", "I",
     "All changes to production deployed via CI/CD pipelines (GitHub Actions / Azure Pipelines) with peer-reviewed PRs and staged environments. No direct production access.",
     "CI/CD pipeline definitions"],
    ["5.2", "Capacity Management", "I",
     "Azure App Service auto-scale rules and Redis Premium tier elasticity. Capacity reviewed against OpenTelemetry metrics monthly.",
     "WebAPI/AdminAPI App Service config"],
    ["5.3", "Malware Protection", "I",
     "Microsoft Defender for Cloud enabled on all Azure resources. Dependency scanning in CI. Container images scanned before push.",
     "Azure Defender / CI pipeline"],
    ["5.4", "Backup & Recovery", "I",
     "SQL geo-redundant backups (35-day retention); Redis persistence (AOF); OneLake recycle-bin policy. RPO ≤ 4 h; RTO ≤ 8 h.",
     "Hosting design / Azure Backup"],
    ["5.5", "Logging & Monitoring", "I",
     "Serilog → Seq / Azure Monitor. OpenTelemetry traces + metrics to OTLP. All privileged events in AuditLog. Alerts on anomalies.",
     "AuditLogService.cs / OpenTelemetry config"],
    ["5.6", "Vulnerability Management", "I",
     "OWASP ZAP scripted scans on every deployment to staging. DESC-authorised VAPT before P3. Monthly dependency scan report.",
     "CI/CD pipeline / VAPT schedule"],
])
desc_section("D.6 Access Control", [
    ["6.1", "Access Control Policy", "I",
     "Least-privilege RBAC via AuthorizationPolicies. No shared accounts; each user has a named ApplicationUser identity.",
     "AuthorizationPolicies.cs"],
    ["6.2", "User Registration & Provisioning", "I",
     "User registration via Auth/Register.razor with email verification. Role assignment by SystemAdministrator only via UserAdminController.",
     "UserAdminController.cs / Register.razor"],
    ["6.3", "Privileged Access Management", "I",
     "SystemAdmin policy restricts to SystemAdministrator role; Hangfire dashboard read-only for admin; AdminAPI not reachable from public internet.",
     "AdminAPI network config"],
    ["6.4", "Authentication", "I",
     "JWT bearer (HS256, zero clock-skew) + federated OIDC. Refresh-token rotation. CAPTCHA on registration and password reset.",
     "WebAPI/Program.cs / CaptchaController.cs"],
    ["6.5", "Session Management", "I",
     "Short-lived access tokens; sliding refresh tokens with family rotation; logout revokes the token family.",
     "RefreshToken / AuthController.cs"],
    ["6.6", "Network Access Control", "I",
     "AdminAPI bound to private VNET (not reachable from public internet). WebAPI behind WAF (OWASP). TLS 1.2+ enforced. CORS allow-list (not wildcard in prod).",
     "WebAPI/Program.cs / hosting network config"],
    ["6.7", "Remote Access", "PI",
     "Developer access to production via Azure Bastion / jump-box with MFA; no direct SSH from developer laptops. Formal remote-access policy in runbook.",
     "Hosting design / operational runbook (P1)"],
])
desc_section("D.7 Cryptography", [
    ["7.1", "Cryptographic Controls", "I",
     "JWT HMAC-SHA256; passwords hashed with ASP.NET Core Identity (PBKDF2 + salt); HMAC-SHA256 unsubscribe tokens; TLS 1.2+; Azure Key Vault for secrets.",
     "WebAPI/Program.cs / HmacUnsubscribeTokenService.cs"],
    ["7.2", "Key Management", "I",
     "JWT signing key from Azure Key Vault (IOptionsMonitor); key rotation without redeployment. Secrets never in appsettings.json in production.",
     "WebAPI/Program.cs / Key Vault binding"],
])
desc_section("D.8 Incident Management", [
    ["8.1", "Incident Response Procedure", "I",
     "EWRS L1–L4 escalation mirrors incident severity; Sev 1/2 alerting via AlertEscalationService. Admin/Playbooks.razor provides DLD step-by-step SOPs.",
     "AlertEscalationService.cs / Admin/Playbooks.razor"],
    ["8.2", "Reporting & Learning", "I",
     "AuditLog records every security event; monthly support report includes incident register, RCA summaries and recommended improvements.",
     "AuditLogService.cs / support SLA §15.3"],
    ["8.3", "Evidence Preservation", "I",
     "AuditLog is append-only; EscrowTransaction is immutable. Logs retained 12 months on the sink.",
     "AuditLogService.cs / EscrowTransaction entity"],
])
desc_section("D.9 Business Continuity", [
    ["9.1", "BCM Planning", "I",
     "RTO ≤ 8 h, RPO ≤ 4 h. Azure cross-AZ replicas for SQL and Redis. OneLake geo-redundancy. App Service slot-swap for zero downtime releases.",
     "Hosting design / Azure SLA documentation"],
    ["9.2", "DR Testing", "P",
     "Annual DR drill scheduled in Q4 of each operating year. Results reported to DLD in the monthly support pack.",
     "DR test plan (with operational runbook)"],
])
desc_section("D.10 Compliance", [
    ["10.1", "UAE Legal & Regulatory Compliance", "I",
     "UAE Federal PDPL honoured through UserConsentService (opt-in/out), minimum-data design, erasure workflow. All data and AI inference in UAE.",
     "UserConsentService.cs / hosting design"],
    ["10.2", "Intellectual Property", "I",
     "All platform IP vests in DLD on payment (RFP §19.1). Open-source components are permissively licensed. Licence SBOM generated by CI.",
     "Licence SBOM / Financial Proposal §19"],
    ["10.3", "Privacy & Personal Data Protection", "I",
     "BeneficialOwner and investor personal data tagged. Every read by an internal user logged with reason code. Right-to-erasure exposed on AccountController.",
     "UserConsentService.cs / AccountController.cs / AuditLogService.cs"],
    ["10.4", "DESC Compliance Reporting", "P",
     "DESC ISR v3 compliance self-assessment submitted with proposal (this annex). Annual external audit scheduled in operating year.",
     "This document / DESC audit schedule"],
])

add_heading(doc, "D.11 Residual Risks & Remediation Plan", 3)
add_table(doc, ["Risk ID", "Description", "State", "Likelihood", "Remediation"], [
    ["R-1", "Partially Implemented controls", "3 controls (1.4, 2.3, 6.7) marked PI", "Medium", "Operational runbook + formal process completed by Phase 1 M1"],
    ["R-2", "Planned controls pre-VAPT",      "3 controls (3.1, 3.2, 9.2) marked P",  "Low",    "Completed before Phase 3 VAPT; tracked in risk register"],
    ["R-3", "AI model provider UAE residency","Depends on provider UAE endpoint",     "Medium", "Provider contractually bound to UAE inference; UAE-hosted OSS fallback (Llama/Mistral)"],
], col_widths_in=[0.6, 2.2, 1.5, 0.8, 2.2])

page_break(doc)

# =====================================================================
# ANNEX E — HARDWARE SIZING
# =====================================================================
add_heading(doc, "Annex E — Hardware Sizing & Hosting Options", 1)
add_para(doc, "Production sizing per hosting option, UAE-resident. Application & data tier (steady-state):", space_after=4)
add_table(doc, ["Component", "Role", "vCPU", "RAM", "Storage", "Inst.", "Azure SKU"], [
    ["IRETP.Web",      "Blazor Server portal",              "4", "16 GB", "64 GB SSD",  "2 (HA)", "D4s_v5"],
    ["IRETP.WebAPI",   "Public + Investor API",             "4", "16 GB", "64 GB SSD",  "2 (HA)", "D4s_v5"],
    ["IRETP.AdminAPI", "Internal DLD operations",            "2", "8 GB",  "64 GB SSD",  "2 (HA)", "D2s_v5"],
    ["Hangfire worker","EWRS · alerts · KPI refresh",       "4", "16 GB", "128 GB SSD", "2 (HA)", "D4s_v5"],
    ["SQL Server",     "Primary OLTP",                       "8", "64 GB", "1 TB SSD",   "1+repl", "MI BC 8 vCore"],
    ["OneLake / Blob", "Lakehouse + exports + CMS",         "—", "—",     "5 TB GRS",   "—",      "Storage v2 GRS"],
    ["Key Vault",      "Secrets · MSI · rotation",           "—", "—",     "—",          "1",      "Premium HSM"],
    ["Front Door + WAF","Edge · OWASP · DDoS",              "—", "—",     "—",          "Global", "Premium"],
    ["Redis Cache",    "Distributed caching",                "—", "8 GB",  "—",          "1",      "Premium P1"],
], col_widths_in=[1.3, 1.6, 0.5, 0.6, 0.95, 0.7, 1.0])

add_heading(doc, "E.1 Total Footprint by Hosting Option", 2)
add_table(doc, ["Option", "Sizing", "Year-0 CAPEX", "Annual OPEX"], [
    ["Option A — Azure UAE North", "Managed PaaS · MSI · GRS backup · App: P1v3 (~12 vCPU / 48 GB) · Data: MI BC 8 vCore · 5 TB storage", "AED 335,500",  "AED 629,200"],
    ["Option B — On-Premises",     "11 prod servers + 11 DR + 4 UAT · K8s + open-source observability · 26 total servers · UPS · F5 WAF", "AED 1,414,500","AED 686,200"],
    ["Option C — Hybrid",          "Compute on-prem + lakehouse in Azure UAE + S2S VPN · 13 servers (compute + DR)",                  "AED 793,500",  "AED 654,200"],
], col_widths_in=[1.4, 3.0, 1.1, 1.1])

page_break(doc)

# =====================================================================
# APPENDIX A — WORKING SYSTEM EVIDENCE (with screenshots)
# =====================================================================
add_heading(doc, "Appendix A — Working System Evidence", 1)
add_para(doc, "The IRETP reference implementation is available for live review by the DLD evaluation panel at any point during the evaluation period. The screenshots in this appendix are rendered from the live system. Build clean (0/0), 88/88 xUnit tests pass.")

add_heading(doc, "A.1 Build and Test Status", 3)
add_bullets(doc, [
    "Build: dotnet build IRETP.sln → 0 warnings, 0 errors.",
    "Tests: 88 of 88 xUnit pass — domain rules, application handlers, infrastructure services, AI guardrail, EWRS engine, SLA computation, audit-log immutability, share-token expiry, notification SLA probe.",
    "Static analysis: CodeQL clean; dotnet list package --vulnerable clean.",
    "Security: OWASP ZAP baseline clean on the public surface.",
])

add_heading(doc, "A.2 Live Page Screenshots", 3)

add_figure(doc, os.path.join(SCREENS, '01_dashboard.png'),
    "A.2.1 — Public Dashboard with KPI cards, market charts, freshness badge.")

add_figure(doc, os.path.join(SCREENS, '02_map.png'),
    "A.2.2 — Interactive GIS Map with 4 heatmap modes and individual project pins.")

add_figure(doc, os.path.join(SCREENS, '03_transactions.png'),
    "A.2.3 — Transactions Explorer with URL-persisted filters and CSV/PDF export.")

add_figure(doc, os.path.join(SCREENS, '04_ai_agent.png'),
    "A.2.4 — Real Estate AI Agent with citation chips, EN/AR language picker, deterministic stats appended.")

add_figure(doc, os.path.join(SCREENS, '05_ewrs.png'),
    "A.2.5 — EWRS Alert Inbox with L1–L4 escalation ladder and immutable threshold-edit audit.")

add_figure(doc, os.path.join(SCREENS, '06_scorecards.png'),
    "A.2.6 — Public Developer Scorecards with 8-criterion composite score.")

add_figure(doc, os.path.join(SCREENS, '07_beneficial_ownership.png'),
    "A.2.7 — Beneficial Ownership with English + Arabic owner names rendered RTL.")

add_figure(doc, os.path.join(SCREENS, '08_mortgage.png'),
    "A.2.8 — Mortgage page with 5 KPIs including Average Mortgage Size.")

add_figure(doc, os.path.join(SCREENS, '09_audit_logs.png'),
    "A.2.9 — Audit Logs view (Admin) — append-only at DbContext, 12-month retention.")

page_break(doc)

# =====================================================================
# APPENDIX B — RFP §16.2
# =====================================================================
add_heading(doc, "Appendix B — Mandatory Proposal Contents Compliance", 1)
add_para(doc, "The following checklist confirms compliance with Section 16.2 of the RFP. All items are included in or attached to this proposal.")
add_table(doc, ["RFP §16.2 Requirement", "Status & Location"], [
    ["Confidentiality Statement",                  "Included — front matter."],
    ["Executive Summary",                          "Included — Executive Summary section."],
    ["Understanding of the Requirement",           "Section 1."],
    ["Proposed Solution Architecture",             "Section 3."],
    ["AI Agent Architecture & Design",             "Section 4."],
    ["Internal Platform Design (EWRS, Developer Performance)", "Section 5."],
    ["Phased Delivery Plan",                       "Section 6 + Annex B."],
    ["Security & DESC Compliance",                 "Section 7 + Annex D (35-control mapping)."],
    ["Visualisation Strategy",                     "Section 8."],
    ["Data Understanding & Accuracy Methodology",  "Section 9."],
    ["Data Integration Reference",                 "Section 10."],
    ["Project Team & Responsibility Matrix",       "Section 11."],
    ["Cost Structure (3 hosting options costed)",  "Section 12 + Annex E."],
    ["Warranty & Technical Support",               "Section 13."],
    ["AI Knowledge Update Approach",               "Section 14."],
    ["Risk Register",                              "Section 15."],
    ["Knowledge Transfer & Exit Plan",             "Section 16."],
    ["Beyond-Brief / GRETI Recommendations",       "Section 17."],
    ["Vendor Commitments (Value-Add)",             "Section 18."],
    ["Quality Posture & Test Strategy",            "Section 19."],
    ["System Architecture Diagram",                "Annex A."],
    ["Phased Delivery Gantt Chart",                "Annex B."],
    ["Functional Requirements Compliance Matrix",  "Annex C (130+ rows, full roll-up)."],
    ["DESC ISR v3 Control Mapping",                "Annex D (35 controls, 10 domains)."],
    ["Hardware Sizing & Hosting Options",          "Annex E."],
    ["Working System / POC Evidence",              "Appendix A (with 9 live screenshots)."],
], col_widths_in=[3.6, 3.4])

page_break(doc)

# =====================================================================
# APPENDIX C — CLA
# =====================================================================
add_heading(doc, "Appendix C — Compliance List Acknowledgement (CLA)", 1)
add_para(doc, "Per-section compliance status against RFP DLD-IRETP-2026-001. Legend: FC = Full Compliance · PC = Partial Compliance · NC = Non-Compliance.")
add_table(doc, ["§", "Section Title", "# Reqs", "Status", "Evidence / Notes"], [
    ["§3",  "External Public Portal",          "12", "FC", "Blazor Server pages live · CMS · bilingual · KPI dashboard"],
    ["§4",  "Slice-and-Dice Analytics",        "9",  "FC", "Saved views (cap 12) · 4 FR-004 charts · Zone Compare (≤5)"],
    ["§5",  "Real Estate AI Agent",            "11", "FC", "UAE-resident · multi-model tier routing · advisory guardrail · 14-case harness"],
    ["§6",  "Investor Notifications",          "8",  "FC", "Email / SMS / in-platform · SLA probe · HMAC unsubscribe · List-Unsubscribe"],
    ["§7",  "Multilingual EN/AR + Currency",  "6",  "FC", "RTL parity · Phase-4 ZH/RU/UR/FR/HI/DE roadmap"],
    ["§8",  "Early Warning Risk System",       "10", "FC", "10 indicators · 4-level auto-escalation · threshold-driven · admin override"],
    ["§9",  "Developer Rating & Escrow",       "8",  "FC", "Public scorecards · §8.4 escrow monitor · weighted composite"],
    ["§10", "NFR, Security & RBAC",            "11", "FC", "P95 ≤ 2 s · RBAC matrix · DESC ISR · CAPTCHA · tiered rate limit"],
    ["§11", "Technology & Hosting",            "6",  "FC", "Three hosting options costed · TLS · WAF · private endpoints · BCDR"],
    ["§12", "Data Familiarisation",            "3",  "FC", "2-week discovery sprint · Pre-Pub Assessment · accuracy SLA"],
    ["§13", "Phased Delivery",                  "4",  "FC", "90-day phased plan with go-live gates per phase"],
    ["§14", "Cost Structure",                   "4",  "FC", "CAPEX + OPEX broken down per RFP format; 6 milestone schedule"],
    ["§15", "Warranty & Support",               "4",  "FC", "12-month warranty; Sev 1–4 SLA matrix; monthly reporting"],
    ["§16", "Vendor Qualifications",            "2",  "FC", "Compliance package complete; qualifications dossier attached"],
    ["§17", "Acceptance Criteria",              "3",  "FC", "Phase-level UAT package; data accuracy gate; universal UAT bar"],
    ["§18", "Proposal Evaluation",              "3",  "FC", "Technical + Financial submissions complete; clarification available"],
    ["§19", "Data Governance & Privacy",        "4",  "FC", "IP vests in DLD · UAE PDPL · audit log + reason code"],
    ["§20", "Beyond-Brief (Recommended)",       "8",  "FC", "All 8 features in repository today (E status)"],
    ["VC",  "Vendor Commitments",               "15", "FC", "All 15 committed in P1 hardening at no extra cost (E status)"],
], col_widths_in=[0.5, 2.3, 0.8, 0.7, 2.7])

add_heading(doc, "Acceptance Statement", 2)
add_para(doc, "The vendor confirms that every requirement of RFP DLD-IRETP-2026-001 is addressed by this proposal. Items marked C are demonstrable today against the IRETP codebase. Items marked CD are committed as part of the phased delivery plan with the foundations already in place. Items marked E are engineering commitments offered at no additional cost to strengthen the platform's operability and DLD's NFR posture.")
add_para(doc, "Submitted with respect to the Dubai Land Department.", italic=True, color=COL_MUTED, space_before=6)

# save
doc.save(OUT)
print(f"WROTE {OUT}")
