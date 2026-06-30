import docx, os, sys
files = [
    r'C:\Users\kalmi\Downloads\TechnicalProposal_IRETP.docx',
    r'C:\Users\kalmi\Downloads\FinancialProposal_IRETP.docx',
    r'C:\Users\kalmi\Downloads\ComplianceMatrix_IRETP.docx',
    r'C:\Users\kalmi\Downloads\AppendixA_DESC_ISR_IRETP.docx',
]
out = []
for f in files:
    d = docx.Document(f)
    name = os.path.basename(f)
    out.append(f'### {name}: {len(d.paragraphs)} paras, {len(d.tables)} tables')
    for p in d.paragraphs:
        s = p.style.name if p.style else ''
        t = p.text.strip()
        if t and ('Head' in s or 'Title' in s):
            lvl = s.replace('Heading ', 'H').replace('Title', 'T')
            out.append(f'  {lvl}: {t[:130]}')
    out.append('')
open(r'C:\Users\kalmi\IRETP\_attachments_outline.txt', 'w', encoding='utf-8').write('\n'.join(out))
print('OK')
