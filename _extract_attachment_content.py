import docx, os, sys

files = [
    r'C:\Users\kalmi\Downloads\TechnicalProposal_IRETP.docx',
    r'C:\Users\kalmi\Downloads\FinancialProposal_IRETP.docx',
    r'C:\Users\kalmi\Downloads\ComplianceMatrix_IRETP.docx',
    r'C:\Users\kalmi\Downloads\AppendixA_DESC_ISR_IRETP.docx',
]
for f in files:
    d = docx.Document(f)
    name = os.path.basename(f).replace('.docx','')
    out = []
    out.append(f'================== {name} ==================')
    # interleave paragraphs and tables in document order
    body = d.element.body
    # python-docx doesn't give an interleaved iterator easily; we'll just dump all paragraphs then tables.
    for p in d.paragraphs:
        t = p.text.strip()
        if not t: continue
        style = p.style.name if p.style else ''
        if 'Heading' in style or 'Title' in style:
            lvl = style.replace('Heading ', 'H').replace('Title', 'T')
            out.append(f'\n## {lvl}: {t}')
        else:
            out.append(t)
    out.append('\n--- TABLES ---')
    for i, t in enumerate(d.tables):
        out.append(f'\n[T{i}] ({len(t.rows)}r x {len(t.columns)}c)')
        for r in t.rows:
            cells = [c.text.replace('\n', ' ').strip() for c in r.cells]
            out.append(' | '.join(cells))
    open(rf'C:\Users\kalmi\IRETP\_att_{name}.txt', 'w', encoding='utf-8').write('\n'.join(out))
    print(f'WROTE _att_{name}.txt')
